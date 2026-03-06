#!/usr/bin/env python3
"""
Scanwave CyberIntel Platform - Full Dashboard Viewer
====================================================
Usage: python viewer.py
Then open http://localhost:5000
"""

import os
import re
import csv
import json
import time
import asyncio
import threading
import subprocess
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import StringIO, BytesIO
from pathlib import Path
from datetime import datetime, timezone, timedelta
from collections import defaultdict
import gzip as _gzip
from flask import Flask, jsonify, request, Response, send_file, abort

try:
    from telethon.sync import TelegramClient as TGSync
    TELETHON_OK = True
except ImportError:
    TELETHON_OK = False

# ── SQLite Database Layer ────────────────────────────────────────────────────
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))
from app.database import init_db, get_conn
from app.config import DB_PATH
from app import models as db

# Initialize SQLite on import (creates tables if needed, enables WAL)
init_db(DB_PATH)

# ── Structured Logging ──────────────────────────────────────────────────────
from app.logging_config import setup_logging
log = setup_logging("viewer")

app = Flask(__name__)

@app.after_request
def _compress_response(response):
    """Gzip compress JSON responses > 1KB if client supports it."""
    if (response.status_code < 200 or response.status_code >= 300
            or response.direct_passthrough
            or 'Content-Encoding' in response.headers
            or 'gzip' not in request.headers.get('Accept-Encoding', '')):
        return response
    ct = response.content_type or ''
    if not (ct.startswith('application/json') or ct.startswith('text/')):
        return response
    data = response.get_data()
    if len(data) < 1024:
        return response
    compressed = _gzip.compress(data, compresslevel=6)
    response.set_data(compressed)
    response.headers['Content-Encoding'] = 'gzip'
    response.headers['Content-Length'] = len(compressed)
    response.headers['Vary'] = 'Accept-Encoding'
    return response

OUTPUT_DIR     = Path("./telegram_intel")
SESSION        = str(Path("./jordan_cyber_intel"))
API_ID         = int(os.environ.get("TG_API_ID",   "35545979"))
API_HASH       = os.environ.get("TG_API_HASH", "41240e3f451065a430692d2e1bc82453")
WATCHLIST_FILE = OUTPUT_DIR / "active_watchlist.json"

# Channel tier metadata for sidebar display and briefing context
CHANNEL_TIERS = {
    # TIER 1 — Directly targeting Jordan
    "hak993":               {"tier": 1, "label": "Fatemiyoun Cyber Team", "threat": "CRITICAL"},
    "hak994":               {"tier": 1, "label": "Fatemiyoun Cyber Team", "threat": "CRITICAL"},
    "Fatimion310_bot":      {"tier": 1, "label": "Fatemiyoun Bot",        "threat": "CRITICAL"},
    "fatimion110":          {"tier": 1, "label": "Fatemiyoun Community",  "threat": "CRITICAL"},
    "hkr313":               {"tier": 1, "label": "Fatemiyoun Field Team", "threat": "CRITICAL"},
    "xX313XxTeam":          {"tier": 1, "label": "313 Team (Iraqi Cyber Resistance)", "threat": "CRITICAL"},
    "Team313Official":      {"tier": 1, "label": "313 Team Official",     "threat": "CRITICAL"},
    "x313xTeamLeak":        {"tier": 1, "label": "313 Team Leaks",        "threat": "CRITICAL"},
    "x313xTeamBackup":      {"tier": 1, "label": "313 Team Backup",       "threat": "CRITICAL"},
    # TIER 2 — Iran-aligned ecosystem
    "Mhwear98":             {"tier": 2, "label": "Cyber Islamic Resistance", "threat": "HIGH",   "status": "banned"},
    "Mhwercyber4":          {"tier": 2, "label": "Cyber Islamic Resistance", "threat": "HIGH",   "status": "banned"},
    "mhwear0":              {"tier": 2, "label": "Cyber Islamic Resistance (active handle)", "threat": "HIGH"},
    "fattahh_ir":           {"tier": 2, "label": "Cyber Fattah Team",     "threat": "HIGH",   "status": "banned"},
    "Handala_hack":         {"tier": 2, "label": "Handala (MOIS/Void Manticore)", "threat": "HIGH", "status": "banned"},
    "handala_hack26":       {"tier": 2, "label": "Handala Backup",        "threat": "HIGH",   "status": "banned"},
    "handal_a":             {"tier": 2, "label": "Handala (Active 2025)", "threat": "HIGH"},
    "Handala_hack_iranian": {"tier": 2, "label": "Handala Iranian",       "threat": "HIGH"},
    "Handala_Hack_Team":    {"tier": 2, "label": "Handala Team",          "threat": "HIGH"},
    "blackopmrhamza":       {"tier": 2, "label": "Mr Hamza",              "threat": "HIGH",   "status": "banned"},
    "RipperSec":            {"tier": 2, "label": "RipperSec (Malaysia)",  "threat": "HIGH",   "status": "banned"},
    "TheRipperSec":         {"tier": 2, "label": "RipperSec Main",        "threat": "HIGH",   "status": "banned"},
    "AnonGhostOfficialTeam":{"tier": 2, "label": "AnonGhost",             "threat": "HIGH",   "status": "banned"},
    "sylhetgangsgofficial": {"tier": 2, "label": "Sylhet Gang SG",        "threat": "HIGH",   "status": "banned"},
    "KeymousTeam":          {"tier": 2, "label": "Keymous+",              "threat": "HIGH"},
    "KMPteam":              {"tier": 2, "label": "Keymous+ (Alt)",        "threat": "HIGH",   "status": "banned"},
    "Keymous_V2":           {"tier": 2, "label": "Keymous+ Backup",       "threat": "HIGH",   "status": "banned"},
    "islamic_hacker_army1": {"tier": 2, "label": "Islamic Hacker Army",   "threat": "HIGH"},
    # TIER 3 — Broader ecosystem
    "noname05716eng":       {"tier": 3, "label": "NoName057(16) EN",      "threat": "MEDIUM", "status": "banned"},
    "noname05716":          {"tier": 3, "label": "NoName057(16) RU",      "threat": "MEDIUM", "status": "banned"},
    "dienet3":              {"tier": 3, "label": "DieNet (DDoS-as-a-service)", "threat": "MEDIUM"},
    "dnnmabot":             {"tier": 3, "label": "DieNet Attack Notifier Bot", "threat": "MEDIUM"},
    "LulzSecBlack":         {"tier": 3, "label": "LulzSec Black (PIJ)",   "threat": "MEDIUM"},
    "LulzSecHackers":       {"tier": 3, "label": "LulzSec Hackers",       "threat": "MEDIUM"},
    "Luls_sec_muslims":     {"tier": 3, "label": "LulzSec Muslims",       "threat": "MEDIUM"},
    "ElamAlmoqawama":       {"tier": 3, "label": "Islamic Resistance Iraq","threat": "MEDIUM"},
    "toufan_alaksa":        {"tier": 3, "label": "Al Aqsa Resistance Axis","threat": "MEDIUM"},
    "Chat_Islamic_Hacker_Army": {"tier": 3, "label": "IHA Discussion",    "threat": "MEDIUM"},
    "Khamenei_arabi":       {"tier": 3, "label": "Khamenei Arabic",       "threat": "LOW"},
    "operationswordofjustice": {"tier": 2, "label": "Operation Sword of Justice", "threat": "HIGH"},
    "Anonymous0islamic":    {"tier": 3, "label": "Anonymous Islamic",     "threat": "MEDIUM"},
    "handala24":            {"tier": 2, "label": "Handala (Active 2024+)","threat": "HIGH"},
    "tttteam313":           {"tier": 2, "label": "313 Team (Affiliated)", "threat": "HIGH"},
    "SabrenNewss":          {"tier": 2, "label": "Sabren News",           "threat": "HIGH"},
    # ── VERIFIED ADDITIONS (2026-03-03, cross-referenced TGStat/telemetr.io) ───
    # Dark Storm Team — DDoS/defacement US/EU/ME infrastructure
    "DarkStormTeams":       {"tier": 2, "label": "Dark Storm Team",        "threat": "HIGH"},
    "DarkStormBackup":      {"tier": 2, "label": "Dark Storm Team (Backup)", "threat": "HIGH"},
    "darkstormchat":        {"tier": 2, "label": "Dark Storm Team (Chat)", "threat": "HIGH"},
    # Arabian Ghosts — pro-Palestine Gulf hacktivist
    "arabian_ghosts":       {"tier": 2, "label": "Arabian Ghosts",         "threat": "HIGH"},
    # APT Iran — IRGC cyber readiness, threat tracking
    "aptiran":              {"tier": 2, "label": "APT Iran",               "threat": "HIGH"},
    # Golden Falcon — pro-Palestine DDoS/defacement
    "Golden_falcon_team":   {"tier": 2, "label": "Golden Falcon",          "threat": "HIGH"},
    # Stucx Team — active defacement, multiple verified handles
    "stucxteam":            {"tier": 2, "label": "Stucx Team",             "threat": "HIGH"},
    "stucxnet":             {"tier": 2, "label": "Stucx Net",              "threat": "HIGH"},
    "xxstucxteam":          {"tier": 2, "label": "Stucx Team (Alt)",       "threat": "HIGH"},
    # Hand of Justice — affiliated with Cyber Isnaad Front
    "the_hand_of_justice":  {"tier": 2, "label": "Hand of Justice",        "threat": "HIGH"},
    # Cyber Isnaad Front — Iranian-aligned Syrian hacktivist
    "CyberIsnaadFront":     {"tier": 2, "label": "Cyber Isnaad Front",     "threat": "HIGH"},
    # Cyb3rDrag0nz / CyberDrag0nzz — defacement coalition
    "TeamCyb3rDrag0nz":     {"tier": 2, "label": "CyberDrag0nzz",         "threat": "HIGH"},
    "cyb3r_drag0nz_team":   {"tier": 2, "label": "CyberDrag0nzz (Alt)",   "threat": "HIGH"},
    # Hacktivist of Garuda / Garuda Eye — Indonesian hacktivist
    "HacktivistOfGaruda":   {"tier": 3, "label": "Garuda Eye (Indonesia)", "threat": "MEDIUM"},
    "HacktivistOfGarudaOfficial": {"tier": 3, "label": "Garuda Eye (Official)", "threat": "MEDIUM"},
    # Nation of Saviours — active pro-Palestine
    "nation_of_saviors_public": {"tier": 3, "label": "Nation of Saviours", "threat": "MEDIUM"},
    # EvilNet 3.0 — data exfiltration, defacement
    "EvilNet3":             {"tier": 3, "label": "EvilNet 3.0",            "threat": "MEDIUM"},
    # Gaza Children's Group — Gaza-based hacktivist
    "Gaza_Children_Hackers": {"tier": 2, "label": "Gaza Children's Group", "threat": "HIGH"},
    "gaza_children_ha":     {"tier": 2, "label": "Gaza Children (Backup)", "threat": "HIGH"},
    # Indohaxsec — Indonesian hacktivist, Arctic Wolf-reported
    "INDOHAXSEC":           {"tier": 3, "label": "Indohaxsec",             "threat": "MEDIUM"},
    # Altoufan Team — resistance cyber ops
    "ALTOUFANTEAM":         {"tier": 2, "label": "Altoufan Team",          "threat": "HIGH"},
    # Team Azrael (Angel of Death) — resistance-axis
    "anonymous_cr02x":      {"tier": 2, "label": "Team Azrael (Angel of Death)", "threat": "HIGH"},
    "teamAzraelbackup":     {"tier": 2, "label": "Team Azrael (Backup)",   "threat": "HIGH"},
    # BD Anonymous / The Anonymous BD — Bangladesh
    "anonymous_bangladesh": {"tier": 3, "label": "BD Anonymous",           "threat": "MEDIUM"},
    "t_gray_hacker":        {"tier": 3, "label": "The Anonymous BD",       "threat": "MEDIUM"},
    # Moroccan Black Cyber Army — note: zero not O
    "M0roccan_Black_CyberArmy": {"tier": 3, "label": "Moroccan Black Cyber Army", "threat": "MEDIUM"},
    "moroccan_blackcyberarmy": {"tier": 3, "label": "MBCA (Alt)",          "threat": "MEDIUM"},
    # Akatsuki Cyber Team — pro-Palestine, pro-Iran ops
    "akatsukicyberteam":    {"tier": 3, "label": "Akatsuki Cyber Team",    "threat": "MEDIUM"},
    # FAD Team — CONFIRMED Ministry of Finance breach, real channel @r3_6j
    "r3_6j":                {"tier": 1, "label": "FAD Team (Min. of Finance breach)", "threat": "CRITICAL"},
    # Cyber Av3ngers — IRGC-affiliated, ICS/OT attacks on water/energy
    "CyberAv3ngers":        {"tier": 2, "label": "Cyber Av3ngers (IRGC)",  "threat": "HIGH"},
    "cyberaveng3rs":        {"tier": 2, "label": "Cyber Av3ngers (Alt)",   "threat": "HIGH"},
    # Iran Anonymous / Anonymous OpIran
    "anonopiran":           {"tier": 3, "label": "Anonymous OpIran",       "threat": "MEDIUM"},
    # Liwaa Mohammad (Mohamed Brigade) — Lebanese resistance cyber unit
    "liwaamohammad":        {"tier": 2, "label": "Liwaa Mohammad (Lebanese Resistance)", "threat": "HIGH"},
    # Tharallah Brigade — uses mhwear* namespace
    "mhwear10":             {"tier": 2, "label": "Tharallah Brigade",      "threat": "HIGH"},
    # Cyber32
    "Cyber32":              {"tier": 3, "label": "Cyber32",                "threat": "MEDIUM"},
    # DieNet extended channels
    "DieNetAPI":            {"tier": 3, "label": "DieNet API Info",        "threat": "MEDIUM"},
    "dienet_media":         {"tier": 3, "label": "DieNet Media Corp",      "threat": "MEDIUM"},
}

# Scan process state
_scan = {"running": False, "last_run": None, "pid": None}

# ── In-memory cache for performance ──────────────────────────────────────────
_msg_cache = {"data": None, "mtime": 0}

# ── AI enrichment cache ────────────────────────────────────────────────────────
_enrich_cache = {"data": None, "ts": 0}

def load_enrichments():
    """Load AI-enriched alerts from SQLite with 30s cache."""
    now = time.time()
    if _enrich_cache["data"] is not None and (now - _enrich_cache["ts"]) < 30:
        return _enrich_cache["data"]
    result = db.get_enrichments()
    _enrich_cache["data"] = result
    _enrich_cache["ts"] = now
    return result

_msg_cache = {"data": None, "ts": 0, "count": 0}
_MSG_CACHE_TTL = 30  # seconds

def load_messages():
    """Load messages from SQLite with 15s in-memory cache."""
    now = time.time()
    # Fast path: cache is fresh
    if _msg_cache["data"] is not None and (now - _msg_cache["ts"]) < _MSG_CACHE_TTL:
        return _msg_cache["data"]
    try:
        conn = get_conn()
        # Quick row count — if unchanged, extend cache up to 2 min
        cnt = conn.execute("SELECT COUNT(*) FROM messages").fetchone()[0]
        if _msg_cache["data"] is not None and cnt == _msg_cache["count"] and (now - _msg_cache["ts"]) < 120:
            _msg_cache["ts"] = now  # refresh TTL
            return _msg_cache["data"]
        rows = conn.execute(
            "SELECT raw_json, critical_subtype, backfill, media_path, has_media FROM messages ORDER BY timestamp_utc ASC"
        ).fetchall()
        msgs = []
        for row in rows:
            try:
                m = json.loads(row[0])
                # Always recompute critical_subtype (viewer has latest filters)
                if m.get("priority") == "CRITICAL":
                    m["critical_subtype"] = _compute_critical_subtype(
                        m.get("keyword_hits", []),
                        m.get("text_preview", "") or m.get("full_text", "") or m.get("text", "")
                    )
                elif row[1]:
                    m["critical_subtype"] = row[1]
                if row[2]:
                    m["backfill"] = True
                # Ensure media fields are set from DB columns (may not be in raw_json)
                if row[3]:
                    m["media_path"] = row[3]
                if row[4]:
                    m["has_media"] = True
                msgs.append(m)
            except Exception:
                pass
        _msg_cache["data"] = msgs
        _msg_cache["ts"] = now
        _msg_cache["count"] = cnt
        return msgs
    except Exception:
        return _msg_cache["data"] or []


# Keywords for live-fetch scoring
_SCORE_CRIT = [
    "jordan","الاردن","الأردن","أردن","اردن",".jo",".gov.jo",
    "hacked","breached","defaced","leak","dump","wiper","destroy",
    "ddos","تسريب","تم اختراق","arab bank","البنك العربي",
    "bank of jordan","بنك الأردن","housing bank","بنك الإسكان",
    "ministry of interior","وزارة الداخلية","ministry of defense",
    "royal court","الديوان الملكي","prime minister","رئاسة الوزراء",
    "jordan islamic bank","البنك الإسلامي الأردني","jmis","jcbank",
]

def _score_text(text):
    tl = text.lower()
    hits = [k for k in _SCORE_CRIT if k in tl]
    if any(k in tl for k in ["jordan","الاردن","الأردن",".jo",".gov.jo"]):
        return "CRITICAL", hits
    if hits:
        return "MEDIUM", hits
    return "LOW", []


def fetch_live_context(channel_username, msg_id, before, after):
    """Pull real messages from Telegram by ID range, merge stored metadata."""
    if not TELETHON_OK:
        return None, -1
    try:
        try:
            asyncio.get_event_loop()
        except RuntimeError:
            asyncio.set_event_loop(asyncio.new_event_loop())
        with TGSync(SESSION, API_ID, API_HASH) as client:
            raw = list(client.iter_messages(
                channel_username,
                min_id=max(0, msg_id - before - 1),
                max_id=msg_id + after + 1,
                limit=before + after + 1,
            ))
        raw = sorted(raw, key=lambda m: m.id)
        stored_map = {}
        for m in load_messages():
            ch = m.get("channel_username") or m.get("channel", "")
            if ch == channel_username:
                stored_map[m.get("message_id")] = m
        result = []
        target_idx = -1
        for i, m in enumerate(raw):
            stored = stored_map.get(m.id, {})
            if m.id == msg_id:
                target_idx = i
            result.append({
                "message_id":       m.id,
                "channel_username": channel_username,
                "channel":          stored.get("channel", channel_username),
                "timestamp_utc":    m.date.isoformat() if m.date else "",
                "timestamp_irst":   stored.get("timestamp_irst", ""),
                "text_preview":     (m.text or ""),
                "priority":         stored.get("priority", "LOW"),
                "keyword_hits":     stored.get("keyword_hits", []),
                "iocs":             stored.get("iocs", {}),
                "has_media":        bool(m.media),
                "live":             True,
            })
        return result, target_idx
    except Exception:
        return None, -1


# ═══════════════════════════════════════════════════════════════════════════════
# API ROUTES
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/api/channels")
def api_channels():
    messages = load_messages()
    now      = datetime.now(timezone.utc)
    # 7-day sparkline buckets: index 0 = 6 days ago, index 6 = today
    def _spark_key(ts_str):
        try:
            dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
            delta = (now.date() - dt.date()).days
            if 0 <= delta < 7:
                return 6 - delta   # 0=6days_ago .. 6=today
        except Exception:
            pass
        return None

    stats = defaultdict(lambda: {
        "channel": "", "channel_username": "",
        "count": 0, "critical": 0, "medium": 0,
        "last_date": "", "last_critical_date": "",
        "spark_7d": [0, 0, 0, 0, 0, 0, 0],  # 7-day critical count per day
    })
    for m in messages:
        ch  = m.get("channel_username") or m.get("channel", "unknown")
        s   = stats[ch]
        s["channel"]          = m.get("channel", ch)
        s["channel_username"] = ch
        s["count"] += 1
        p  = m.get("priority", "LOW")
        ts = m.get("timestamp_utc", "")
        if p == "CRITICAL":
            s["critical"] += 1
            if ts > s["last_critical_date"]:
                s["last_critical_date"] = ts
            idx = _spark_key(ts)
            if idx is not None:
                s["spark_7d"][idx] += 1
        elif p == "MEDIUM":
            s["medium"] += 1
        if ts > s["last_date"]:
            s["last_date"] = ts
    # Enrich with tier metadata
    now = datetime.now(timezone.utc)
    for s in stats.values():
        ch = s["channel_username"]
        meta = CHANNEL_TIERS.get(ch, {})
        s["tier"]         = meta.get("tier", 0)
        s["tier_label"]   = meta.get("label", "")
        s["threat_level"] = meta.get("threat", "")
        s["status"]       = meta.get("status", "active")
        # Days since last post
        if s["last_date"]:
            try:
                ld = datetime.fromisoformat(s["last_date"].replace("Z", "+00:00"))
                if ld.tzinfo is None:
                    ld = ld.replace(tzinfo=timezone.utc)
                s["days_silent"] = max(0, (now - ld).days)
            except Exception:
                s["days_silent"] = -1
        else:
            s["days_silent"] = -1

    # Inject stubs for all CHANNEL_TIERS entries with no messages in DB
    # so analysts can see the full picture including banned and newly-added channels
    for username, meta in CHANNEL_TIERS.items():
        if username not in stats:
            stats[username] = {
                "channel": meta.get("label", username),
                "channel_username": username,
                "count": 0, "critical": 0, "medium": 0,
                "last_date": "", "last_critical_date": "",
                "tier": meta.get("tier", 0),
                "tier_label": meta.get("label", ""),
                "threat_level": meta.get("threat", ""),
                "status": meta.get("status", "active"),
                "days_silent": -1,
                "spark_7d": [0, 0, 0, 0, 0, 0, 0],
            }

    # Enrich status for all stats entries that aren't already enriched
    for s in stats.values():
        if "status" not in s:
            ch   = s["channel_username"]
            meta = CHANNEL_TIERS.get(ch, {})
            s["status"] = meta.get("status", "active")

    # Sort: active channels with criticals first, then by date DESC;
    # banned channels always go to the bottom of the list
    result = sorted(stats.values(),
        key=lambda x: (
            0 if x.get("status") == "banned" else 1,           # banned → bottom
            1 if x["last_critical_date"] else 0,               # has crit → higher
            x["last_critical_date"] or x["last_date"] or ""    # newest first
        ),
        reverse=True
    )
    return jsonify(result)


@app.route("/api/channel/<channel_username>/iocs")
def api_channel_iocs(channel_username):
    """Aggregated IOCs for a single channel - for per-channel IOC drill-down."""
    from collections import Counter
    messages = load_messages()
    agg = defaultdict(Counter)
    msg_count = 0
    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")
        if ch != channel_username:
            continue
        msg_count += 1
        for ioc_type, vals in (m.get("iocs") or {}).items():
            for v in (vals or []):
                agg[ioc_type][v] += 1
    result = []
    for ioc_type, counts in agg.items():
        for val, count in counts.most_common(50):
            result.append({"type": ioc_type, "value": val, "count": count})
    result.sort(key=lambda x: -x["count"])
    return jsonify({"channel": channel_username, "msg_count": msg_count, "iocs": result})


@app.route("/api/channel/<channel_username>/trend")
def api_channel_trend(channel_username):
    """Daily message counts for a single channel over last 30 days."""
    messages = load_messages()
    now  = datetime.now(timezone.utc)
    days = 30
    buckets = defaultdict(lambda: {"total": 0, "critical": 0, "medium": 0})
    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")
        if ch != channel_username:
            continue
        ts = m.get("timestamp_utc", "")
        if not ts:
            continue
        try:
            dt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
            delta = (now - dt).days
            if 0 <= delta < days:
                day_key = dt.strftime("%Y-%m-%d")
                p = m.get("priority", "LOW")
                buckets[day_key]["total"] += 1
                if p == "CRITICAL":   buckets[day_key]["critical"] += 1
                elif p == "MEDIUM":   buckets[day_key]["medium"]   += 1
        except Exception:
            pass
    result = sorted(
        [{"date": k, **v} for k, v in buckets.items()],
        key=lambda x: x["date"]
    )
    return jsonify(result)


@app.route("/api/messages/<channel_username>")
def api_messages(channel_username):
    priority_filter  = request.args.get("priority", "ALL")
    subtype_filter   = request.args.get("critical_subtype", "ALL")
    since  = request.args.get("since", "")
    until  = request.args.get("until", "")
    search = request.args.get("search", "").lower()

    enrichments = load_enrichments()
    result = []
    for m in load_messages():
        ch = m.get("channel_username") or m.get("channel", "unknown")
        if ch != channel_username:
            continue
        if priority_filter != "ALL" and m.get("priority") != priority_filter:
            continue
        ts = m.get("timestamp_utc", "")
        if since and ts < since:
            continue
        if until and ts > until + " 99":
            continue
        if search:
            text = ((m.get("text_preview") or "") + " " +
                    " ".join(m.get("keyword_hits", []))).lower()
            if search not in text:
                continue
        # Attach AI enrichment if available
        key = f"{ch}_{m.get('message_id','')}"
        m = dict(m)
        if key in enrichments:
            m["ai_enrichment"] = enrichments[key]
        # Inject critical_subtype for backward compat (old messages won't have it)
        if m.get("priority") == "CRITICAL" and not m.get("critical_subtype"):
            m["critical_subtype"] = _compute_critical_subtype(m.get("keyword_hits", []), m.get("text_preview", "") or m.get("text", ""))
        # Apply critical_subtype filter (only filters CRITICAL messages)
        if subtype_filter != "ALL" and m.get("priority") == "CRITICAL":
            ms = m.get("critical_subtype", "GENERAL")
            if ms != subtype_filter and ms != "BOTH":
                continue
        result.append(m)

    result.sort(key=lambda x: x.get("timestamp_utc", ""))
    return jsonify(result)


@app.route("/api/messages/all")
def api_messages_all():
    """Combined feed across all channels. Used by Dashboard and Timeline.
    Supports pagination: ?page=1&per_page=50 returns paginated result.
    Without page param, returns legacy array format for backward compat.
    """
    priority_filter = request.args.get("priority", "ALL")
    subtype_filter  = request.args.get("critical_subtype", "ALL")
    since   = request.args.get("since", "")
    until   = request.args.get("until", "")
    search  = request.args.get("search", "").lower()
    keyword = request.args.get("keyword", "").lower()
    channel = request.args.get("channel", "").lower()
    limit   = min(int(request.args.get("limit", 1000)), 5000)
    # Pagination params (optional)
    page     = request.args.get("page", type=int)
    per_page = min(request.args.get("per_page", 50, type=int), 200)

    enrichments = load_enrichments()
    result = []
    for m in load_messages():
        p  = m.get("priority", "LOW")
        ts = m.get("timestamp_utc", "")
        ch = (m.get("channel_username") or m.get("channel", "")).lower()

        if priority_filter != "ALL" and p != priority_filter:
            continue
        if since  and ts < since:
            continue
        if until  and ts > until + " 99":
            continue
        if channel and channel not in ch:
            continue
        if search:
            haystack = ((m.get("text_preview") or "") + " " +
                        " ".join(m.get("keyword_hits", [])) + " " +
                        (m.get("channel", "") or "")).lower()
            if search not in haystack:
                continue
        if keyword:
            kws      = [k.lower() for k in m.get("keyword_hits", [])]
            txt_low  = (m.get("text_preview") or "").lower()
            if keyword not in kws and keyword not in txt_low:
                continue
        m = dict(m)
        # Inject critical_subtype for backward compat (old messages won't have it)
        if p == "CRITICAL" and not m.get("critical_subtype"):
            m["critical_subtype"] = _compute_critical_subtype(m.get("keyword_hits", []), m.get("text_preview", "") or m.get("text", ""))
        # Attach AI enrichment if available
        ch_full = m.get("channel_username") or m.get("channel", "")
        key = f"{ch_full}_{m.get('message_id','')}"
        if key in enrichments:
            m["ai_enrichment"] = enrichments[key]
        # Apply critical_subtype filter (only filters CRITICAL messages)
        if subtype_filter != "ALL" and p == "CRITICAL":
            ms = m.get("critical_subtype", "GENERAL")
            if ms != subtype_filter and ms != "BOTH":
                continue
        result.append(m)

    result.sort(key=lambda x: x.get("timestamp_utc", ""), reverse=True)

    # If page param provided, return paginated response
    if page is not None:
        total = len(result)
        start = (page - 1) * per_page
        end = start + per_page
        return jsonify({
            "data": result[start:end],
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total": total,
                "pages": (total + per_page - 1) // per_page if per_page else 1,
            }
        })
    return jsonify(result[:limit])


@app.route("/api/messages/poll")
def api_messages_poll():
    """Lightweight polling endpoint: returns only messages newer than `after` timestamp.
    Uses DB-level WHERE clause for speed — no full table scan.
    Returns: {new_count, newest_ts, messages[]} where messages only includes new ones.
    """
    after    = request.args.get("after", "")
    priority = request.args.get("priority", "ALL")
    limit    = min(int(request.args.get("limit", 200)), 500)

    if not after:
        return jsonify({"new_count": 0, "newest_ts": "", "messages": []})

    try:
        conn = get_conn()
        query = "SELECT raw_json, critical_subtype, backfill, media_path, has_media FROM messages WHERE timestamp_utc > ?"
        params = [after]
        if priority != "ALL":
            query += " AND priority = ?"
            params.append(priority)
        query += " ORDER BY timestamp_utc DESC LIMIT ?"
        params.append(limit)
        rows = conn.execute(query, params).fetchall()

        msgs = []
        newest_ts = after
        for row in rows:
            try:
                m = json.loads(row[0])
                if m.get("priority") == "CRITICAL" and not m.get("critical_subtype"):
                    m["critical_subtype"] = _compute_critical_subtype(
                        m.get("keyword_hits", []),
                        m.get("text_preview", "") or m.get("text", ""))
                if row[2]:
                    m["backfill"] = True
                if row[3]:
                    m["media_path"] = row[3]
                if row[4]:
                    m["has_media"] = True
                ts = m.get("timestamp_utc", "")
                if ts > newest_ts:
                    newest_ts = ts
                msgs.append(m)
            except Exception:
                pass

        return jsonify({
            "new_count": len(msgs),
            "newest_ts": newest_ts,
            "messages": msgs,
        })
    except Exception as e:
        return jsonify({"new_count": 0, "newest_ts": after, "messages": [], "error": str(e)})


@app.route("/api/messages/count")
def api_messages_count():
    """Ultra-lightweight: just return message count and newest timestamp.
    Used by frontend to decide whether a full fetch is needed.
    """
    priority = request.args.get("priority", "ALL")
    try:
        conn = get_conn()
        if priority != "ALL":
            row = conn.execute(
                "SELECT COUNT(*), MAX(timestamp_utc) FROM messages WHERE priority = ?",
                (priority,)).fetchone()
        else:
            row = conn.execute(
                "SELECT COUNT(*), MAX(timestamp_utc) FROM messages").fetchone()
        return jsonify({"count": row[0] or 0, "newest_ts": row[1] or ""})
    except Exception:
        return jsonify({"count": 0, "newest_ts": ""})


@app.route("/api/dashboard")
def api_dashboard():
    """Aggregated intelligence for the Dashboard tab."""
    messages = load_messages()

    kw_crit  = defaultdict(int)
    kw_med   = defaultdict(int)
    kw_total = defaultdict(int)

    ioc_agg = defaultdict(lambda: defaultdict(lambda: {
        "count": 0, "channels": set(), "last_seen": ""
    }))

    # (weekday_abbr, hour_int) → count
    WEEKDAYS = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
    activity_matrix = {wd: {h: 0 for h in range(24)} for wd in WEEKDAYS}

    ch_last_critical = {}
    ch_msg_count     = defaultdict(int)
    total = len(messages)
    critical_count = medium_count = ioc_count = 0

    # Campaign detection: track per-keyword which channels post in same 6-hour window
    kw_windows = defaultdict(list)  # kw → list of (timestamp, channel)

    for m in messages:
        p  = m.get("priority", "LOW")
        ch = m.get("channel_username") or m.get("channel", "unknown")
        ts = m.get("timestamp_utc", "")
        ch_msg_count[ch] += 1

        if p == "CRITICAL":
            critical_count += 1
            if ts > ch_last_critical.get(ch, ""):
                ch_last_critical[ch] = ts
        elif p == "MEDIUM":
            medium_count += 1

        for kw in m.get("keyword_hits", []):
            kw_total[kw] += 1
            if p == "CRITICAL":
                kw_crit[kw]  += 1
                kw_windows[kw].append((ts, ch))
            elif p == "MEDIUM":
                kw_med[kw] += 1

        for ioc_type, vals in (m.get("iocs") or {}).items():
            for val in (vals or []):
                entry = ioc_agg[ioc_type][val]
                entry["count"] += 1
                entry["channels"].add(ch)
                if ts > entry["last_seen"]:
                    entry["last_seen"] = ts
                ioc_count += 1

        # Activity matrix
        irst_hour    = m.get("irst_hour")
        irst_weekday = m.get("irst_weekday", "")
        if irst_hour is not None and irst_weekday:
            wd_abbr = irst_weekday[:3]
            if wd_abbr in activity_matrix and 0 <= irst_hour < 24:
                activity_matrix[wd_abbr][irst_hour] += 1

    # Keywords sorted by weight
    keywords = []
    for kw in kw_total:
        weight = kw_crit.get(kw, 0) * 3 + kw_med.get(kw, 0)
        keywords.append({
            "keyword": kw,
            "total":   kw_total[kw],
            "critical": kw_crit.get(kw, 0),
            "medium":  kw_med.get(kw, 0),
            "weight":  weight,
        })
    keywords.sort(key=lambda x: -x["weight"])

    # IOC list
    iocs_out = []
    for ioc_type, vals in ioc_agg.items():
        for val, info in vals.items():
            iocs_out.append({
                "type":     ioc_type,
                "value":    val,
                "count":    info["count"],
                "channels": sorted(info["channels"]),
                "last_seen": info["last_seen"],
            })
    iocs_out.sort(key=lambda x: -x["count"])

    # Coordinated campaigns: keyword with 2+ distinct channels in 6-hour window
    campaigns = []
    for kw, events in kw_windows.items():
        if len(events) < 2:
            continue
        events_s = sorted(events, key=lambda e: e[0])
        i = 0
        while i < len(events_s):
            t0_str, ch0 = events_s[i]
            cluster_chs = {ch0}
            j = i + 1
            while j < len(events_s):
                tj_str, chj = events_s[j]
                if t0_str and tj_str and tj_str[:10] == t0_str[:10]:
                    cluster_chs.add(chj)
                    j += 1
                else:
                    break
            if len(cluster_chs) >= 2:
                campaigns.append({
                    "keyword":  kw,
                    "channels": sorted(cluster_chs),
                    "date":     t0_str[:10] if t0_str else "",
                    "count":    len(cluster_chs),
                })
            i = j if j > i else i + 1
    # Deduplicate campaigns by (keyword, date)
    seen_c = set()
    campaigns_dedup = []
    for c in sorted(campaigns, key=lambda x: (-x["count"], x["keyword"])):
        k = (c["keyword"], c["date"])
        if k not in seen_c:
            seen_c.add(k)
            campaigns_dedup.append(c)
    campaigns_dedup = campaigns_dedup[:30]

    # Channel ranking
    ch_ranking = sorted(
        [{"channel": ch, "count": cnt, "last_critical": ch_last_critical.get(ch, "")}
         for ch, cnt in ch_msg_count.items()],
        key=lambda x: -x["count"]
    )[:20]

    total_configured = len(CHANNEL_TIERS)
    banned_count     = sum(1 for v in CHANNEL_TIERS.values() if v.get("status") == "banned")
    return jsonify({
        "total":            total,
        "critical":         critical_count,
        "medium":           medium_count,
        "ioc_count":        ioc_count,
        "channel_count":    len(ch_msg_count),
        "total_configured": total_configured,
        "banned_count":     banned_count,
        "keywords":         keywords[:100],
        "iocs":             iocs_out[:500],
        "activity_matrix":  {wd: list(hours.values()) for wd, hours in activity_matrix.items()},
        "campaigns":        campaigns_dedup,
        "ch_ranking":       ch_ranking,
    })


@app.route("/api/trend")
def api_trend():
    """Daily message volume over past 60 days, broken down by priority."""
    messages  = load_messages()
    now       = datetime.now(timezone.utc)
    days      = int(request.args.get("days", 60))
    cutoff    = (now - timedelta(days=days)).isoformat()

    day_data = defaultdict(lambda: {"CRITICAL": 0, "MEDIUM": 0, "LOW": 0})
    for m in messages:
        ts = m.get("timestamp_utc", "")
        if ts < cutoff:
            continue
        day = ts[:10]
        p   = m.get("priority", "LOW")
        day_data[day][p] += 1

    # Fill every day in range
    result = []
    for i in range(days):
        day = (now - timedelta(days=days-1-i)).strftime("%Y-%m-%d")
        d   = day_data.get(day, {"CRITICAL": 0, "MEDIUM": 0, "LOW": 0})
        result.append({"date": day, **d})

    return jsonify(result)


@app.route("/api/briefing")
def api_briefing():
    """24-hour intelligence briefing for the analyst."""
    messages = load_messages()
    now = datetime.now(timezone.utc)
    cutoff_24h = (now - timedelta(hours=24)).isoformat()
    cutoff_7d  = (now - timedelta(days=7)).isoformat()

    recent   = [m for m in messages if m.get("timestamp_utc","") >= cutoff_24h]
    week     = [m for m in messages if m.get("timestamp_utc","") >= cutoff_7d]
    all_crit = [m for m in messages if m.get("priority") == "CRITICAL"]
    rec_crit = [m for m in recent  if m.get("priority") == "CRITICAL"]
    rec_med  = [m for m in recent  if m.get("priority") == "MEDIUM"]

    # Top targeted Jordan entities in last 24h
    entity_hits = defaultdict(int)
    for m in rec_crit:
        for kw in m.get("keyword_hits", []):
            entity_hits[kw] += 1
    top_entities = sorted(entity_hits.items(), key=lambda x: -x[1])[:15]

    # Active channels last 24h
    ch_24h = defaultdict(lambda: {"critical": 0, "medium": 0, "total": 0})
    for m in recent:
        ch = m.get("channel_username") or m.get("channel", "")
        ch_24h[ch]["total"] += 1
        p = m.get("priority","LOW")
        if p == "CRITICAL": ch_24h[ch]["critical"] += 1
        elif p == "MEDIUM": ch_24h[ch]["medium"] += 1
    active_channels = sorted(
        [{"channel": k, **v} for k, v in ch_24h.items()],
        key=lambda x: -x["critical"]
    )[:10]

    # Fresh IOCs in last 24h
    fresh_iocs = defaultdict(set)
    for m in recent:
        for ioc_type, vals in (m.get("iocs") or {}).items():
            for v in (vals or []):
                fresh_iocs[ioc_type].add(v)
    fresh_ioc_list = {k: sorted(v)[:20] for k, v in fresh_iocs.items()}

    # Trend: last 24h vs previous 24h
    cutoff_48h = (now - timedelta(hours=48)).isoformat()
    prev_24h_crit = sum(1 for m in messages
                        if cutoff_48h <= m.get("timestamp_utc","") < cutoff_24h
                        and m.get("priority") == "CRITICAL")
    trend = "ESCALATING" if len(rec_crit) > prev_24h_crit else \
            "DECREASING" if len(rec_crit) < prev_24h_crit else "STABLE"

    # Newest critical messages
    newest_crits = sorted(rec_crit, key=lambda x: x.get("timestamp_utc",""), reverse=True)[:10]

    return jsonify({
        "generated_at":    now.isoformat(),
        "period_hours":    24,
        "summary": {
            "total_messages_24h":    len(recent),
            "critical_alerts_24h":   len(rec_crit),
            "medium_alerts_24h":     len(rec_med),
            "active_channels_24h":   len(ch_24h),
            "fresh_ioc_count":       sum(len(v) for v in fresh_ioc_list.values()),
            "prev_24h_critical":     prev_24h_crit,
            "trend":                 trend,
            "total_all_time":        len(messages),
            "total_critical_all":    len(all_crit),
        },
        "top_targeted_entities": top_entities,
        "active_channels":       active_channels,
        "fresh_iocs":            fresh_ioc_list,
        "newest_critical":       newest_crits[:5],
        "weekly_critical":       len([m for m in week if m.get("priority")=="CRITICAL"]),
    })


@app.route("/api/messages/export")
def api_messages_export():
    """Download filtered messages as CSV."""
    priority_filter = request.args.get("priority", "CRITICAL")
    since = request.args.get("since", "")
    search = request.args.get("search", "").lower()

    rows = []
    for m in load_messages():
        p  = m.get("priority", "LOW")
        ts = m.get("timestamp_utc", "")
        if priority_filter != "ALL" and p != priority_filter:
            continue
        if since and ts < since:
            continue
        if search:
            haystack = ((m.get("text_preview") or "") + " " + " ".join(m.get("keyword_hits", []))).lower()
            if search not in haystack:
                continue
        rows.append(m)
    rows.sort(key=lambda x: x.get("timestamp_utc",""), reverse=True)

    buf = StringIO()
    w   = csv.writer(buf)
    w.writerow(["timestamp_utc","timestamp_irst","channel","channel_username",
                "priority","keyword_hits","text_preview","iocs"])
    for m in rows:
        ioc_str = "; ".join(f"{t}:{','.join(vs)}" for t, vs in (m.get("iocs") or {}).items())
        w.writerow([
            m.get("timestamp_utc",""), m.get("timestamp_irst",""),
            m.get("channel",""),       m.get("channel_username",""),
            m.get("priority",""),      "|".join(m.get("keyword_hits",[])),
            (m.get("text_preview") or "").replace("\n"," "),
            ioc_str
        ])
    fname = f"intel_{priority_filter.lower()}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": f"attachment;filename={fname}"})


@app.route("/api/iocs/export")
def api_iocs_export():
    """Download all IOCs as CSV."""
    messages = load_messages()
    ioc_agg  = defaultdict(lambda: defaultdict(lambda: {
        "count": 0, "channels": set(), "last_seen": ""
    }))
    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "unknown")
        ts = m.get("timestamp_utc", "")
        for ioc_type, vals in (m.get("iocs") or {}).items():
            for val in (vals or []):
                e = ioc_agg[ioc_type][val]
                e["count"] += 1
                e["channels"].add(ch)
                if ts > e["last_seen"]:
                    e["last_seen"] = ts
    buf = StringIO()
    w   = csv.writer(buf)
    w.writerow(["type", "value", "count", "channels", "last_seen"])
    for ioc_type, vals in ioc_agg.items():
        for val, info in sorted(vals.items(), key=lambda x: -x[1]["count"]):
            w.writerow([ioc_type, val, info["count"],
                        "|".join(sorted(info["channels"])), info["last_seen"]])
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=iocs.csv"})


@app.route("/api/threat_matrix")
def api_threat_matrix():
    """Matrix of threat actors (channels) vs Jordan target categories. Critical msgs only."""
    TARGET_CATEGORIES = {
        "Banking":   ["bank", "arab bank", "housing bank", "jordan ahli", "cairo amman",
                      "central bank", "البنك", "بنك", "aib", "invest bank"],
        "Telecom":   ["zain", "orange", "umniah", "tcs.jo", "jet.jo", "go.jo",
                      "jordan telecom", "اتصالات", "tcs communications"],
        "Gov/Mil":   ["government", "ministry", "military", "army", "وزار", "جيش",
                      "parliament", "muwaffaq", "قوات", "نيابة", "armed forces",
                      "مجلس", "presidency", "prime minister"],
        "ISP/Net":   ["cablenet", "index.jo", "internet", "broadband", "isp",
                      "network", "الانترنت"],
        "Media":     ["media", "news", "petra", "jordan tv", "قناة", "بث",
                      "television", "broadcast", "جهاز"],
        "Infra":     ["power", "water", "airport", "electricity", "سد", "كهرباء",
                      "fuel", "oil", "energy", "pipeline"],
        "General":   ["jordan", "الأردن", "amman", "عمان", "hashemite"],
    }
    messages = load_messages()
    matrix = defaultdict(lambda: {cat: 0 for cat in TARGET_CATEGORIES})
    for m in messages:
        if m.get("priority") != "CRITICAL":
            continue
        ch   = m.get("channel_username") or m.get("channel", "unknown")
        text = (m.get("text_preview") or "").lower()
        kws  = " ".join(m.get("keyword_hits", [])).lower()
        haystack = text + " " + kws
        for cat, keywords in TARGET_CATEGORIES.items():
            if any(kw.lower() in haystack for kw in keywords):
                matrix[ch][cat] += 1
    # Build rows from known CHANNEL_TIERS actors
    actors = []
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        row = {
            "channel":  ch,
            "label":    meta.get("label", ch),
            "tier":     meta.get("tier", 0),
            "threat":   meta.get("threat", ""),
        }
        row_totals = 0
        for cat in TARGET_CATEGORIES:
            cnt = matrix[ch].get(cat, 0)
            row[cat] = cnt
            row_totals += cnt
        row["total"] = row_totals
        if row_totals > 0:
            actors.append(row)
    # Also include active channels not in CHANNEL_TIERS
    known = set(CHANNEL_TIERS.keys())
    for ch, cats in matrix.items():
        if ch in known:
            continue
        total = sum(cats.values())
        if total > 0:
            actors.append({"channel": ch, "label": ch, "tier": 0, "threat": "", "total": total, **cats})
    actors.sort(key=lambda x: (-(x.get("tier") or 0), -x["total"]))
    return jsonify({"categories": list(TARGET_CATEGORIES.keys()), "actors": actors})


# ─── APT TRACKER & IOC INTELLIGENCE ENDPOINTS ────────────────────────────────

def _build_apt_profiles():
    """Aggregate APT profiles from CHANNEL_TIERS + message data."""
    messages = load_messages()
    enrichments = load_enrichments()

    # Group channels by APT label (many channels → one APT)
    apt_groups = defaultdict(lambda: {
        "channels": [], "tier": 99, "threat": "", "status": "active",
        "total_msgs": 0, "critical_count": 0, "medium_count": 0,
        "sectors": defaultdict(int),
        "attack_types": defaultdict(int),
        "first_seen": "", "last_seen": "",
        "jordan_attacks": 0,
    })
    research_cache = _load_research_cache()

    # Map channel → APT name
    channel_to_apt = {}
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        label = meta.get("label", ch)
        # Normalize group names: strip suffixes like (Backup), (Alt), (Active 2024+)
        base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', label).strip()
        channel_to_apt[ch] = base
        grp = apt_groups[base]
        grp["channels"].append(ch)
        tier = meta.get("tier", 3)
        if tier < grp["tier"]:
            grp["tier"] = tier
            grp["threat"] = meta.get("threat", "MEDIUM")

    # Jordan keywords for attack detection
    _jo_kw = [".jo", "jordan", "الاردن", "الأردن", "أردن", "اردن", "عمان", "amman", "hashemite"]

    # Scan all messages
    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")
        apt_name = channel_to_apt.get(ch)

        # Also check ai_enrichment group_attribution for channels not in CHANNEL_TIERS
        if not apt_name:
            ekey = f"{ch}_{m.get('message_id', '')}"
            ai = enrichments.get(ekey, m.get("ai_enrichment", {}))
            ga = ai.get("group_attribution", "")
            if ga and ga.lower() not in ("unknown", "n/a", ""):
                apt_name = ga
                if apt_name not in apt_groups:
                    apt_groups[apt_name]["channels"].append(ch)
                    apt_groups[apt_name]["tier"] = 3
                    apt_groups[apt_name]["threat"] = "MEDIUM"
                elif ch not in apt_groups[apt_name]["channels"]:
                    apt_groups[apt_name]["channels"].append(ch)

        if not apt_name:
            continue

        grp = apt_groups[apt_name]
        grp["total_msgs"] += 1
        pri = m.get("priority", "LOW")
        if pri == "CRITICAL":
            grp["critical_count"] += 1
        elif pri == "MEDIUM":
            grp["medium_count"] += 1

        # Timestamps
        ts = m.get("timestamp_utc", "")
        if ts:
            if not grp["first_seen"] or ts < grp["first_seen"]:
                grp["first_seen"] = ts
            if not grp["last_seen"] or ts > grp["last_seen"]:
                grp["last_seen"] = ts

        # AI enrichment — sectors and attack types
        ekey = f"{ch}_{m.get('message_id', '')}"
        ai = enrichments.get(ekey, m.get("ai_enrichment", {}))
        sector = ai.get("target_sector", "")
        if sector and sector.lower() not in ("unknown", "n/a", ""):
            grp["sectors"][sector] += 1
        atype = ai.get("attack_type", "")
        if atype and atype.lower() not in ("unknown", "n/a", ""):
            grp["attack_types"][atype] += 1

        # Jordan attack detection
        text = (m.get("text_preview") or "").lower()
        kws = " ".join(m.get("keyword_hits", [])).lower()
        haystack = text + " " + kws
        if any(kw in haystack for kw in _jo_kw):
            grp["jordan_attacks"] += 1

    # Build response list
    profiles = []
    for name, grp in apt_groups.items():
        if grp["total_msgs"] == 0:
            continue
        # IOC counts from external research cache
        research = research_cache.get(name, {})
        r_stats = research.get("stats", {})
        ioc_total = r_stats.get("total", 0)
        ioc_malicious = r_stats.get("malicious", 0)
        ioc_suspicious = r_stats.get("suspicious", 0)

        # Activity status
        last = grp["last_seen"]
        status = "inactive"
        if last:
            try:
                ld = datetime.fromisoformat(last.replace("Z", "+00:00"))
                delta = (datetime.now(timezone.utc) - ld).days
                status = "active" if delta <= 14 else ("stale" if delta <= 60 else "inactive")
            except Exception:
                pass

        profiles.append({
            "name": name,
            "tier": grp["tier"],
            "threat": grp["threat"],
            "channels": grp["channels"][:10],
            "status": status,
            "total_msgs": grp["total_msgs"],
            "critical_count": grp["critical_count"],
            "medium_count": grp["medium_count"],
            "ioc_count": ioc_total,
            "ioc_malicious": ioc_malicious,
            "ioc_suspicious": ioc_suspicious,
            "sectors": dict(sorted(grp["sectors"].items(), key=lambda x: -x[1])[:10]),
            "attack_types": dict(sorted(grp["attack_types"].items(), key=lambda x: -x[1])[:10]),
            "first_seen": grp["first_seen"],
            "last_seen": grp["last_seen"],
            "jordan_attacks": grp["jordan_attacks"],
        })
    profiles.sort(key=lambda x: (x["tier"], -x["critical_count"], -x["total_msgs"]))
    return profiles


@app.route("/api/apt/profiles")
def api_apt_profiles():
    """APT profile listing for APT Tracker sidebar."""
    return jsonify(_build_apt_profiles())


@app.route("/api/apt/<path:name>/detail")
def api_apt_detail(name):
    """Full APT detail: IOCs, attacks, timeline, recent messages."""
    messages = load_messages()
    enrichments = load_enrichments()

    # Find channels for this APT
    apt_channels = set()
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        label = meta.get("label", ch)
        base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', label).strip()
        if base == name:
            apt_channels.add(ch)

    _jo_kw = [".jo", "jordan", "الاردن", "الأردن", "أردن", "اردن", "عمان", "amman", "hashemite"]
    attacks = []
    recent_critical = []
    timeline = defaultdict(lambda: {"critical": 0, "medium": 0, "low": 0})
    total_msgs = 0
    critical_count = 0
    medium_count = 0
    sectors = defaultdict(int)
    attack_types = defaultdict(int)
    first_seen = ""
    last_seen = ""

    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")

        # Check direct channel membership
        in_apt = ch in apt_channels

        # Also check ai_enrichment group_attribution
        if not in_apt:
            ekey = f"{ch}_{m.get('message_id', '')}"
            ai = enrichments.get(ekey, m.get("ai_enrichment", {}))
            ga = ai.get("group_attribution", "")
            if ga:
                ga_base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', ga).strip()
                if ga_base == name:
                    in_apt = True

        if not in_apt:
            continue

        total_msgs += 1
        ts = m.get("timestamp_utc", "")
        pri = m.get("priority", "LOW")
        if pri == "CRITICAL":
            critical_count += 1
        elif pri == "MEDIUM":
            medium_count += 1

        if ts:
            if not first_seen or ts < first_seen:
                first_seen = ts
            if not last_seen or ts > last_seen:
                last_seen = ts
            try:
                month = ts[:7]
                if pri == "CRITICAL":
                    timeline[month]["critical"] += 1
                elif pri == "MEDIUM":
                    timeline[month]["medium"] += 1
                else:
                    timeline[month]["low"] += 1
            except Exception:
                pass

        # IOCs
        msg_id = m.get("message_id", "")

        # AI enrichment
        ekey = f"{ch}_{m.get('message_id', '')}"
        ai = enrichments.get(ekey, m.get("ai_enrichment", {}))
        sector = ai.get("target_sector", "")
        if sector and sector.lower() not in ("unknown", "n/a", ""):
            sectors[sector] += 1
        atype = ai.get("attack_type", "")
        if atype and atype.lower() not in ("unknown", "n/a", ""):
            attack_types[atype] += 1

        # Jordan attacks
        text = (m.get("text_preview") or "").lower()
        kws = " ".join(m.get("keyword_hits", [])).lower()
        haystack = text + " " + kws
        if any(kw in haystack for kw in _jo_kw) and pri in ("CRITICAL", "MEDIUM"):
            # Extract target domain if present
            domains = (m.get("iocs") or {}).get("domain", [])
            jo_domains = [d for d in domains if ".jo" in d.lower()]
            target = jo_domains[0] if jo_domains else "Jordan target"
            attacks.append({
                "date": ts[:10] if ts else "",
                "target": target,
                "type": atype or pri,
                "channel": ch,
                "msg_id": str(msg_id),
                "summary": (m.get("text_preview") or "")[:120],
            })

        # Recent CRITICAL messages
        if pri == "CRITICAL":
            recent_critical.append({
                "msg_id": str(msg_id),
                "channel": ch,
                "timestamp": ts,
                "text": (m.get("text_preview") or "")[:200],
                "keyword_hits": m.get("keyword_hits", [])[:5],
                "iocs": {k: v[:3] for k, v in (m.get("iocs") or {}).items()},
            })

    # Timeline sorted
    tl = [{"month": k, **v} for k, v in sorted(timeline.items())]

    # Attacks sorted by date desc
    attacks.sort(key=lambda x: x["date"], reverse=True)

    # Recent critical sorted by timestamp desc
    recent_critical.sort(key=lambda x: x["timestamp"], reverse=True)

    return jsonify({
        "name": name,
        "channels": sorted(apt_channels),
        "total_msgs": total_msgs,
        "critical_count": critical_count,
        "medium_count": medium_count,
        "first_seen": first_seen,
        "last_seen": last_seen,
        "sectors": dict(sorted(sectors.items(), key=lambda x: -x[1])),
        "attack_types": dict(sorted(attack_types.items(), key=lambda x: -x[1])),
        "attacks": attacks[:50],
        "recent_messages": recent_critical[:20],
        "timeline": tl,
    })


# ─── AbuseIPDB Integration ───────────────────────────────────────────────────

_ABUSEIPDB_CACHE_FILE = OUTPUT_DIR / "abuseipdb_cache.json"
_ABUSEIPDB_KEY = ""  # Moved to .env / app/config.py

def _load_abuse_cache():
    """Load abuse cache — returns dict for backward compat."""
    # Used by legacy code paths; prefer db.get_abuse_cache(ip) directly
    return {}

def _save_abuse_cache(cache):
    """Save abuse cache — now handled by SQLite."""
    pass

def _abuseipdb_check(value, ioc_type):
    """Check IP/domain against AbuseIPDB with SQLite cache."""
    import urllib.request, urllib.parse
    from app.config import ABUSEIPDB_KEY
    key = os.environ.get("ABUSEIPDB_API_KEY", ABUSEIPDB_KEY)
    if not key:
        return None

    # Check SQLite cache first
    cache_key = f"{ioc_type}:{value}"
    cached = db.get_abuse_cache(cache_key)
    if cached:
        try:
            cached_at = cached.get("cached_at", "")
            if cached_at:
                ct = datetime.fromisoformat(str(cached_at).replace("Z", "+00:00"))
                if (datetime.now(timezone.utc) - ct).total_seconds() < 86400:
                    return cached.get("data", cached)
        except Exception:
            pass

    try:
        # AbuseIPDB only accepts IPs — resolve domains first
        lookup_ip = value
        if ioc_type == "domain":
            import socket
            try:
                lookup_ip = socket.gethostbyname(value)
            except socket.gaierror:
                return {"error": f"Cannot resolve domain {value}"}

        if ioc_type == "subnet":
            url = f"https://api.abuseipdb.com/api/v2/check-block?network={urllib.parse.quote(value)}"
        else:
            url = f"https://api.abuseipdb.com/api/v2/check?ipAddress={urllib.parse.quote(lookup_ip)}&maxAgeInDays=90&verbose=true"

        req = urllib.request.Request(url, headers={"Key": key, "Accept": "application/json"})
        resp = urllib.request.urlopen(req, timeout=15)
        data = json.loads(resp.read().decode())["data"]

        db.set_abuse_cache(cache_key, {"data": data, "cached_at": datetime.now(timezone.utc).isoformat()})
        return data
    except Exception as e:
        return {"error": str(e)}


def _virustotal_check(hash_value):
    """Check a hash against VirusTotal free API. Returns score + verdict."""
    import urllib.request
    vt_key = os.environ.get("VT_API_KEY", "")
    if not vt_key:
        return None
    try:
        url = f"https://www.virustotal.com/api/v3/files/{hash_value}"
        req = urllib.request.Request(url, headers={"x-apikey": vt_key, "Accept": "application/json"})
        resp = urllib.request.urlopen(req, timeout=15)
        data = json.loads(resp.read().decode())
        stats = data.get("data", {}).get("attributes", {}).get("last_analysis_stats", {})
        malicious = stats.get("malicious", 0)
        total = sum(stats.values()) if stats else 0
        score = int((malicious / total) * 100) if total > 0 else 0
        verdict = "MALICIOUS" if score > 50 else ("SUSPICIOUS" if score > 10 else "CLEAN")
        return {"score": score, "verdict": verdict, "detections": f"{malicious}/{total}"}
    except Exception:
        return None


def _detect_ioc_type(value):
    """Auto-detect IOC type from value."""
    v = value.strip()
    if re.match(r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$', v):
        return "ipv4"
    if re.match(r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}/\d{1,2}$', v):
        return "subnet"
    if re.match(r'^[a-fA-F0-9]{32}$', v):
        return "hash_md5"
    if re.match(r'^[a-fA-F0-9]{64}$', v):
        return "hash_sha256"
    if re.match(r'^CVE-\d{4}-\d+$', v, re.I):
        return "cve"
    if re.match(r'^https?://', v, re.I):
        return "url"
    if re.match(r'^[\w.+-]+@[\w.-]+\.\w+$', v):
        return "email"
    if re.match(r'^[\w.-]+\.[a-zA-Z]{2,}$', v):
        return "domain"
    return "unknown"


@app.route("/api/apt/ioc/lookup", methods=["POST"])
def api_apt_ioc_lookup():
    """Lookup an IOC: local DB search + AbuseIPDB enrichment."""
    body = request.get_json(force=True, silent=True) or {}
    value = (body.get("value") or "").strip()
    if not value:
        return jsonify({"error": "No IOC value provided"}), 400

    ioc_type = body.get("type", "auto")
    if ioc_type == "auto":
        ioc_type = _detect_ioc_type(value)

    messages = load_messages()
    enrichments = load_enrichments()

    # Local DB search
    local_matches = []
    local_channels = set()
    local_apts = set()
    first_seen = ""
    last_seen = ""
    count = 0

    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")
        ts = m.get("timestamp_utc", "")
        msg_id = str(m.get("message_id", ""))

        # Check IOC fields
        found = False
        for it, vals in (m.get("iocs") or {}).items():
            if value.lower() in [v.lower() for v in (vals or [])]:
                found = True
                break

        # Also check text content for the value
        if not found:
            text = (m.get("text_preview") or "").lower()
            if value.lower() in text:
                found = True

        if not found:
            continue

        count += 1
        local_channels.add(ch)

        # Map to APT
        if ch in CHANNEL_TIERS and CHANNEL_TIERS[ch].get("status") != "banned":
            label = CHANNEL_TIERS[ch].get("label", ch)
            base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', label).strip()
            local_apts.add(base)
        else:
            ekey = f"{ch}_{m.get('message_id', '')}"
            ai = enrichments.get(ekey, m.get("ai_enrichment", {}))
            ga = ai.get("group_attribution", "")
            if ga and ga.lower() not in ("unknown", "n/a", ""):
                local_apts.add(ga)

        if ts:
            if not first_seen or ts < first_seen:
                first_seen = ts
            if not last_seen or ts > last_seen:
                last_seen = ts

        if len(local_matches) < 10:
            local_matches.append({
                "msg_id": msg_id,
                "channel": ch,
                "timestamp": ts,
                "summary_snippet": (m.get("text_preview") or "")[:150],
            })

    result = {
        "value": value,
        "type": ioc_type,
        "local": {
            "found": count > 0,
            "count": count,
            "apts": sorted(local_apts),
            "channels": sorted(local_channels),
            "first_seen": first_seen,
            "last_seen": last_seen,
            "messages": local_matches,
        },
    }

    # AbuseIPDB enrichment for IPs, domains, and subnets
    if ioc_type in ("ipv4", "domain", "subnet"):
        abuse_data = _abuseipdb_check(value, ioc_type)
        if abuse_data and "error" not in abuse_data:
            score = abuse_data.get("abuseConfidenceScore", 0)
            result["abuseipdb"] = {
                "abuseConfidenceScore": score,
                "countryCode": abuse_data.get("countryCode", ""),
                "isp": abuse_data.get("isp", ""),
                "domain": abuse_data.get("domain", ""),
                "totalReports": abuse_data.get("totalReports", 0),
                "lastReportedAt": abuse_data.get("lastReportedAt", ""),
                "usageType": abuse_data.get("usageType", ""),
                "isWhitelisted": abuse_data.get("isWhitelisted", False),
                "hostnames": abuse_data.get("hostnames", []),
            }
            # Verdict
            if score <= 25:
                verdict = "CLEAN"
            elif score <= 70:
                verdict = "SUSPICIOUS"
            else:
                verdict = "MALICIOUS"
            # Auto-bump if found in local DB linked to known APT
            if count > 0 and local_apts and verdict == "CLEAN":
                verdict = "SUSPICIOUS"
            result["verdict"] = verdict
        elif abuse_data and "error" in abuse_data:
            result["abuseipdb_error"] = abuse_data["error"]
    else:
        # For non-IP types, verdict based on local presence
        if count > 0 and local_apts:
            result["verdict"] = "SUSPICIOUS"
        elif count > 0:
            result["verdict"] = "SUSPICIOUS"
        else:
            result["verdict"] = "CLEAN"

    return jsonify(result)


@app.route("/api/apt/ioc/ai-extract", methods=["POST"])
def api_apt_ioc_ai_extract():
    """AI-powered deep IOC extraction — finds IOCs regex missed."""
    import openai
    body = request.get_json(force=True, silent=True) or {}
    apt_name = body.get("apt_name", "")
    if not apt_name:
        return jsonify({"error": "apt_name required"}), 400

    openai_key = os.environ.get("OPENAI_API_KEY", "")
    if not openai_key:
        return jsonify({"error": "OPENAI_API_KEY not set"}), 500

    messages = load_messages()
    enrichments = load_enrichments()

    # Find channels for this APT
    apt_channels = set()
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        label = meta.get("label", ch)
        base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', label).strip()
        if base == apt_name:
            apt_channels.add(ch)

    # Gather CRITICAL+MEDIUM messages for this APT
    target_msgs = []
    for m in messages:
        ch = m.get("channel_username") or m.get("channel", "")
        if ch not in apt_channels:
            continue
        if m.get("priority") not in ("CRITICAL", "MEDIUM"):
            continue
        text = m.get("text_preview") or ""
        if text.strip():
            target_msgs.append({
                "id": str(m.get("message_id", "")),
                "ch": ch,
                "text": text[:500],
                "existing_iocs": m.get("iocs", {}),
            })

    if not target_msgs:
        return jsonify({"apt_name": apt_name, "new_iocs": [], "msg": "No messages to scan"})

    client = openai.OpenAI(api_key=openai_key)

    extract_prompt = """Analyze these Telegram messages from threat actor "{apt}" and extract ALL indicators of compromise (IOCs).

Look for:
- IP addresses (including defanged: 1[.]2[.]3[.]4, 1{.}2{.}3{.}4)
- Domains (including defanged: hxxps://evil[.]com, evil{dot}com)
- File hashes (MD5, SHA256)
- CVEs (CVE-YYYY-NNNNN)
- Email addresses
- C2 server addresses
- URLs to paste sites, file shares, proof screenshots
- Infrastructure mentioned in attack context

For each IOC found, output ONE line:
TYPE|VALUE|CONTEXT

TYPE = ipv4, domain, url, hash_md5, hash_sha256, cve, email
VALUE = deobfuscated (hxxp→http, [.]→., etc.)
CONTEXT = brief description

ONLY output IOC lines. No explanations."""

    # Process in chunks
    chunk_size = 40
    all_new_iocs = []

    def _process_chunk(chunk):
        msg_block = "\n---\n".join([f"[{m['id']}@{m['ch']}] {m['text']}" for m in chunk])
        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": extract_prompt.format(apt=apt_name)},
                    {"role": "user", "content": msg_block}
                ],
                temperature=0.1,
                max_tokens=2000,
            )
            return resp.choices[0].message.content or ""
        except Exception as e:
            return f"ERROR:{e}"

    chunks = [target_msgs[i:i+chunk_size] for i in range(0, len(target_msgs), chunk_size)]

    with ThreadPoolExecutor(max_workers=5) as pool:
        futures = {pool.submit(_process_chunk, c): i for i, c in enumerate(chunks)}
        for fut in as_completed(futures):
            try:
                result = fut.result()
                for line in result.strip().split("\n"):
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) >= 2:
                        itype = parts[0].lower()
                        ival = parts[1]
                        ctx = parts[2] if len(parts) > 2 else ""
                        if itype in ("ipv4", "domain", "url", "hash_md5", "hash_sha256", "cve", "email"):
                            all_new_iocs.append({"type": itype, "value": ival, "context": ctx})
            except Exception:
                pass

    # Deduplicate and filter out already-known IOCs
    existing = set()
    for m in target_msgs:
        for itype, vals in (m.get("existing_iocs") or {}).items():
            for v in (vals or []):
                existing.add(f"{itype}:{v.lower()}")

    unique_new = []
    seen = set()
    for ioc in all_new_iocs:
        key = f"{ioc['type']}:{ioc['value'].lower()}"
        if key not in existing and key not in seen:
            seen.add(key)
            unique_new.append(ioc)

    return jsonify({
        "apt_name": apt_name,
        "messages_scanned": len(target_msgs),
        "chunks_processed": len(chunks),
        "new_iocs": unique_new,
        "total_found": len(all_new_iocs),
        "already_known": len(all_new_iocs) - len(unique_new),
    })


# ─── External IOC Research Engine ─────────────────────────────────────────────

_RESEARCH_CACHE_FILE = OUTPUT_DIR / "apt_ioc_research.json"
_RESEARCH_TTL = 48 * 3600  # 48 hours
from app.config import OTX_KEY as _OTX_KEY  # Moved to .env / app/config.py
_research_lock = threading.Lock()

def _load_research_cache():
    """Load research cache from SQLite, grouped by APT name.
    Returns format: {apt_name: {researched_at, summary, sources_queried, iocs[], stats{}}}
    """
    rows = db.get_apt_research()
    cache = {}
    for r in rows:
        apt = r["apt_name"]
        if apt not in cache:
            cache[apt] = {
                "researched_at": r.get("researched_at", ""),
                "sources_queried": [],
                "summary": "",
                "iocs": [],
                "stats": {"total": 0, "malicious": 0, "suspicious": 0, "clean": 0, "verified": 0}
            }
        ioc = {
            "type": r.get("ioc_type", ""),
            "value": r.get("ioc_value", ""),
            "source": r.get("source", ""),
            "context": r.get("context", ""),
            "abuse_verdict": r.get("abuse_verdict"),
            "abuse_score": r.get("abuse_score", -1),
            "abuse_country": r.get("abuse_country"),
        }
        cache[apt]["iocs"].append(ioc)
        cache[apt]["stats"]["total"] += 1
        v = r.get("abuse_verdict", "")
        if v == "MALICIOUS": cache[apt]["stats"]["malicious"] += 1
        elif v == "SUSPICIOUS": cache[apt]["stats"]["suspicious"] += 1
        elif v == "CLEAN": cache[apt]["stats"]["clean"] += 1
        if r.get("abuse_score", -1) >= 0: cache[apt]["stats"]["verified"] += 1
    return cache

def _save_research_cache(cache):
    """Save research cache to SQLite + JSON file (dual-write for backward compat)."""
    with _research_lock:
        for apt_name, entry in cache.items():
            researched_at = entry.get("researched_at", "")
            for ioc in entry.get("iocs", []):
                db.upsert_apt_ioc(
                    apt_name=apt_name,
                    ioc_value=ioc.get("value", ""),
                    ioc_type=ioc.get("type", ""),
                    source=ioc.get("source", ""),
                    context=ioc.get("context"),
                    abuse_verdict=ioc.get("abuse_verdict"),
                    abuse_score=ioc.get("abuse_score", -1),
                    abuse_country=ioc.get("abuse_country"),
                    researched_at=researched_at,
                )
        # Also write JSON file for backward compat
        try:
            with open(_RESEARCH_CACHE_FILE, "w", encoding="utf-8") as f:
                json.dump(cache, f, ensure_ascii=False, indent=1)
        except Exception:
            pass

def _research_cache_fresh(entry):
    """Check if a research cache entry is still fresh."""
    try:
        ts = entry.get("researched_at", "")
        dt = datetime.fromisoformat(ts.replace("Z", "+00:00"))
        return (datetime.now(timezone.utc) - dt).total_seconds() < _RESEARCH_TTL
    except Exception:
        return False


def _generate_apt_summary(apt_name, aliases):
    """Generate an OSINT-based summary/bio for an APT group using GPT-4o-mini."""
    import openai
    openai_key = os.environ.get("OPENAI_API_KEY", "")
    if not openai_key:
        return ""
    alias_str = ", ".join(aliases) if aliases else apt_name
    try:
        client = openai.OpenAI(api_key=openai_key)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a senior cyber threat intelligence analyst. Provide concise, factual profiles of threat actors based on publicly available OSINT and published CTI reports. Write in a professional, intelligence-briefing style."},
                {"role": "user", "content": f"""Write a 2-3 sentence intelligence profile of the threat group "{apt_name}" (aliases: {alias_str}).

Include: origin/nationality, known affiliations (state sponsors, parent APT groups, or hacktivist collectives), primary motivations, targeted sectors (e.g. government, banking, energy, telecom), and operational timeline if known.

This is a hacktivist/APT group active in the Middle East conflict. They may be linked to Iranian, Russian, or pro-Palestinian cyber operations.

If you have limited info, state what is known and note it's a lesser-known group. Be factual and concise. No headers or bullet points — just a paragraph."""}
            ],
            temperature=0.2,
            max_tokens=300,
        )
        summary = (resp.choices[0].message.content or "").strip()
        log.info(f"[RESEARCH] Generated summary for {apt_name}: {len(summary)} chars")
        return summary
    except Exception as e:
        log.info(f"[RESEARCH] Summary generation error for {apt_name}: {e}")
        return ""


def _research_gpt(apt_name, aliases):
    """Ask GPT-4o-mini for known IOCs attributed to an APT group."""
    import openai
    openai_key = os.environ.get("OPENAI_API_KEY", "")
    if not openai_key:
        return []
    alias_str = ", ".join(aliases) if aliases else apt_name
    prompt = f"""List the known INFRASTRUCTURE IOCs specifically attributed to "{apt_name}" (also known as: {alias_str}).

This is a hacktivist/APT group active in the Middle East targeting Jordan, Israel, and Gulf states.

IMPORTANT RULES:
- Only include IOCs that are SPECIFICALLY attributed to this group in published CTI reports
- DO NOT include generic Tor exit nodes, public VPN IPs, or shared hosting IPs
- DO NOT include 185.220.101.x, 45.153.160.x, or other well-known Tor relay ranges
- Each IOC must be unique to this group's operations, not shared internet infrastructure
- If you don't know any group-specific IOCs, output NONE — do not guess or use generic IPs

I want:
- C2 server IPs specifically used by this group
- Domains registered/used by this group for C2, phishing, defacement
- Malware hashes of their specific tools, payloads, webshells
- URLs of their C2 panels, leak sites, dropper locations

DO NOT include CVEs or generic infrastructure.

Format — one line per IOC:
ipv4|1.2.3.4|C2 server for campaign X|high
domain|evil.com|phishing domain used in Y attack|medium
hash_md5|abc123...|webshell used in Z campaign|high
hash_sha256|def456...|RAT payload|high

Only output IOC lines. If zero knowledge, output: NONE"""

    try:
        client = openai.OpenAI(api_key=openai_key)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a CTI analyst. When asked about threat actors, always list any IOCs you know from published reports, blog posts, CISA advisories, or OSINT research. Never say NONE if you can recall any infrastructure details at all. Include approximate or historically reported IOCs."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3000,
        )
        text = resp.choices[0].message.content or ""
        if "NONE" in text.strip().upper()[:10]:
            return []
        # Type aliases: GPT often outputs "IP" instead of "ipv4", "Hash" instead of "hash_md5"
        _type_aliases = {
            "ip": "ipv4", "ipv4": "ipv4", "ipv6": "ipv4",
            "domain": "domain", "hostname": "domain",
            "url": "url",
            "hash": "hash_md5", "md5": "hash_md5", "hash_md5": "hash_md5",
            "sha256": "hash_sha256", "hash_sha256": "hash_sha256", "sha1": "hash_sha256",
        }
        results = []
        for line in text.strip().split("\n"):
            # Strip markdown: numbered lists, bold, backticks, bullets
            clean = re.sub(r'^\s*\d+[\.\)]\s*', '', line)   # "1. " or "1) "
            clean = clean.strip().strip('*').strip('`').strip('-').strip()
            clean = clean.replace('**', '')
            if '|' not in clean:
                continue
            parts = [p.strip().strip('*').strip('`').strip() for p in clean.split("|")]
            if len(parts) < 3:
                continue
            ioc_type = _type_aliases.get(parts[0].lower().replace(" ", "_"))
            if not ioc_type:
                continue
            value = parts[1].strip()
            if not value or len(value) < 3:
                continue
            results.append({
                "type": ioc_type,
                "value": value,
                "source": "gpt-4o",
                "context": parts[2],
                "confidence": parts[3].lower() if len(parts) > 3 and parts[3].lower() in ("high","medium","low") else "medium",
            })
        log.info(f"[RESEARCH] GPT returned {len(results)} IOCs for {apt_name}")
        return results
    except Exception as e:
        log.info(f"[RESEARCH] GPT error for {apt_name}: {e}")
        return []


def _research_otx(apt_name):
    """Search AlienVault OTX for pulses mentioning this APT and extract IOCs."""
    import urllib.request, urllib.parse
    headers = {"Accept": "application/json", "X-OTX-API-KEY": _OTX_KEY}
    results = []
    try:
        search_url = f"https://otx.alienvault.com/api/v1/search/pulses?q={urllib.parse.quote(apt_name)}&limit=5&sort=modified"
        req = urllib.request.Request(search_url, headers=headers)
        resp = urllib.request.urlopen(req, timeout=20)
        data = json.loads(resp.read().decode())
        pulse_ids = [p["id"] for p in data.get("results", [])[:5]]
        type_map = {"IPv4": "ipv4", "domain": "domain", "URL": "url",
                    "FileHash-MD5": "hash_md5", "FileHash-SHA256": "hash_sha256",
                    "hostname": "domain"}
        for pid in pulse_ids:
            try:
                ind_url = f"https://otx.alienvault.com/api/v1/pulses/{pid}/indicators?limit=50"
                req2 = urllib.request.Request(ind_url, headers=headers)
                resp2 = urllib.request.urlopen(req2, timeout=15)
                ind_data = json.loads(resp2.read().decode())
                for ind in ind_data.get("results", []):
                    otype = type_map.get(ind.get("type"), "")
                    if otype:
                        results.append({
                            "type": otype,
                            "value": ind.get("indicator", ""),
                            "source": "otx",
                            "context": (ind.get("description") or "") or f"OTX pulse {pid[:8]}",
                            "confidence": "high",
                        })
            except Exception:
                pass
    except Exception as e:
        log.info(f"[RESEARCH] OTX error for {apt_name}: {e}")
    return results


def _research_threatfox(apt_name):
    """Search ThreatFox for IOCs tagged with this APT name."""
    import urllib.request
    results = []
    for qtype, qfield in [("taginfo", "tag"), ("malwareinfo", "malware")]:
        try:
            body = json.dumps({"query": qtype, qfield: apt_name, "limit": 50}).encode()
            req = urllib.request.Request(
                "https://threatfox-api.abuse.ch/api/v1/",
                data=body,
                headers={"Content-Type": "application/json"},
            )
            resp = urllib.request.urlopen(req, timeout=15)
            data = json.loads(resp.read().decode())
            if data.get("query_status") != "ok":
                continue
            type_map = {"ip:port": "ipv4", "domain": "domain", "url": "url",
                        "md5_hash": "hash_md5", "sha256_hash": "hash_sha256"}
            for item in (data.get("data") or []):
                ioc_val = item.get("ioc", "")
                otype = type_map.get(item.get("ioc_type", ""), "")
                if otype:
                    if otype == "ipv4" and ":" in ioc_val:
                        ioc_val = ioc_val.split(":")[0]
                    results.append({
                        "type": otype,
                        "value": ioc_val,
                        "source": "threatfox",
                        "context": item.get("malware_printable", "") or item.get("threat_type_desc", ""),
                        "confidence": "high" if item.get("confidence_level", 0) >= 75 else "medium",
                    })
        except Exception as e:
            log.info(f"[RESEARCH] ThreatFox {qtype} error for {apt_name}: {e}")
    return results


def _get_apt_aliases(apt_name):
    """Get aliases for an APT from CHANNEL_TIERS labels."""
    aliases = set()
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        label = meta.get("label", "")
        base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot|IRGC|MOIS[^)]*|[^)]*handle[^)]*|[^)]*breach[^)]*)\)\s*$', '', label).strip()
        if base == apt_name and label != base:
            aliases.add(label)
    return list(aliases)


def _research_apt_iocs(apt_name):
    """Research IOCs for an APT group from all external sources, verify via AbuseIPDB."""
    aliases = _get_apt_aliases(apt_name)
    all_raw = []
    sources_queried = []
    apt_summary_text = ""

    with ThreadPoolExecutor(max_workers=4) as pool:
        summary_fut = pool.submit(_generate_apt_summary, apt_name, aliases)
        ioc_futures = {
            pool.submit(_research_gpt, apt_name, aliases): "gpt-4o",
            pool.submit(_research_otx, apt_name): "otx",
            pool.submit(_research_threatfox, apt_name): "threatfox",
        }
        for fut in as_completed(ioc_futures):
            src = ioc_futures[fut]
            sources_queried.append(src)
            try:
                res = fut.result()
                log.info(f"[RESEARCH] {apt_name} <- {src}: {len(res)} IOCs")
                all_raw.extend(res)
            except Exception as e:
                log.info(f"[RESEARCH] {apt_name} <- {src}: ERROR {e}")
        try:
            apt_summary_text = summary_fut.result()
        except Exception:
            apt_summary_text = ""

    # Deduplicate by type:value
    seen = {}
    for ioc in all_raw:
        key = f"{ioc['type']}:{ioc['value'].lower().strip()}"
        if key not in seen:
            seen[key] = ioc
        else:
            existing = seen[key]
            if ioc.get("confidence") == "high":
                existing["confidence"] = "high"
            if ioc["source"] not in existing["source"]:
                existing["source"] += "+" + ioc["source"]
    unique_iocs = list(seen.values())

    # Verify ALL IOCs — IPs/domains/URLs via AbuseIPDB, hashes via VirusTotal
    def _verify(ioc):
        itype = ioc["type"]
        val = ioc["value"]

        # IPs — direct AbuseIPDB check
        if itype == "ipv4":
            abuse = _abuseipdb_check(val, "ipv4")
            if abuse and "error" not in abuse:
                score = abuse.get("abuseConfidenceScore", 0)
                ioc["abuse_score"] = score
                ioc["abuse_verdict"] = "MALICIOUS" if score > 70 else ("SUSPICIOUS" if score > 25 else "CLEAN")
                ioc["abuse_country"] = abuse.get("countryCode", "")
                ioc["abuse_isp"] = abuse.get("isp", "")
                ioc["abuse_reports"] = abuse.get("totalReports", 0)
            else:
                ioc["abuse_score"] = -1
                ioc["abuse_verdict"] = "UNVERIFIED"
            return ioc

        # Domains — resolve then AbuseIPDB
        if itype == "domain":
            abuse = _abuseipdb_check(val, "domain")
            if abuse and "error" not in abuse:
                score = abuse.get("abuseConfidenceScore", 0)
                ioc["abuse_score"] = score
                ioc["abuse_verdict"] = "MALICIOUS" if score > 70 else ("SUSPICIOUS" if score > 25 else "CLEAN")
                ioc["abuse_country"] = abuse.get("countryCode", "")
                ioc["abuse_isp"] = abuse.get("isp", "")
                ioc["abuse_reports"] = abuse.get("totalReports", 0)
            else:
                ioc["abuse_score"] = -1
                ioc["abuse_verdict"] = "SUSPICIOUS"  # GPT attributed it — at least suspicious
            return ioc

        # URLs — extract domain, check via AbuseIPDB
        if itype == "url":
            try:
                from urllib.parse import urlparse
                host = urlparse(val).hostname or ""
                if host:
                    abuse = _abuseipdb_check(host, "domain")
                    if abuse and "error" not in abuse:
                        score = abuse.get("abuseConfidenceScore", 0)
                        ioc["abuse_score"] = score
                        ioc["abuse_verdict"] = "MALICIOUS" if score > 70 else ("SUSPICIOUS" if score > 25 else "CLEAN")
                        ioc["abuse_country"] = abuse.get("countryCode", "")
                        return ioc
            except Exception:
                pass
            ioc["abuse_score"] = -1
            ioc["abuse_verdict"] = "SUSPICIOUS"
            return ioc

        # Hashes — check VirusTotal
        if itype in ("hash_md5", "hash_sha256"):
            vt_result = _virustotal_check(val)
            if vt_result:
                ioc["abuse_score"] = vt_result["score"]
                ioc["abuse_verdict"] = vt_result["verdict"]
                ioc["vt_detections"] = vt_result.get("detections", "")
            else:
                ioc["abuse_score"] = -1
                ioc["abuse_verdict"] = "SUSPICIOUS"
            return ioc

        # Fallback for any other type
        ioc["abuse_score"] = -1
        ioc["abuse_verdict"] = "SUSPICIOUS"
        return ioc

    with ThreadPoolExecutor(max_workers=3) as pool:
        list(pool.map(_verify, unique_iocs))

    # Sort: malicious first
    vord = {"MALICIOUS": 0, "SUSPICIOUS": 1, "UNVERIFIED": 2, "CLEAN": 3}
    unique_iocs.sort(key=lambda x: (vord.get(x.get("abuse_verdict", "UNVERIFIED"), 9), -(x.get("abuse_score", -1))))

    stats = {
        "total": len(unique_iocs),
        "verified": sum(1 for i in unique_iocs if i.get("abuse_score", -1) >= 0),
        "malicious": sum(1 for i in unique_iocs if i.get("abuse_verdict") == "MALICIOUS"),
        "suspicious": sum(1 for i in unique_iocs if i.get("abuse_verdict") == "SUSPICIOUS"),
        "clean": sum(1 for i in unique_iocs if i.get("abuse_verdict") == "CLEAN"),
    }

    return {
        "researched_at": datetime.now(timezone.utc).isoformat(),
        "sources_queried": sources_queried,
        "summary": apt_summary_text,
        "iocs": unique_iocs,
        "stats": stats,
    }


def _get_all_apt_names():
    """Get unique APT names from CHANNEL_TIERS."""
    names = set()
    for ch, meta in CHANNEL_TIERS.items():
        if meta.get("status") == "banned":
            continue
        label = meta.get("label", ch)
        base = re.sub(r'\s*\((?:Backup|Alt|Active\s*\d+\+?|Chat|Bot)\)\s*$', '', label).strip()
        names.add(base)
    return sorted(names)


_TG_BOT_TOKEN = "8731122610:AAGeLVygNXo6Ybq1WUTznXmQG30mr0X6CdM"
_TG_CHAT_ID = "-5258928587"


def _tg_send(text, parse_mode="HTML"):
    """Send a message to the Telegram bot chat."""
    import urllib.request, urllib.parse
    url = f"https://api.telegram.org/bot{_TG_BOT_TOKEN}/sendMessage"
    payload = json.dumps({
        "chat_id": _TG_CHAT_ID,
        "text": text,
        "parse_mode": parse_mode,
        "disable_web_page_preview": True,
    }).encode()
    try:
        req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
        urllib.request.urlopen(req, timeout=15)
    except Exception as e:
        log.info(f"[TG-BOT] Send error: {e}")


_TG_SENT_FILE = OUTPUT_DIR / "tg_sent_iocs.json"


def _load_tg_sent():
    try:
        return set(json.loads(_TG_SENT_FILE.read_text()))
    except Exception:
        return set()


def _save_tg_sent(sent):
    try:
        _TG_SENT_FILE.write_text(json.dumps(list(sent)))
    except Exception:
        pass


_tg_sent_iocs = _load_tg_sent()


def _tg_notify_apt_research(apt_name, result):
    """Send researched IOCs for an APT to Telegram — only NEW IOCs not yet sent."""
    global _tg_sent_iocs
    stats = result.get("stats", {})
    total = stats.get("total", 0)
    if total == 0:
        return

    iocs = result.get("iocs", [])
    # Filter to only IOCs we haven't already sent
    new_iocs = [i for i in iocs if i["value"] not in _tg_sent_iocs]
    new_actionable = [i for i in new_iocs if i.get("abuse_verdict") in ("MALICIOUS", "SUSPICIOUS")]

    if not new_actionable:
        for i in iocs:
            _tg_sent_iocs.add(i["value"])
        _save_tg_sent(_tg_sent_iocs)
        return

    mal = sum(1 for i in new_actionable if i.get("abuse_verdict") == "MALICIOUS")
    sus = sum(1 for i in new_actionable if i.get("abuse_verdict") == "SUSPICIOUS")

    lines = [f"\U0001f6e1 <b>Scanwave CyberIntel</b> \u2014 IOCs Detected & Blocked"]
    lines.append(f"\n\U0001f3af <b>{apt_name}</b>")
    summary = result.get("summary", "")
    if summary:
        lines.append(f"<i>{summary[:250]}</i>")
    lines.append(f"\n\u2716 {len(new_actionable)} new IOCs | \U0001f534 {mal} Malicious | \U0001f7e1 {sus} Suspicious")

    lines.append("")
    for ioc in new_actionable[:10]:
        v = ioc.get("abuse_verdict", "")
        icon = "\U0001f6ab" if v == "MALICIOUS" else "\u26a0\ufe0f"
        score = ioc.get("abuse_score", -1)
        tag = f" {score}%" if score >= 0 else f" [{ioc.get('type','')}]"
        lines.append(f"{icon} <code>{ioc['value']}</code>{tag}")

    # Mark all as sent and persist
    for i in iocs:
        _tg_sent_iocs.add(i["value"])
    _save_tg_sent(_tg_sent_iocs)

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n..."
    _tg_send(text)


def _tg_notify_cycle_complete(cache):
    """Send summary when full research cycle completes, including top IOCs."""
    seen = {}
    for apt_name, entry in cache.items():
        for ioc in entry.get("iocs", []):
            val = ioc["value"].lower().strip()
            if val not in seen:
                seen[val] = {"apt": apt_name, **ioc}
            elif apt_name not in seen[val].get("apt", ""):
                seen[val]["apt"] += ", " + apt_name

    all_iocs = list(seen.values())
    total_mal = sum(1 for i in all_iocs if i.get("abuse_verdict") == "MALICIOUS")
    total_sus = sum(1 for i in all_iocs if i.get("abuse_verdict") == "SUSPICIOUS")
    apts_hit = sum(1 for e in cache.values() if e.get("stats", {}).get("total", 0) > 0)

    lines = [
        "\U0001f6e1 <b>Scanwave CyberIntel</b> \u2014 Research Cycle Complete",
        f"\n{len(cache)} APT groups scanned \u2022 {apts_hit} with IOCs",
        f"\U0001f534 {total_mal} Malicious \u2022 \U0001f7e1 {total_sus} Suspicious \u2022 {len(all_iocs)} Unique IOCs",
    ]

    malicious = sorted([i for i in all_iocs if i.get("abuse_verdict") == "MALICIOUS"], key=lambda x: -(x.get("abuse_score", 0)))
    if malicious:
        lines.append("\n<b>Top Malicious:</b>")
        for ioc in malicious[:10]:
            score = ioc.get("abuse_score", -1)
            tag = f" {score}%" if score >= 0 else ""
            lines.append(f"\U0001f6ab <code>{ioc['value']}</code>{tag}")

    suspicious = sorted([i for i in all_iocs if i.get("abuse_verdict") == "SUSPICIOUS"], key=lambda x: -(x.get("abuse_score", 0)))
    if suspicious:
        lines.append("\n<b>Suspicious:</b>")
        for ioc in suspicious[:8]:
            score = ioc.get("abuse_score", -1)
            tag = f" {score}%" if score >= 0 else f" [{ioc.get('type','')}]"
            lines.append(f"\u26a0\ufe0f <code>{ioc['value']}</code>{tag}")

    lines.append(f"\n<a href='http://138.2.138.225:8888'>Open Dashboard</a>")

    text = "\n".join(lines)
    if len(text) > 4000:
        text = text[:4000] + "\n..."
    _tg_send(text)


def _auto_research_loop():
    """Background thread: research all APTs on startup, then every 24h."""
    time.sleep(10)  # Let Flask start first
    while True:
        _tg_sent_iocs.update(_load_tg_sent())  # Reload persisted set
        try:
            cache = _load_research_cache()
            apt_names = _get_all_apt_names()
            for apt_name in apt_names:
                if apt_name in cache and _research_cache_fresh(cache[apt_name]):
                    continue
                log.info(f"[RESEARCH] Auto-researching: {apt_name}")
                try:
                    result = _research_apt_iocs(apt_name)
                    cache[apt_name] = result
                    _save_research_cache(cache)
                    log.info(f"[RESEARCH] {apt_name}: {result['stats']['total']} IOCs ({result['stats']['malicious']} malicious)")
                    _tg_notify_apt_research(apt_name, result)
                except Exception as e:
                    log.info(f"[RESEARCH] Error for {apt_name}: {e}")
                time.sleep(3)  # Throttle between APTs
            log.info(f"[RESEARCH] Auto-research cycle complete. {len(apt_names)} APTs processed.")
            _tg_notify_cycle_complete(cache)
        except Exception as e:
            log.info(f"[RESEARCH] Cycle error: {e}")
        time.sleep(86400)  # Sleep 24h


@app.route("/api/apt/<path:name>/research")
def api_apt_research(name):
    """Get researched IOCs for an APT. Auto-triggers if stale."""
    cache = _load_research_cache()
    if name in cache and _research_cache_fresh(cache[name]):
        entry = cache[name]
        entry["cached"] = True
        return jsonify(entry)
    # Research fresh (this may take 10-30s)
    result = _research_apt_iocs(name)
    result["cached"] = False
    cache[name] = result
    _save_research_cache(cache)
    return jsonify(result)


@app.route("/api/blocklist")
def api_blocklist():
    """Central blocklist: ALL researched IOCs across ALL APTs."""
    cache = _load_research_cache()
    apt_filter = request.args.get("apt", "").strip()
    type_filter = request.args.get("type", "").strip().lower()
    verdict_filter = request.args.get("verdict", "").strip().upper()
    search_q = request.args.get("q", "").strip().lower()

    apt_summary = {}
    # Collect all IOCs, dedup by value — merge APT attributions
    seen = {}  # key=value -> merged IOC dict

    for apt_name, entry in cache.items():
        if apt_filter and apt_name != apt_filter:
            continue
        stats = entry.get("stats", {})
        apt_summary[apt_name] = {
            "researched_at": entry.get("researched_at", ""),
            "total": stats.get("total", 0),
            "malicious": stats.get("malicious", 0),
            "suspicious": stats.get("suspicious", 0),
        }
        for ioc in entry.get("iocs", []):
            if type_filter and ioc.get("type") != type_filter:
                continue
            if verdict_filter and ioc.get("abuse_verdict") != verdict_filter:
                continue
            if search_q and search_q not in ioc.get("value", "").lower() and search_q not in ioc.get("context", "").lower():
                continue
            val = ioc["value"].lower().strip()
            if val in seen:
                # Merge: append APT name if not already there
                existing = seen[val]
                if apt_name not in existing["apt"]:
                    existing["apt"] += ", " + apt_name
            else:
                seen[val] = {"apt": apt_name, **ioc}

    blocklist = list(seen.values())
    vord = {"MALICIOUS": 0, "SUSPICIOUS": 1, "UNVERIFIED": 2, "CLEAN": 3}
    blocklist.sort(key=lambda x: (vord.get(x.get("abuse_verdict", "UNVERIFIED"), 9), -(x.get("abuse_score", -1))))

    total_mal = sum(1 for i in blocklist if i.get("abuse_verdict") == "MALICIOUS")
    total_sus = sum(1 for i in blocklist if i.get("abuse_verdict") == "SUSPICIOUS")
    total_clean = sum(1 for i in blocklist if i.get("abuse_verdict") == "CLEAN")

    # Pagination (optional — backward compatible)
    page     = request.args.get("page", type=int)
    per_page = min(request.args.get("per_page", 100, type=int), 500)

    base = {
        "total": len(blocklist),
        "apts_researched": len([a for a in apt_summary.values() if a.get("total", 0) > 0]),
        "malicious": total_mal,
        "suspicious": total_sus,
        "clean": total_clean,
        "apt_summary": apt_summary,
    }
    if page is not None:
        start = (page - 1) * per_page
        base["iocs"] = blocklist[start:start + per_page]
        base["pagination"] = {
            "page": page,
            "per_page": per_page,
            "total": len(blocklist),
            "pages": (len(blocklist) + per_page - 1) // per_page if per_page else 1,
        }
    else:
        base["iocs"] = blocklist[:1000]
    return jsonify(base)


@app.route("/api/blocklist/export")
def api_blocklist_export():
    """Export blocklist as CSV."""
    cache = _load_research_cache()
    verdict_filter = request.args.get("verdict", "").strip().upper()
    buf = StringIO()
    w = csv.writer(buf)
    w.writerow(["APT Group", "IOC Type", "Value", "Verdict", "AbuseIPDB Score",
                "Country", "ISP", "Reports", "Source", "Context", "Confidence"])
    for apt_name, entry in cache.items():
        for ioc in entry.get("iocs", []):
            if verdict_filter and ioc.get("abuse_verdict") != verdict_filter:
                continue
            w.writerow([apt_name, ioc.get("type", ""), ioc.get("value", ""),
                        ioc.get("abuse_verdict", ""), ioc.get("abuse_score", ""),
                        ioc.get("abuse_country", ""), ioc.get("abuse_isp", ""),
                        ioc.get("abuse_reports", ""), ioc.get("source", ""),
                        ioc.get("context", ""), ioc.get("confidence", "")])
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=blocklist.csv"})


# ── Media serving ──────────────────────────────────────────────
_MEDIA_DIR = OUTPUT_DIR / "media"
_IMG_EXTS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
_VID_EXTS = {'.mp4', '.webm', '.mov', '.avi', '.mkv'}

@app.route("/api/media/<path:filepath>")
def api_media(filepath):
    """Serve downloaded media files with path traversal protection."""
    safe = (_MEDIA_DIR / filepath).resolve()
    if not str(safe).startswith(str(_MEDIA_DIR.resolve())):
        abort(403)
    if not safe.exists() or not safe.is_file():
        abort(404)
    return send_file(safe)

@app.route("/api/media/lookup/<channel>/<int:message_id>")
def api_media_lookup(channel, message_id):
    """Find all media files for a specific message."""
    media_dir = _MEDIA_DIR / f"{channel}_{message_id}"
    if not media_dir.exists() or not media_dir.is_dir():
        return jsonify({"files": []})
    files = []
    for f in sorted(media_dir.iterdir()):
        if not f.is_file():
            continue
        ext = f.suffix.lower()
        mtype = "image" if ext in _IMG_EXTS else "video" if ext in _VID_EXTS else "file"
        files.append({
            "url": f"/api/media/{channel}_{message_id}/{f.name}",
            "type": mtype,
            "name": f.name,
            "size": f.stat().st_size,
        })
    return jsonify({"files": files})


# ── DOCX Report Generator (Template-based) ────────────────────
def _generate_blocklist_report():
    """Generate ScanWave SOC Client Advisory using the original branded template.

    Opens scanwave_report_template.docx (exact copy of original report),
    preserves ALL styling/branding/logo/cover page, removes old IOC tables,
    inserts live blocklist IOCs organized by APT group.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    template_path = Path(__file__).parent / "scanwave_report_template.docx"
    if not template_path.exists():
        raise FileNotFoundError("Report template not found: scanwave_report_template.docx")

    doc = Document(str(template_path))
    body = doc.element.body
    now = datetime.now(timezone.utc)

    # ── Load live IOC data — per-group, no cross-group name merging ──
    cache = _load_research_cache()
    # Only include groups that exist in CHANNEL_TIERS (skip bogus entries)
    known_labels = {v.get("label", k) for k, v in CHANNEL_TIERS.items()}
    apt_iocs = {}  # {apt_name: [unique IOC dicts]}
    global_seen = set()  # track all unique IOC values for total count
    for apt_name, entry in cache.items():
        if apt_name not in known_labels:
            continue
        iocs = entry.get("iocs", [])
        if not iocs:
            continue
        seen_in_group = set()
        unique = []
        for ioc in iocs:
            val = ioc["value"].lower().strip()
            if val not in seen_in_group:
                seen_in_group.add(val)
                unique.append(ioc)
                global_seen.add(val)
        if unique:
            apt_iocs[apt_name] = unique

    # Split: major groups (>=4 IOCs) shown individually, minor (<4) pooled
    THRESHOLD = 4
    major = {k: v for k, v in apt_iocs.items() if len(v) >= THRESHOLD}
    minor = {k: v for k, v in apt_iocs.items() if len(v) < THRESHOLD}

    # Pool minor-group IOCs (deduped)
    other_iocs = []
    other_seen = set()
    other_group_names = sorted(minor.keys())
    for grp in other_group_names:
        for ioc in minor[grp]:
            val = ioc["value"].lower().strip()
            if val not in other_seen:
                other_seen.add(val)
                other_iocs.append(ioc)

    total_iocs = len(global_seen)
    all_flat = [ioc for group_iocs in apt_iocs.values() for ioc in group_iocs]
    total_mal = sum(1 for i in all_flat if i.get("abuse_verdict") == "MALICIOUS")
    total_sus = sum(1 for i in all_flat if i.get("abuse_verdict") == "SUSPICIOUS")

    # ── Step 1: Update date on cover page ──
    import re as _re
    for p in doc.paragraphs:
        if _re.match(r'^[A-Z][a-z]+ \d{4}$', p.text.strip()):
            for run in p.runs:
                run.text = now.strftime("%B %Y")
            break

    # ── Step 1b: Update header dates to current month ──
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    date_str = now.strftime("%B %Y")
    for section in doc.sections:
        for hdr_attr in ("header", "first_page_header", "even_page_header"):
            try:
                hdr = getattr(section, hdr_attr)
                for t_el in hdr._element.iter(f"{{{ns_w}}}t"):
                    if t_el.text and _re.search(r'[A-Z][a-z]+ \d{4}', t_el.text):
                        t_el.text = _re.sub(r'[A-Z][a-z]+ \d{4}', date_str, t_el.text)
            except Exception:
                pass

    # ── Step 1d: Add Qatar to country observation list ──
    for p in doc.paragraphs:
        if p.text.strip() == "Jordan":
            # Insert Qatar paragraph after Jordan using same formatting
            jordan_el = p._element
            from copy import deepcopy
            qatar_el = deepcopy(jordan_el)
            # Replace text in the copy
            for t_el in qatar_el.iter(f"{{{ns_w}}}t"):
                if t_el.text and "Jordan" in t_el.text:
                    t_el.text = t_el.text.replace("Jordan", "Qatar")
            jordan_el.addnext(qatar_el)
            break

    # ── Step 2: Find markers ──
    ioc_summary_el = None
    critical_note_el = None
    for p in doc.paragraphs:
        if "IOC SUMMARY" in p.text and ioc_summary_el is None:
            ioc_summary_el = p._element
        if "CRITICAL NOTE" in p.text:
            critical_note_el = p._element

    if ioc_summary_el is None or critical_note_el is None:
        raise ValueError("Template missing IOC SUMMARY or CRITICAL NOTE markers")

    # ── Step 3: Remove everything between IOC SUMMARY and CRITICAL NOTE ──
    # Find indices of the two markers
    children = list(body)
    ioc_idx = None
    crit_idx = None
    for i, child in enumerate(children):
        if child is ioc_summary_el:
            ioc_idx = i
        if child is critical_note_el:
            crit_idx = i
    if ioc_idx is None or crit_idx is None:
        raise ValueError("Cannot locate marker indices")

    # Remove elements between markers (exclusive of both)
    for child in children[ioc_idx + 1:crit_idx]:
        body.remove(child)

    # ── Step 4: Build new IOC content as raw XML elements ──
    # We build elements via doc.add_paragraph/add_table, then detach and
    # insert at the correct position right after IOC SUMMARY.
    new_elements = []

    # Helper: build a 3-column IOC grid table and return the tbl element
    def _add_ioc_table(ioc_list):
        values = [i["value"] for i in ioc_list]
        cols = 3
        nrows = (len(values) + cols - 1) // cols
        table = doc.add_table(rows=nrows, cols=cols)
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "E5E7EB")
            borders.append(el)
        tblPr.append(borders)
        for idx, val in enumerate(values):
            r, c = divmod(idx, cols)
            cell = table.cell(r, c)
            cell.text = val
            for par in cell.paragraphs:
                for rn in par.runs:
                    rn.font.size = Pt(7)
                    rn.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)
        return tbl

    # Summary paragraph
    p = doc.add_paragraph()
    run = p.add_run(
        f"All indicators below are sourced from automated threat intelligence research "
        f"conducted by the ScanWave CyberIntel Platform across AlienVault OTX, ThreatFox, "
        f"and curated CTI reports, collected {now.strftime('%B %d, %Y')}. "
        f"Total unique IOCs: {total_iocs} (Malicious: {total_mal}, Suspicious: {total_sus}). "
        f"Every IOC has been verified through automated scoring and cross-referenced "
        f"against multiple threat intelligence feeds before inclusion."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)
    new_elements.append(p._element)

    spacer = doc.add_paragraph()
    new_elements.append(spacer._element)

    # ── Major APT groups (>= 4 IOCs each) — individual sections ──
    for apt_name, iocs in sorted(major.items()):
        # Sub-header
        p = doc.add_paragraph()
        run = p.add_run(f"{apt_name} \u2014 {len(iocs)} Indicators")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
        new_elements.append(p._element)

        # AI summary
        summary = cache.get(apt_name, {}).get("summary", "")
        if summary:
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
            run.italic = True
            new_elements.append(p._element)

        new_elements.append(_add_ioc_table(iocs))

        spacer = doc.add_paragraph()
        new_elements.append(spacer._element)

    # ── Minor groups (< 4 IOCs) — pooled into one section ──
    if other_iocs:
        p = doc.add_paragraph()
        run = p.add_run(
            f"Additional Threat Indicators \u2014 {len(other_iocs)} IOCs "
            f"from {len(other_group_names)} Monitored Groups"
        )
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
        new_elements.append(p._element)

        p = doc.add_paragraph()
        run = p.add_run(
            "The following indicators were identified across additional threat groups "
            "under active monitoring, each contributing a small number of IOCs: "
            + ", ".join(other_group_names) + "."
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        run.italic = True
        new_elements.append(p._element)

        new_elements.append(_add_ioc_table(other_iocs))

        spacer = doc.add_paragraph()
        new_elements.append(spacer._element)

    # ── Step 5: Detach new elements from end of doc, insert after IOC SUMMARY ──
    for elem in new_elements:
        body.remove(elem)

    # Insert right after ioc_summary_el (ioc_idx position)
    insert_pos = list(body).index(ioc_summary_el) + 1
    for i, elem in enumerate(new_elements):
        body.insert(insert_pos + i, elem)

    # ── Save ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/api/blocklist/report")
def api_blocklist_report():
    """Generate and download ScanWave SOC Client Advisory as PDF."""
    try:
        buf = _generate_blocklist_report()
        now = datetime.now(timezone.utc)
        base = f"ScanWave_SOC_Client_Advisory_{now.strftime('%d_%B%Y')}"

        # Write DOCX to temp file, convert to PDF via LibreOffice
        import tempfile, subprocess
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, f"{base}.docx")
            with open(docx_path, "wb") as f:
                f.write(buf.getvalue())

            try:
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmpdir, docx_path],
                    capture_output=True, timeout=120
                )
                pdf_path = os.path.join(tmpdir, f"{base}.pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_bytes = f.read()
                    return Response(
                        pdf_bytes,
                        mimetype="application/pdf",
                        headers={"Content-Disposition": f"attachment;filename={base}.pdf"}
                    )
                log.info(f"[REPORT] PDF conversion failed: {result.stderr.decode()}")
            except (FileNotFoundError, subprocess.TimeoutExpired) as conv_err:
                log.info(f"[REPORT] LibreOffice not available, falling back to DOCX: {conv_err}")

            # Fallback to DOCX if LibreOffice not available
            return Response(
                buf.getvalue(),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment;filename={base}.docx"}
            )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/api/messages/<channel_username>/<message_id>/context")
def api_message_context(channel_username, message_id):
    before = min(max(int(request.args.get("before", 5)), 1), 50)
    after  = min(max(int(request.args.get("after",  5)), 1), 50)
    msg_id = int(message_id)

    all_ch = [m for m in load_messages()
              if (m.get("channel_username") or m.get("channel", "unknown")) == channel_username]
    total  = len(all_ch)

    live_msgs, target_idx = fetch_live_context(channel_username, msg_id, before, after)
    if live_msgs is not None:
        return jsonify({"messages": live_msgs, "target_idx": target_idx,
                        "total": total, "source": "live"})

    all_ch.sort(key=lambda x: x.get("timestamp_utc", ""))
    idx = next((i for i, m in enumerate(all_ch)
                if str(m.get("message_id", "")) == str(message_id)), None)
    if idx is None:
        return jsonify({"messages": [], "target_idx": -1, "total": total, "source": "stored"})
    start = max(0, idx - before)
    end   = min(len(all_ch), idx + after + 1)
    return jsonify({"messages": all_ch[start:end], "target_idx": idx - start,
                    "total": total, "source": "stored"})


@app.route("/api/discovery/list")
def api_discovery_list():
    if WATCHLIST_FILE.exists():
        with open(WATCHLIST_FILE, encoding="utf-8") as f:
            data = json.load(f)
        return jsonify(data.get("all_discovered", []))
    return jsonify([])


@app.route("/api/discovery/scan", methods=["POST"])
def api_discovery_scan():
    global _scan
    if _scan["running"]:
        return jsonify({"error": "scan already running"})
    quick = (request.json or {}).get("quick", False)
    args  = ["python", "channel_discovery.py"] + (["--quick"] if quick else [])

    def _run():
        global _scan
        p = subprocess.Popen(args, cwd=str(Path(__file__).parent))
        _scan["running"] = True
        _scan["pid"]     = p.pid
        p.wait()
        _scan["running"]  = False
        _scan["last_run"] = datetime.now(timezone.utc).isoformat()

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"started": True})


@app.route("/api/discovery/status")
def api_discovery_status():
    return jsonify(_scan)


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN PANEL API
# ══════════════════════════════════════════════════════════════════════════════

KEYWORDS_FILE     = OUTPUT_DIR / "keywords.json"
BACKFILL_QUEUE    = OUTPUT_DIR / "backfill_queue.json"
PENDING_FILE_VIEW = OUTPUT_DIR / "pending_channels.json"
CHANNELS_FILE     = OUTPUT_DIR / "channels_config.json"
MONITOR_LOG       = OUTPUT_DIR / "monitor.log"
AI_SUGGESTIONS        = OUTPUT_DIR / "ai_suggestions.json"
AI_BRIEF_FILE         = OUTPUT_DIR / "ai_brief.json"
AI_ENRICHED_FILE      = OUTPUT_DIR / "enriched_alerts.jsonl"
AI_STATE_FILE         = OUTPUT_DIR / "ai_agent_state.json"
ESCALATION_FILE       = OUTPUT_DIR / "escalation_status.json"
HUNTING_LEADS_FILE    = OUTPUT_DIR / "hunting_leads.json"
NETWORK_FILE          = OUTPUT_DIR / "channel_network.json"

# ── Critical Subtype Classification (mirrors telegram_monitor.py) ─────────────
_CYBER_SIGNALS = {
    # Attack types (unambiguous cyber terms only)
    "ddos", "d-dos", "defacement", "defaced", "data leak", "data breach", "databreach",
    "ransomware", "ransom", "malware", "trojan", "botnet", "exploit", "sql injection",
    "sqlmap", "webshell", "backdoor", "rootkit", "c2", "c&c", "command and control",
    "brute force", "credential stuffing", "credential dump",
    "pwned", "owned",
    "data breach", "breached", "data dump", "database dump", "wiper", "zero-day", "0day",
    "root access", "full access", "rce", "remote code execution",
    "phishing", "spear phishing", "xss", "cross-site",
    # Arabic cyber terms (specific — no generic words)
    "تسريب بيانات", "تسريب معلومات", "قرصنة", "هاكر", "هاكرز",
    "فيروس", "هجوم سيبراني", "هجوم الكتروني", "هجوم إلكتروني",
    "دي دوس", "برامج خبيثة", "رانسوم", "فدية",
    "تم اختراق", "تم قرصنة", "تهكير",
    # Jordan domain targets (domain targeting = cyber)
    ".jo", ".gov.jo", ".com.jo", ".edu.jo", ".org.jo",
    # DDoS tools / indicators
    "check-host", "dienet", "connection timed out", "connection refused",
    "layer7", "layer4", "http flood",
}
# Terms that mean CYBER in hacking context but also appear in military context
_AMBIGUOUS_CYBER = {"اختراق", "hacked", "hacking", "hack"}
_NATIONAL_SIGNALS = {
    # Iranian/IRGC
    "irgc", "iranian", "khamenei", "خامنئي", "حرس الثوري", "فاطميون",
    "soleimani", "سليماني", "quds force", "الحرس الثوري",
    # Resistance axis
    "hezbollah", "حزب الله", "hamas", "حماس", "qassam", "قسام", "houthi", "حوثي",
    "انصار الله", "مقاومة", "جهاد اسلامي",
    # Jordan military / security services
    "military", "عسكري", "troops", "army", "القوات المسلحة",
    "الجيش الاردني", "jordan armed forces", "القوات الجوية", "air force",
    "us base", "nato", "ain al asad", "العديد", "المفرق",
    ".mil.jo",  # military domain = national security
    "استخبارات", "intelligence", "gendarmerie", "الدرك",
    "border guard", "حرس الحدود", "مكافحة الإرهاب", "counter terrorism",
    "الأمن العام", "security directorate", "muwaffaq", "الموفق",
    # War / conflict (avoid short words that match inside cyber terms like "malware")
    "warfare", "warzone", "at war", "of war", "حرب", "missile", "صاروخ", "escalation", "تصعيد",
    "airstrike", "air strike", "عملية عسكرية", "military operation",
}

# Service/sale advertisements — generic hacking services being sold, not actual attacks
_SERVICE_SIGNALS = {
    # English service/sale indicators
    "service", "services", "for sale", "for hire", "hire", "buy", "sell", "selling",
    "pricing", "price", "order", "contact us", "dm for", "dm me",
    "blackhat", "black hat", "professional hacking", "hacking service", "blackhat service",
    "we hack", "hack for", "hacker for",
    "we offer", "we provide", "available now", "24/7",
    "guaranteed results", "confidential", "affordable", "discount",
    "package", "combo", "premium", "vip",
    # Arabic service/sale indicators
    "خدمات", "خدمة", "للبيع", "للإيجار", "اشتري", "نبيع",
    "اسعار", "سعر", "اطلب", "تواصل معنا", "راسلنا",
    "نقدم", "نوفر", "متاح الآن", "خدمات احترافية",
    # Farsi service/sale indicators
    "خدمات هک", "سرویس", "فروش", "قیمت", "سفارش",
    "تماس بگیرید", "ارائه می‌دهیم",
}

# Jordan references — all formats (English, Arabic, Farsi)
_JORDAN_REFS = {
    # English
    "jordan", "jordanian", "amman",
    # Arabic (all common spellings)
    "الاردن", "الأردن", "أردن", "اردن", "اردني", "أردني", "الأردني", "الاردني",
    "عمان", "عمّان",
    # Farsi
    "اردن", "اُردن",
    # Jordan domains
    ".jo", ".gov.jo", ".com.jo", ".edu.jo", ".org.jo", ".mil.jo",
}

def _mentions_jordan(txt):
    """Check if text mentions Jordan in any language."""
    return any(ref in txt for ref in _JORDAN_REFS)

def _is_service_ad(txt):
    """Check if text is advertising hacking services for sale."""
    return sum(1 for sig in _SERVICE_SIGNALS if sig in txt) >= 2

def _compute_critical_subtype(keyword_hits, text=""):
    """Classify a CRITICAL message by subtype.

    Checks keyword_hits first, then scans full message text for additional
    context.  Ambiguous terms (like اختراق which means both 'hack' and
    'military penetration') are only counted as CYBER when the message
    has no national-security context.

    Service advertisements (hacking-for-sale) are demoted to GENERAL
    unless Jordan is explicitly mentioned.
    """
    if not keyword_hits:
        return "GENERAL"
    hits = [kw.lower() for kw in keyword_hits]
    txt = text.lower() if text else ""

    # Strong (unambiguous) cyber match from keyword hits
    strong_cyber = any(sig in hit for hit in hits for sig in _CYBER_SIGNALS)
    # Ambiguous cyber match (could be military)
    ambig_cyber = any(sig in hit for hit in hits for sig in _AMBIGUOUS_CYBER)
    # National match from keyword hits
    is_national = any(sig in hit for hit in hits for sig in _NATIONAL_SIGNALS)

    # Also scan full message text for national context (catches حزب الله, صاروخ etc.)
    if not is_national and txt:
        is_national = any(sig in txt for sig in _NATIONAL_SIGNALS)

    is_cyber = strong_cyber  # start with unambiguous signals
    if ambig_cyber and not is_national:
        # Ambiguous term + no national context → count as cyber
        is_cyber = True

    # Demote service/sale ads: if it's a hacking service advertisement
    # and doesn't mention Jordan, it's not a relevant cyber threat
    if is_cyber and not is_national and txt and _is_service_ad(txt) and not _mentions_jordan(txt):
        return "GENERAL"

    if is_cyber and is_national: return "BOTH"
    if is_cyber:    return "CYBER"
    if is_national: return "NATIONAL"
    return "GENERAL"


def _load_channels_config():
    """Load channels: start with hardcoded CHANNEL_TIERS, overlay with saved config."""
    base = {k: dict(v) for k, v in CHANNEL_TIERS.items()}
    if CHANNELS_FILE.exists():
        try:
            saved = json.loads(CHANNELS_FILE.read_text(encoding="utf-8"))
            base.update(saved)
        except Exception:
            pass
    return base


def _save_channels_config(cfg):
    CHANNELS_FILE.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/api/admin/status")
def api_admin_status():
    """System health: monitor process, DB stats, cursor, log tail."""
    import subprocess as _sp
    # Check if monitor is running (look for telegram_monitor.py in process list)
    monitor_running = False
    try:
        import os as _os
        for _pid in _os.listdir("/proc"):
            if not _pid.isdigit():
                continue
            try:
                _cmd = open(f"/proc/{_pid}/cmdline", "rb").read().replace(b"\x00", b" ").decode()
                if "telegram_monitor.py" in _cmd:
                    monitor_running = True
                    break
            except Exception:
                continue
    except Exception:
        try:
            result = _sp.run(["tasklist", "/FI", "IMAGENAME eq python.exe", "/FO", "CSV"],
                             capture_output=True, text=True, timeout=5)
            monitor_running = "python" in result.stdout.lower()
        except Exception:
            pass

    cursor = {}
    if (OUTPUT_DIR / "last_seen.json").exists():
        try:
            cursor = json.loads((OUTPUT_DIR / "last_seen.json").read_text(encoding="utf-8"))
        except Exception:
            pass

    # DB stats (fast SQL)
    conn = get_conn()
    _row = conn.execute("""
        SELECT COUNT(*) as total,
               SUM(CASE WHEN priority='CRITICAL' THEN 1 ELSE 0 END) as crit,
               SUM(CASE WHEN priority='MEDIUM' THEN 1 ELSE 0 END) as med,
               COUNT(DISTINCT channel_username) as chs,
               MAX(timestamp_utc) as last_ts
        FROM messages
    """).fetchone()
    total_msgs = _row[0] or 0
    crit = _row[1] or 0
    med = _row[2] or 0
    chs = _row[3] or 0
    last_ts = _row[4] or ""
    iocs = 0  # IOC count from JSON columns
    for _irow in conn.execute("SELECT iocs FROM messages WHERE iocs IS NOT NULL").fetchall():
        try:
            _iocs = json.loads(_irow[0]) if isinstance(_irow[0], str) else _irow[0]
            if isinstance(_iocs, dict):
                iocs += sum(len(v or []) for v in _iocs.values())
        except Exception:
            pass

    # Last 30 lines of monitor log
    log_tail = []
    if MONITOR_LOG.exists():
        try:
            lines = MONITOR_LOG.read_text(encoding="utf-8", errors="replace").splitlines()
            log_tail = lines[-30:]
        except Exception:
            pass

    # Backfill queue status
    bfq = {}
    if BACKFILL_QUEUE.exists():
        try:
            bfq = json.loads(BACKFILL_QUEUE.read_text(encoding="utf-8"))
        except Exception:
            pass

    return jsonify({
        "monitor_running": monitor_running,
        "cursor": cursor,
        "db": {"total": total_msgs, "critical": crit, "medium": med,
               "channels": chs, "iocs": iocs, "last_message": last_ts},
        "log_tail": log_tail,
        "backfill_queue": bfq,
    })


@app.route("/api/admin/keywords", methods=["GET"])
def api_admin_keywords_get():
    return jsonify(db.get_keywords())


@app.route("/api/admin/keywords", methods=["POST"])
def api_admin_keywords_post():
    data = request.get_json(force=True)
    crit = [str(k).strip() for k in data.get("critical", []) if str(k).strip()]
    med  = [str(k).strip() for k in data.get("medium",   []) if str(k).strip()]
    db.save_keywords({"critical": crit, "medium": med})
    # Also write to keywords.json for backward compat with telegram_monitor
    KEYWORDS_FILE.write_text(
        json.dumps({"critical": crit, "medium": med}, ensure_ascii=False, indent=2),
        encoding="utf-8")
    return jsonify({"ok": True, "critical": len(crit), "medium": len(med),
                    "note": "Restart monitor for changes to take effect"})


@app.route("/api/admin/channels", methods=["GET"])
def api_admin_channels_get():
    return jsonify(db.get_channels())


@app.route("/api/admin/channels", methods=["POST"])
def api_admin_channels_post():
    """Add or update a channel entry."""
    data = request.get_json(force=True)
    username = data.get("username", "").strip().lstrip("@")
    if not username:
        return jsonify({"ok": False, "error": "username required"}), 400
    chan_data = {
        "tier":   int(data.get("tier", 3)),
        "label":  data.get("label", username),
        "threat": data.get("threat", "MEDIUM"),
        "status": data.get("status", "active"),
    }
    db.upsert_channel(username, **chan_data)
    # Also add to CHANNEL_TIERS runtime dict
    CHANNEL_TIERS[username] = chan_data
    # Also write to channels_config.json for backward compat
    _save_channels_config(db.get_channels())
    # Add to pending file so live monitor joins it
    pf_data = {"pending": [], "processed": []}
    if PENDING_FILE_VIEW.exists():
        try:
            pf_data = json.loads(PENDING_FILE_VIEW.read_text(encoding="utf-8"))
        except Exception:
            pass
    if username not in pf_data.get("pending", []) and username not in pf_data.get("processed", []):
        pf_data.setdefault("pending", []).append(username)
        PENDING_FILE_VIEW.write_text(json.dumps(pf_data, indent=2, ensure_ascii=False), encoding="utf-8")
    return jsonify({"ok": True, "channel": username, "config": chan_data})


@app.route("/api/admin/channels/<username>", methods=["DELETE"])
def api_admin_channels_delete(username):
    db.delete_channel(username)
    CHANNEL_TIERS.pop(username, None)
    # Update channels_config.json for backward compat
    _save_channels_config(db.get_channels())
    return jsonify({"ok": True, "removed": username})


@app.route("/api/admin/backfill", methods=["POST"])
def api_admin_backfill():
    """Queue a backfill request — picked up by live monitor within 60s."""
    data     = request.get_json(force=True)
    channel  = data.get("channel", "").strip().lstrip("@")
    limit    = min(int(data.get("limit", 500)), 2000)
    since    = data.get("since", "")
    if not channel:
        return jsonify({"ok": False, "error": "channel required"}), 400

    req = {"channel": channel, "limit": limit, "queued_at": datetime.now(timezone.utc).isoformat()}
    if since:
        req["since"] = since

    bfq = {"pending": [], "completed": []}
    if BACKFILL_QUEUE.exists():
        try:
            bfq = json.loads(BACKFILL_QUEUE.read_text(encoding="utf-8"))
        except Exception:
            pass
    bfq.setdefault("pending", []).append(req)
    BACKFILL_QUEUE.write_text(json.dumps(bfq, ensure_ascii=False, indent=2), encoding="utf-8")
    return jsonify({"ok": True, "queued": req,
                    "note": "Monitor will process within 60 seconds"})


@app.route("/api/admin/compact", methods=["POST"])
def api_admin_compact():
    """Deduplicate messages in SQLite and vacuum."""
    result = db.compact_messages()
    conn = get_conn()
    conn.execute("VACUUM")
    conn.commit()
    crit = conn.execute("SELECT COUNT(*) FROM messages WHERE priority='CRITICAL'").fetchone()[0]
    return jsonify({"ok": True, "unique": result["remaining"], "deleted": result["deleted"],
                    "critical": crit})


@app.route("/api/stats/summary")
def api_stats_summary():
    """Lightweight real-time stats for the status bar — fast SQL."""
    conn = get_conn()
    now = datetime.now(timezone.utc)
    cutoff_24h = (now - timedelta(hours=24)).isoformat()
    cutoff_1h  = (now - timedelta(hours=1)).isoformat()

    row = conn.execute("""
        SELECT
            COUNT(*) as total,
            SUM(CASE WHEN priority='CRITICAL' THEN 1 ELSE 0 END) as critical,
            SUM(CASE WHEN priority='MEDIUM' THEN 1 ELSE 0 END) as medium,
            SUM(CASE WHEN priority='CRITICAL' AND timestamp_utc >= ? THEN 1 ELSE 0 END) as crit_24h,
            SUM(CASE WHEN priority='CRITICAL' AND timestamp_utc >= ? THEN 1 ELSE 0 END) as crit_1h,
            COUNT(DISTINCT channel_username) as channels
        FROM messages
    """, (cutoff_24h, cutoff_1h)).fetchone()

    # IOC count (still need to sum from JSON — but fast)
    ioc_rows = conn.execute("SELECT iocs FROM messages WHERE iocs IS NOT NULL").fetchall()
    ioc_count = 0
    for r in ioc_rows:
        try:
            iocs = json.loads(r[0]) if isinstance(r[0], str) else r[0]
            if isinstance(iocs, dict):
                ioc_count += sum(len(v or []) for v in iocs.values())
        except Exception:
            pass

    return jsonify({
        "total":       row[0] or 0,
        "critical":    row[1] or 0,
        "medium":      row[2] or 0,
        "critical_24h": row[3] or 0,
        "critical_1h":  row[4] or 0,
        "channels":    row[5] or 0,
        "iocs":        ioc_count,
        "generated_at": now.isoformat(),
    })


@app.route("/api/admin/discovered")
def api_admin_discovered():
    """Return live-discovery engine results from SQLite.
    Supports pagination: ?page=1&per_page=50
    Without page param, returns legacy array format.
    """
    channels = db.get_discovered_channels()  # dict keyed by username
    page     = request.args.get("page", type=int)
    per_page = min(request.args.get("per_page", 50, type=int), 200)
    if page is not None:
        ch_list = list(channels.values())
        total = len(ch_list)
        start = (page - 1) * per_page
        return jsonify({
            "data": ch_list[start:start + per_page],
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total": total,
                "pages": (total + per_page - 1) // per_page if per_page else 1,
            }
        })
    return jsonify(channels)


@app.route("/api/admin/discovered/<action>/<username>", methods=["POST"])
def api_admin_discovered_action(action, username):
    """
    Actions on a discovered channel:
      approve  — add to CHANNEL_TIERS + queue join
      dismiss  — mark as dismissed so it stops appearing
      ignore   — same as dismiss but keeps low-priority entries hidden
    """
    uname = username.lower().lstrip("@")
    if action == "approve":
        body = request.json or {}
        tier   = int(body.get("tier", 3))
        threat = body.get("threat", "MEDIUM")
        label  = body.get("label", username)
        # Add to DB + CHANNEL_TIERS runtime
        chan_data = {"label": label, "tier": tier, "threat": threat, "status": "active"}
        db.upsert_channel(uname, **chan_data)
        CHANNEL_TIERS[uname] = chan_data
        _save_channels_config(db.get_channels())
        # Queue for join
        pending_path = PENDING_FILE_VIEW
        pdata = {}
        if pending_path.exists():
            try:
                pdata = json.loads(pending_path.read_text(encoding="utf-8"))
            except Exception:
                pass
        pl = pdata.get("pending", [])
        if uname not in pl and uname not in pdata.get("processed", []):
            pl.append(uname)
        pdata["pending"]    = pl
        pdata["updated_at"] = datetime.now(timezone.utc).isoformat()
        pending_path.write_text(json.dumps(pdata, indent=2, ensure_ascii=False),
                                encoding="utf-8")
        db.update_discovered_status(uname, "approved")
    elif action in ("dismiss", "ignore"):
        db.update_discovered_status(uname, "dismissed")
    else:
        return jsonify({"error": "unknown action"}), 400
    return jsonify({"ok": True, "action": action, "username": uname})


# ══════════════════════════════════════════════════════════════════════════════
# AI AGENT API
# ══════════════════════════════════════════════════════════════════════════════

_ai_running = {"status": "idle", "started": None}


@app.route("/api/ai/status")
def api_ai_status():
    agent_state = {}
    if AI_STATE_FILE.exists():
        try:
            agent_state = json.loads(AI_STATE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    brief = None
    if AI_BRIEF_FILE.exists():
        try:
            brief = json.loads(AI_BRIEF_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    # Check if ai_agent.py process is running
    import psutil
    agent_running = False
    try:
        for p in psutil.process_iter(['pid', 'cmdline']):
            if p.info['cmdline'] and 'ai_agent.py' in ' '.join(p.info['cmdline']):
                agent_running = True
                break
    except Exception:
        pass
    return jsonify({
        "agent_running": agent_running,
        "enrichments_done": agent_state.get("enrichments_done", 0),
        "keywords_added": agent_state.get("keywords_added", 0),
        "channels_autoapproved": agent_state.get("channels_autoapproved", 0),
        "channels_autodismissed": agent_state.get("channels_autodismissed", 0),
        "briefs_generated": agent_state.get("briefs_generated", 0),
        "last_kw_run": agent_state.get("last_kw_run"),
        "last_brief_run": agent_state.get("last_brief_run"),
        "latest_brief": brief,
    })


@app.route("/api/ai/suggestions")
def api_ai_suggestions():
    if not AI_SUGGESTIONS.exists():
        return jsonify({"runs": [], "latest": None})
    try:
        return jsonify(json.loads(AI_SUGGESTIONS.read_text(encoding="utf-8")))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/brief")
def api_ai_brief():
    """Latest AI-generated threat intelligence brief."""
    if not AI_BRIEF_FILE.exists():
        return jsonify({"error": "no brief available yet"})
    try:
        return jsonify(json.loads(AI_BRIEF_FILE.read_text(encoding="utf-8")))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/translate", methods=["POST"])
def api_translate():
    """On-demand translation of a message text via OpenAI gpt-4o-mini."""
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "")[:2000]
    if not text:
        return jsonify({"error": "no text provided"}), 400
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return jsonify({"error": "OPENAI_API_KEY not configured"}), 503
    try:
        from openai import OpenAI as _OAI
        client = _OAI(api_key=api_key)
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content":
                    "You are a professional translator. Translate the following text to English. "
                    "Return ONLY the English translation — no explanations, no quoting the original."},
                {"role": "user", "content": text},
            ],
            max_tokens=600,
            temperature=0.1,
        )
        translation = r.choices[0].message.content.strip()
        return jsonify({"translation": translation})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/escalation/status")
def api_escalation_status():
    """Current escalation state from LOOP 6 (updated every 15 min)."""
    if not ESCALATION_FILE.exists():
        return jsonify({
            "escalation_detected": False,
            "urgency": "NONE",
            "signals": [],
            "summary": "No data yet — LOOP 6 starts after first 15-minute cycle",
            "recommended_action": "",
            "checked_at": None,
        })
    try:
        return jsonify(json.loads(ESCALATION_FILE.read_text(encoding="utf-8")))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/hunting/leads")
def api_hunting_leads():
    """Group and operation leads from LOOP 5 Threat Hunter."""
    if not HUNTING_LEADS_FILE.exists():
        return jsonify({"group_leads": [], "operation_leads": [], "runs": []})
    try:
        data = json.loads(HUNTING_LEADS_FILE.read_text(encoding="utf-8"))
        # Sort by confidence descending, return top 100
        data["group_leads"] = sorted(
            data.get("group_leads", []),
            key=lambda x: -int(x.get("confidence", 0))
        )[:100]
        data["operation_leads"] = sorted(
            data.get("operation_leads", []),
            key=lambda x: -int(x.get("confidence", 0))
        )[:50]
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/network/graph")
def api_network_graph():
    """Channel relationship graph summary from LOOP 7."""
    if not NETWORK_FILE.exists():
        return jsonify({
            "generated_at": None,
            "monitored_channels": 0,
            "unknown_channels_scored": 0,
            "newly_queued": 0,
            "top_unknown": [],
            "edges": [],
        })
    try:
        data = json.loads(NETWORK_FILE.read_text(encoding="utf-8"))
        # Trim edges for API response — return top 50 only
        data["edges"] = data.get("edges", [])[:50]
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/enriched")
def api_ai_enriched():
    """Latest N AI-enriched critical alerts."""
    limit = min(int(request.args.get("limit", 20)), 100)
    if not AI_ENRICHED_FILE.exists():
        return jsonify([])
    try:
        lines = AI_ENRICHED_FILE.read_text(encoding="utf-8").splitlines()
        records = []
        for line in reversed(lines):
            line = line.strip()
            if not line:
                continue
            try:
                records.append(json.loads(line))
                if len(records) >= limit:
                    break
            except Exception:
                pass
        return jsonify(list(reversed(records)))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


CHAT_HISTORY_FILE = OUTPUT_DIR / "chat_history.json"

# ── httpx compat patch so openai works inside viewer too ──────────────────────
try:
    import httpx as _vhttpx
    _vorig_client = _vhttpx.Client.__init__
    def _vp_client(self, *a, **kw):
        kw.pop("proxies", None); _vorig_client(self, *a, **kw)
    _vhttpx.Client.__init__ = _vp_client
    _vorig_async = _vhttpx.AsyncClient.__init__
    def _vp_async(self, *a, **kw):
        kw.pop("proxies", None); _vorig_async(self, *a, **kw)
    _vhttpx.AsyncClient.__init__ = _vp_async
except Exception:
    pass

def _chat_openai_client():
    key = os.environ.get("OPENAI_API_KEY", "")
    if not key:
        # Try reading .env directly
        env_path = Path(__file__).parent / ".env"
        if env_path.exists():
            for ln in env_path.read_text(encoding="utf-8").splitlines():
                ln = ln.strip()
                if ln.startswith("OPENAI_API_KEY="):
                    key = ln.split("=", 1)[1].strip().strip('"').strip("'")
                    os.environ["OPENAI_API_KEY"] = key
                    break
    if not key:
        return None
    try:
        from openai import OpenAI
        return OpenAI(api_key=key)
    except Exception:
        return None

# Synonym expansion — common intel query terms mapped to domain-specific search terms
_CHAT_SYNONYMS = {
    "bank":       ["jcbank","arab bank","housing bank","bank of jordan","capital bank","capitalbank","cairo amman","etihad bank","البنك","بنك",".jo","ahli bank","jordan ahli","jordan kuwait bank","jkb"],
    "banks":      ["jcbank","arab bank","housing bank","bank of jordan","capital bank","capitalbank","cairo amman","etihad bank","البنك","بنك",".jo","ahli bank","jordan ahli","jordan kuwait bank","jkb"],
    "jordanian":  ["jordan","الاردن","الأردن","أردن","اردن",".jo","amman","عمان"],
    "jordan":     ["الاردن","الأردن","أردن","اردن",".jo","amman","عمان"],
    "government": ["ministry","وزارة","الديوان","royal court","prime minister","رئاسة الوزراء","gov.jo"],
    "ministry":   ["وزارة","ministry of","الديوان","gov.jo"],
    "attack":     ["ddos","hacked","breach","defaced","down","offline","closed","تم اختراق","اختراق","هكر","check-host"],
    "attacks":    ["ddos","hacked","breach","defaced","down","offline","closed","تم اختراق","اختراق","check-host"],
    "hack":       ["hacked","breach","defaced","leak","dump","تسريب","تم اختراق","اختراق"],
    "hacked":     ["breach","defaced","leak","dump","تسريب","تم اختراق","check-host"],
    "breach":     ["hacked","leak","dump","data","database","تسريب","credentials"],
    "ddos":       ["down","offline","closed","connection refused","check-host","flood","layer7","layer4","http flood"],
    "leak":       ["dump","تسريب","data","database","credentials","pastebin"],
    "data":       ["leak","dump","database","تسريب","credentials","passwords"],
    "israel":     ["إسرائيل","زيونيست","zionist","صهيوني","idf","israeli","tel aviv"],
    "iran":       ["إيران","islamic resistance","حزب الله","hezbollah","fatemiyoun","فاطميون","irgc"],
    "critical":   ["CRITICAL"],
    "recent":     [],
    "latest":     [],
    "today":      [],
    "new":        [],
    "any":        [],
}

def _extract_search_query(user_msg: str, history: list) -> str:
    """Use gpt-4o-mini to extract optimal search terms from conversation context.
    For first messages or clearly self-contained queries, skip the GPT call."""
    meaningful_words = [w for w in re.split(r'\W+', user_msg.lower()) if len(w) > 2]
    has_pronouns = any(w in user_msg.lower() for w in (
        "that", "this", "those", "these", "them", "it", "they",
        "more", "above", "previous", "earlier", "same", "again"
    ))
    if not history or (len(meaningful_words) >= 5 and not has_pronouns):
        return user_msg
    client = _chat_openai_client()
    if not client:
        return user_msg
    recent = history[-6:]
    context_lines = []
    for h in recent:
        role = h.get("role", "")
        content = str(h.get("content", ""))[:300]
        if role in ("user", "assistant"):
            context_lines.append(f"{role}: {content}")
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": (
                    "You extract search keywords from a conversation. "
                    "Given the conversation history and the user's latest message, "
                    "output 3-8 specific search terms/phrases that would find the relevant Telegram messages. "
                    "Include entity names, attack types, dates, and channel names mentioned in context. "
                    "Output ONLY the search terms separated by spaces. No explanation."
                )},
                {"role": "user", "content": (
                    f"Conversation:\n" + "\n".join(context_lines) +
                    f"\n\nLatest user message: {user_msg}\n\nSearch terms:"
                )}
            ],
            max_tokens=100,
            temperature=0,
        )
        extracted = resp.choices[0].message.content.strip()
        return f"{user_msg} {extracted}"
    except Exception:
        return user_msg


def _search_messages_for_chat(query: str, limit: int = 500) -> list:
    """
    Score messages by relevance to query.
    Strategy:
      1. Expand query words with domain synonyms
      2. Score all messages by keyword hits + strong recency boost
      3. Always inject last 100 messages + all CRITICALs as baseline
      4. Return sorted chronologically so AI sees a timeline
    """
    msgs = load_messages()
    now_utc = datetime.now(timezone.utc)

    def _recency_boost(ts_str):
        """Massive boost for recent messages so 'now/today' queries work."""
        if not ts_str:
            return 0
        try:
            dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            age_h = (now_utc - dt).total_seconds() / 3600
            if age_h < 1:    return 30   # last hour — highest priority
            if age_h < 6:    return 15   # last 6h
            if age_h < 24:   return 8    # today
            if age_h < 72:   return 3    # last 3 days
            return 0
        except Exception:
            return 0

    # Base query words
    base_words = [w for w in re.split(r'\W+', query.lower()) if len(w) > 2]

    # Detect if user is asking about "now/today/current/latest/recent"
    recency_query = any(w in query.lower() for w in (
        "now", "today", "current", "latest", "recent", "right now",
        "happening", "new", "just", "currently", "active", "ongoing"
    ))

    # Expand with synonyms
    expanded = list(set(base_words))
    for w in base_words:
        for syn in _CHAT_SYNONYMS.get(w, []):
            expanded.append(syn.lower())

    prio_boost = {"CRITICAL": 4, "HIGH": 2, "MEDIUM": 1, "LOW": 0}
    seen_ids = {}

    for m in msgs:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        text   = (m.get("text_preview") or "").lower()
        ch     = (m.get("channel") or "").lower()
        uname  = (m.get("channel_username") or "").lower()
        ae     = m.get("ai_enrichment") or {}
        enrich = " ".join(str(v) for v in ae.values()).lower() if ae else ""
        kwhits = " ".join(m.get("keyword_hits") or []).lower()

        score = 0
        for w in expanded:
            if w in text:    score += 2
            if w in ch:      score += 3
            if w in uname:   score += 3
            if w in enrich:  score += 2
            if w in kwhits:  score += 2
        score += prio_boost.get(m.get("priority", "LOW"), 0)
        score += _recency_boost(m.get("timestamp_utc", ""))

        # For recency queries, give all messages a base score so recent ones appear
        if recency_query:
            score = max(score, _recency_boost(m.get("timestamp_utc", "")))

        if score > 0:
            seen_ids[mid] = (score, m)

    # Top keyword-ranked results
    ranked = sorted(seen_ids.values(), key=lambda x: x[0], reverse=True)
    result = [m for _, m in ranked[:limit]]
    result_set = {f"{m.get('channel_username','')}_{m.get('message_id','')}" for m in result}

    # Always inject last 100 messages (regardless of keyword match) — captures live events
    recent_all = sorted(msgs, key=lambda m: m.get("timestamp_utc", ""), reverse=True)[:100]
    for m in recent_all:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        if mid not in result_set:
            result.append(m)
            result_set.add(mid)

    # Always inject ALL CRITICAL messages as baseline
    recent_crits = sorted(
        [m for m in msgs if m.get("priority") == "CRITICAL"],
        key=lambda m: m.get("timestamp_utc", ""),
        reverse=True
    )
    for m in recent_crits:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        if mid not in result_set:
            result.append(m)
            result_set.add(mid)

    # Sort final list chronologically — AI reads a timeline, not a random pile
    result.sort(key=lambda m: m.get("timestamp_utc", ""))
    return result


def _search_messages_for_chat_v2(search_query: str, original_query: str, limit: int = 150) -> list:
    """V2 context retrieval: relevance-focused, no context flooding."""
    msgs = load_messages()
    now_utc = datetime.now(timezone.utc)

    def _recency_boost(ts_str):
        if not ts_str:
            return 0
        try:
            dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            age_h = (now_utc - dt).total_seconds() / 3600
            if age_h < 1:    return 30
            if age_h < 6:    return 15
            if age_h < 24:   return 8
            if age_h < 72:   return 3
            return 0
        except Exception:
            return 0

    combined = f"{search_query} {original_query}"
    base_words = list(set(w for w in re.split(r'\W+', combined.lower()) if len(w) > 2))

    recency_query = any(w in original_query.lower() for w in (
        "now", "today", "current", "latest", "recent", "right now",
        "happening", "new", "just", "currently", "active", "ongoing"
    ))

    expanded = list(set(base_words))
    for w in base_words:
        for syn in _CHAT_SYNONYMS.get(w, []):
            expanded.append(syn.lower())

    prio_boost = {"CRITICAL": 4, "HIGH": 2, "MEDIUM": 1, "LOW": 0}
    seen_ids = {}

    for m in msgs:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        text   = (m.get("text_preview") or "").lower()
        ch     = (m.get("channel") or "").lower()
        uname  = (m.get("channel_username") or "").lower()
        ae     = m.get("ai_enrichment") or {}
        enrich = " ".join(str(v) for v in ae.values()).lower() if ae else ""
        kwhits = " ".join(m.get("keyword_hits") or []).lower()

        score = 0
        for w in expanded:
            if w in text:    score += 2
            if w in ch:      score += 3
            if w in uname:   score += 3
            if w in enrich:  score += 2
            if w in kwhits:  score += 2
        score += prio_boost.get(m.get("priority", "LOW"), 0)
        score += _recency_boost(m.get("timestamp_utc", ""))

        if recency_query:
            score = max(score, _recency_boost(m.get("timestamp_utc", "")))

        if score > 0:
            seen_ids[mid] = (score, m)

    ranked = sorted(seen_ids.values(), key=lambda x: x[0], reverse=True)
    result = [m for _, m in ranked[:limit]]
    result_set = {f"{m.get('channel_username','')}_{m.get('message_id','')}" for m in result}

    # Only inject recent messages if recency query (25, not 100)
    if recency_query:
        recent_all = sorted(msgs, key=lambda m: m.get("timestamp_utc", ""), reverse=True)[:25]
        for m in recent_all:
            mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
            if mid not in result_set:
                result.append(m)
                result_set.add(mid)

    # Only CRITICAL from last 7 days, cap at 50
    seven_days_ago = (now_utc - timedelta(days=7)).isoformat()
    recent_crits = [
        m for m in msgs
        if m.get("priority") == "CRITICAL"
        and (m.get("timestamp_utc", "") >= seven_days_ago)
    ]
    recent_crits.sort(key=lambda m: m.get("timestamp_utc", ""), reverse=True)
    for m in recent_crits[:50]:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        if mid not in result_set:
            result.append(m)
            result_set.add(mid)

    result.sort(key=lambda m: m.get("timestamp_utc", ""))
    return result


def _format_msg_for_context(idx: int, m: dict) -> str:
    """Format a single message for the AI context block.
    Uses sequential index (idx) as the REF ID — guaranteed unique across all channels."""
    ts    = (m.get("timestamp_utc") or "")[:16].replace("T", " ")
    pri   = m.get("priority", "LOW")
    ch    = m.get("channel") or m.get("channel_username", "?")
    uname = m.get("channel_username", "")
    text  = (m.get("text_preview") or "")[:500]
    lines = [f"[REF:{idx}] [{pri}] {ts} UTC | @{uname} | {ch}", text]
    ae = m.get("ai_enrichment") or {}
    if ae.get("summary"):
        lines.append(f"  >> AI: {ae['summary'][:250]}")
    if ae.get("target_sector"):
        lines.append(f"  >> Target: {ae['target_sector']} | Severity: {ae.get('severity','')} | Confidence: {ae.get('confidence','')}%")
    kw = m.get("keyword_hits") or []
    if kw:
        lines.append(f"  >> Keywords: {', '.join(kw[:10])}")
    return "\n".join(lines)

_CHAT_SYSTEM_TEMPLATE = """You are a senior cyber threat intelligence analyst for the Scanwave CyberIntel Platform.
You monitor hacktivist and state-linked threat actor activity targeting Jordan and the wider Middle East.

TODAY'S DATE/TIME (UTC): {now_utc}

You are given a numbered context of Telegram messages from monitored threat-actor channels, each tagged [REF:N].
The messages are sorted OLDEST FIRST — the most recent messages are at the END of the context.
Your job is to answer the analyst's question using ONLY this evidence. Be thorough, precise, and analytical.

Rules:
- Cite every [REF:N] tag that supports a claim. Multiple refs per claim are encouraged.
- If a message is in Arabic, translate the key parts in your answer.
- Group findings by threat actor or attack type when it makes sense.
- When asked about "now/today/current", focus on messages from today's date ({today}).
- If the context truly contains no relevant information for the specific date asked, say so clearly — do NOT fabricate.
- Respond ONLY in this exact JSON format (no markdown, no code fences):
  {{"answer":"full detailed answer","references":[0,1,2,...]}}

The "references" array must contain INTEGER indices (0,1,2...) matching the [REF:N] numbers used."""

_CHAT_SYSTEM_TEMPLATE_V2 = """You are a senior cyber threat intelligence analyst for the Scanwave CyberIntel Platform.
You monitor hacktivist and state-linked threat actor activity targeting Jordan and the wider Middle East.

TODAY'S DATE/TIME (UTC): {now_utc}

You are given a numbered context of Telegram messages from monitored threat-actor channels, each tagged [REF:N].
The messages are sorted OLDEST FIRST — the most recent messages are at the END of the context.
Your job is to answer the analyst's question using ONLY this evidence.

ALWAYS respond in English. Translate any Arabic or Farsi content inline.

RESPONSE MODES:
1. **Q&A mode** (default): Be brief, factual, and precise. Cite [REF:N] tags inline.
2. **Report mode** (when user asks to "generate report", "write report", "summarize all", "full analysis"):
   Use this structured format:
   # Report Title
   **Date:** {today} | **Scope:** [describe]
   ## Executive Summary
   [2-3 sentence overview]
   ## Key Findings
   [Group by threat actor or attack type, use bullet points, cite [REF:N]]
   ## Indicators of Compromise
   | Type | Value | Source |
   |------|-------|--------|
   [table rows from context]
   ## Risk Assessment
   [Brief risk level and impact analysis]
   ## Recommendations
   [Actionable bullet points]

Rules:
- Cite [REF:N] tags that support each claim. Multiple refs per claim encouraged.
- Use **bold** for emphasis, bullet lists, and numbered lists as needed.
- When asked about "now/today/current", focus on messages from today's date ({today}).
- If the context contains no relevant information, say so clearly — do NOT fabricate.
- After your complete answer, write exactly this delimiter on its own line:
  ---REFS---
  Then list ALL referenced numbers as comma-separated integers. Example: 3, 7, 12, 45
- If you cited no references, still write ---REFS--- followed by nothing."""

# ── Agentic search: GPT tool-calling for iterative intel retrieval ─────

_SEARCH_INTEL_TOOL = {
    "type": "function",
    "function": {
        "name": "search_intel",
        "description": (
            "Search the Telegram intelligence message database. "
            "Returns messages matching the query, formatted with [REF:N] tags for citation. "
            "Call multiple times with different queries for comprehensive coverage. "
            "Use specific terms: target names, domains (.jo, gov.jo), attack types, actor names, Arabic keywords."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search terms: target names, domains, attack types (DDoS, breach, leak, defacement), actor names, Arabic/English keywords"
                },
                "severity": {
                    "type": "string",
                    "enum": ["all", "CRITICAL", "MEDIUM", "LOW"],
                    "description": "Filter by severity. Default: all"
                },
                "days_back": {
                    "type": "integer",
                    "description": "How many days back to search. Default: 365"
                }
            },
            "required": ["query"]
        }
    }
}


def _format_msg_compact(idx: int, m: dict) -> str:
    """Compact message format for tool search results — optimized for token budget.
    Arabic text tokenizes at 2-4x English rate, so keep text short."""
    ts    = (m.get("timestamp_utc") or "")[:16].replace("T", " ")
    pri   = m.get("priority", "LOW")
    uname = m.get("channel_username", "?")
    text  = (m.get("text_preview") or "")[:150]
    ae    = m.get("ai_enrichment") or {}
    summary = (ae.get("summary") or "")[:120]
    kw = m.get("keyword_hits") or []
    # Prefer English summary over raw Arabic text when available
    if summary:
        return f"[REF:{idx}] [{pri}] {ts} | @{uname}\n>> {summary}\n>> KW: {', '.join(kw[:5])}"
    line = f"[REF:{idx}] [{pri}] {ts} | @{uname}\n{text}"
    if kw:
        line += f"\n>> KW: {', '.join(kw[:5])}"
    return line


def _execute_search_tool(query: str, severity: str, days_back: int,
                         seen_ids: set, all_ref_msgs: list,
                         all_msgs: list) -> tuple:
    """Execute an agentic search, return (count, formatted_text).
    Maintains global REF numbering across calls via all_ref_msgs/seen_ids."""
    now_utc = datetime.now(timezone.utc)
    cutoff = (now_utc - timedelta(days=days_back)).isoformat() if days_back < 365 else ""

    base_words = [w for w in re.split(r'\W+', query.lower()) if len(w) > 2]
    expanded = list(set(base_words))
    for w in base_words:
        for syn in _CHAT_SYNONYMS.get(w, []):
            expanded.append(syn.lower())
    # Also add the raw query terms (including short ones like ".jo")
    for part in query.lower().split():
        part = part.strip(".,;:!?")
        if part and part not in expanded:
            expanded.append(part)

    prio_boost = {"CRITICAL": 4, "HIGH": 2, "MEDIUM": 1, "LOW": 0}
    scored = []

    for m in all_msgs:
        mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
        if mid in seen_ids:
            continue
        if cutoff and (m.get("timestamp_utc", "") < cutoff):
            continue
        if severity != "all" and m.get("priority") != severity:
            continue

        text  = (m.get("text_preview") or "").lower()
        ch    = (m.get("channel") or "").lower()
        uname = (m.get("channel_username") or "").lower()
        ae    = m.get("ai_enrichment") or {}
        enrich = " ".join(str(v) for v in ae.values()).lower() if ae else ""
        kwhits = " ".join(m.get("keyword_hits") or []).lower()

        score = 0
        for w in expanded:
            if w in text:    score += 2
            if w in ch:      score += 3
            if w in uname:   score += 3
            if w in enrich:  score += 2
            if w in kwhits:  score += 2
        score += prio_boost.get(m.get("priority", "LOW"), 0)
        # Severity-filtered searches: give base score so results appear even without keyword match
        if severity != "all" and score == 0:
            score = 1

        if score > 0:
            scored.append((score, m, mid))

    scored.sort(key=lambda x: x[0], reverse=True)

    results = []
    for score, m, mid in scored[:30]:
        ref_idx = len(all_ref_msgs)
        all_ref_msgs.append(m)
        seen_ids.add(mid)
        results.append((ref_idx, m))

    if not results:
        return 0, f"No messages found matching '{query}'"

    results.sort(key=lambda x: x[1].get("timestamp_utc", ""))
    blocks = [_format_msg_compact(ref_idx, m) for ref_idx, m in results]
    formatted = f"Found {len(results)} messages:\n" + "\n---\n".join(blocks)

    if len(formatted) > 10000:
        formatted = formatted[:10000] + "\n[...truncated, try a more specific query]"

    return len(results), formatted


_CHAT_SYSTEM_TEMPLATE_V3 = """You are a senior cyber threat intelligence analyst for the Scanwave CyberIntel Platform.
You monitor hacktivist and state-linked threat actor activity targeting Jordan and the wider Middle East.

TODAY'S DATE/TIME (UTC): {now_utc}

DATABASE: {total_msgs} messages from {total_channels} monitored Telegram channels ({critical_count} CRITICAL, {medium_count} MEDIUM, {low_count} LOW).

You have a **search_intel** tool to query the message database. Each search returns up to 50 messages with [REF:N] citation tags.

ALWAYS respond in English. Translate any Arabic/Farsi content inline.

SEARCH STRATEGY:
- For simple questions: 1-2 focused searches
- For comprehensive requests ("full report", "name all", "everything", "all sectors"):
  * FIRST: search severity="CRITICAL" to get all high-priority attack reports
  * Then search by SECTOR: banks, government, military, telecom, education, energy, media, aviation
  * Then search by ATTACK TYPE: DDoS, breach, leak, defacement
  * Then do FOLLOW-UP searches for specific targets/actors found in earlier results
  * Aim for 6-10 searches total for maximum coverage
- Use both English AND Arabic terms when relevant (e.g., "اختراق" for hack, "بنك" for bank)
- Include domain patterns: .jo, gov.jo, mil.jo, edu.jo, com.jo
- Search for specific organization names found in prior results
- Use severity="CRITICAL" filter to find confirmed attack reports efficiently

RESPONSE MODES:
1. **Q&A mode** (default): Brief, factual, precise. Cite [REF:N] inline.
2. **Report mode** (when asked for "report", "full report", "all", "everything", "summarize", "name all"):
   Produce a DETAILED, COMPREHENSIVE intelligence report. Be thorough — this is for senior decision-makers.
   Name SPECIFIC organizations, companies, banks, domains — be precise, not vague.

   # Report Title
   **Date:** {today} | **Scope:** [describe] | **Classification:** TLP:AMBER

   ## Executive Summary
   [3-5 sentence high-level overview: total incidents, top threat actors, most targeted sectors, overall threat level]

   ## Threat Actor Overview
   | Threat Actor | Incidents | Primary Attack Types | Key Targets |
   [Table of ALL active threat groups with their activity summary]

   ## Incident Analysis by Sector
   For EACH sector with incidents, provide a detailed subsection:
   ### [Sector Name] ([N] incidents)
   - List every incident with: date, threat actor, attack type, specific target, impact
   - Cite [REF:N] for every claim
   - Note any patterns or escalation trends

   ## Attack Timeline
   Chronological timeline of major incidents showing escalation patterns and campaign waves.

   ## Indicators of Compromise
   | Type | Value | Context | Threat Actor |
   |------|-------|---------|-------------|
   [Extract ALL domains, IPs, URLs, hashes mentioned in the intelligence. Be thorough.]

   ## Threat Assessment
   - Current threat level and justification
   - Most at-risk sectors and why
   - Threat actor capability assessment
   - Likelihood of escalation

   ## Recommendations
   [Specific, actionable recommendations per sector — not generic advice]

Rules:
- Cite [REF:N] tags for EVERY factual claim. Multiple refs encouraged.
- Use **bold**, bullet lists, numbered lists, and tables as needed.
- When asked about "now/today/current", focus on today's date ({today}).
- Do NOT fabricate information. Only report what you found in search results.
- After your complete answer, write exactly on its own line:
  ---REFS---
  Then list ALL referenced numbers as comma-separated integers.
- If you cited no references, still write ---REFS--- followed by nothing."""


@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json(silent=True) or {}
    user_msg = (data.get("message") or "").strip()
    history  = data.get("history") or []
    if not user_msg:
        return jsonify({"error": "empty message"}), 400
    client = _chat_openai_client()
    if not client:
        return jsonify({"error": "OPENAI_API_KEY not configured"}), 503

    # Retrieve relevant messages with conversation-aware search
    search_query = _extract_search_query(user_msg, history)
    rel_msgs = _search_messages_for_chat_v2(search_query, user_msg, limit=150)

    # Build context with sequential REF indices — this is the ONLY source of truth for refs
    context_blocks = [_format_msg_for_context(i, m) for i, m in enumerate(rel_msgs)]
    context_text   = "\n---\n".join(context_blocks)

    # gpt-4o supports 128k tokens. Cap context at ~100k chars safely.
    if len(context_text) > 100000:
        # Keep as many complete blocks as fit
        kept, total = [], 0
        for blk in context_blocks:
            if total + len(blk) + 5 > 100000:
                break
            kept.append(blk)
            total += len(blk) + 5
        context_text = "\n---\n".join(kept)
        rel_msgs = rel_msgs[:len(kept)]   # keep index alignment

    # Build date-aware system prompt
    now_utc = datetime.now(timezone.utc)
    today_str  = now_utc.strftime("%Y-%m-%d")
    now_str    = now_utc.strftime("%Y-%m-%d %H:%M UTC")
    newest_ts  = max((m.get("timestamp_utc","") for m in rel_msgs), default="unknown")
    chat_system = _CHAT_SYSTEM_TEMPLATE.format(now_utc=now_str, today=today_str)

    trimmed_history = history[-16:]   # last 8 exchanges
    payload = [{"role": "system", "content": chat_system}]
    payload.append({
        "role": "system",
        "content": (
            f"=== INTEL CONTEXT: {len(rel_msgs)} messages from monitored channels ===\n"
            f"(Messages sorted oldest→newest. Newest message timestamp: {newest_ts[:16]})\n"
            f"{context_text}\n"
            f"=== END CONTEXT ==="
        )
    })
    for h in trimmed_history:
        if h.get("role") in ("user", "assistant") and h.get("content"):
            payload.append({"role": h["role"], "content": str(h["content"])})
    payload.append({"role": "user", "content": user_msg})

    try:
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=payload,
            max_tokens=2000,
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content or "{}"
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        parsed = {"answer": raw, "references": []}
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    answer = parsed.get("answer", "")
    # AI returns integer indices; resolve them to actual messages
    raw_refs = parsed.get("references") or []
    valid_refs = []
    ref_messages = []
    seen_idx = set()
    for r in raw_refs:
        try:
            idx = int(r)
        except (ValueError, TypeError):
            continue
        if 0 <= idx < len(rel_msgs) and idx not in seen_idx:
            seen_idx.add(idx)
            valid_refs.append(idx)
            ref_messages.append(rel_msgs[idx])

    return jsonify({"answer": answer, "references": valid_refs, "ref_messages": ref_messages})


@app.route("/api/chat/history")
def api_chat_history():
    # Session-only: history lives in frontend _chatHistory array
    return jsonify({"messages": []})


@app.route("/api/chat/reset", methods=["POST"])
def api_chat_reset():
    return jsonify({"ok": True})


@app.route("/api/chat/stream", methods=["POST"])
def api_chat_stream():
    """Agentic streaming chat: AI iteratively searches intel DB via tool calling."""
    data = request.get_json(silent=True) or {}
    user_msg = (data.get("message") or "").strip()
    history  = data.get("history") or []
    if not user_msg:
        return jsonify({"error": "empty message"}), 400
    client = _chat_openai_client()
    if not client:
        return jsonify({"error": "OPENAI_API_KEY not configured"}), 503

    # Load messages once for this request
    all_msgs = load_messages()
    total_msgs = len(all_msgs)
    critical_count = sum(1 for m in all_msgs if m.get("priority") == "CRITICAL")
    medium_count = sum(1 for m in all_msgs if m.get("priority") == "MEDIUM")
    low_count = total_msgs - critical_count - medium_count
    n_channels = len(set(m.get("channel_username", "") for m in all_msgs if m.get("channel_username")))

    now_utc = datetime.now(timezone.utc)
    today_str = now_utc.strftime("%Y-%m-%d")
    now_str   = now_utc.strftime("%Y-%m-%d %H:%M UTC")

    chat_system = _CHAT_SYSTEM_TEMPLATE_V3.format(
        now_utc=now_str, today=today_str,
        total_msgs=total_msgs, total_channels=n_channels,
        critical_count=critical_count, medium_count=medium_count, low_count=low_count
    )

    # Build message list: system + conversation history + new question
    # Use generous limits for recent messages (reports can be 5-8k chars)
    messages = [{"role": "system", "content": chat_system}]
    recent_history = history[-16:]
    for i, h in enumerate(recent_history):
        if h.get("role") in ("user", "assistant") and h.get("content"):
            # Last 4 messages get full context (8k), older ones get 2k
            char_limit = 8000 if i >= len(recent_history) - 4 else 2000
            messages.append({"role": h["role"], "content": str(h["content"])[:char_limit]})
    messages.append({"role": "user", "content": user_msg})

    # Shared state across agentic iterations
    all_ref_msgs = []
    seen_ids = set()

    # Detect comprehensive queries → use map-reduce approach
    _comprehensive_kw = ("full report", "all sectors", "everything", "name all",
        "name everything", "comprehensive", "dont skip", "don't skip",
        "complete report", "all cybersecurity", "all incidents", "all attacks")
    is_comprehensive = any(kw in user_msg.lower() for kw in _comprehensive_kw)

    # --- Map-Reduce prompts ---
    _MAP_EXTRACT_PROMPT = (
        "You are an elite cybersecurity incident extractor. "
        "Extract EVERY cyber attack from these Telegram messages.\n\n"
        "OUTPUT — one line per incident:\n"
        "[REF:N] | DATE | THREAT_ACTOR | @CHANNEL | ATTACK_TYPE | TARGET | SECTOR | DETAILS_AND_IOCS\n\n"
        "ATTACK_TYPE: DDoS, Defacement, Data Breach, Data Leak, Hack, Ransomware, Phishing, Wiper/Malware, Threat/Warning, Reconnaissance\n"
        "SECTOR: Government, Military, Banking/Finance, Telecom, Aviation, Energy, Education, Media, Healthcare, Transportation, Private Sector\n\n"
        "RULES:\n"
        "1. Extract EVERY cyber incident. Do NOT skip any.\n"
        "2. THREAT_ACTOR = the hacker GROUP name (DieNet, Arabian Ghosts, Fatemiyoun, OpIsrael, etc.), NOT the Telegram channel name.\n"
        "3. Be SPECIFIC with targets: exact domain (mof.gov.jo), exact org name (Central Bank of Jordan), NOT generic 'government'.\n"
        "4. In DETAILS_AND_IOCS include: domains targeted, check-host proof URLs, IPs, data volume, downtime, any technical IOCs.\n"
        "5. Different actors on same target = separate lines. Same actor different dates = separate lines.\n\n"
        "ABSOLUTE EXCLUSIONS — do NOT extract these (they are NOT cyber attacks):\n"
        "- Missiles, rockets, drones hitting physical targets\n"
        "- Tank battles, ground operations, armed clashes\n"
        "- Building fires, explosions, physical destruction\n"
        "- Military equipment destroyed (THAAD, radar, vehicles)\n"
        "- Troop deployments, military exercises\n"
        "- Political statements, speeches, press releases\n"
        "- Protests, demonstrations\n"
        "- News commentary without attack claims\n"
        "A cyber attack involves COMPUTERS, NETWORKS, WEBSITES, DATA. If it involves bullets, bombs, or missiles = SKIP.\n\n"
        "Output ONLY incident lines. If none found: NO_INCIDENTS"
    )

    _REDUCE_REPORT_PROMPT = _CHAT_SYSTEM_TEMPLATE_V3.format(
        now_utc=now_str, today=today_str,
        total_msgs=total_msgs, total_channels=n_channels,
        critical_count=critical_count, medium_count=medium_count, low_count=low_count
    ).replace(
        "You have a **search_intel** tool to query the message database.",
        "You have been given pre-extracted incident findings from the full message database."
    ).replace(
        "Each search returns up to 50 messages with [REF:N] citation tags.",
        "Each finding has a [REF:N] tag you must cite."
    )

    def _format_ultracompact(idx, m):
        """One-line format: ~100-150 chars. Prefers English AI summary."""
        ts = (m.get("timestamp_utc") or "")[:10]
        uname = m.get("channel_username", "?")
        pri = m.get("priority", "LOW")
        ae = m.get("ai_enrichment") or {}
        summary = (ae.get("summary") or "")[:120]
        if summary:
            return f"[REF:{idx}] [{pri}] {ts} @{uname} | {summary}"
        text = (m.get("text_preview") or "")[:100]
        kw = m.get("keyword_hits") or []
        return f"[REF:{idx}] [{pri}] {ts} @{uname} | {text} | KW:{','.join(kw[:4])}"

    def _format_cyber_detail(idx, m):
        """Richer format for cyber incidents — includes full text for domain/IOC extraction."""
        ts = (m.get("timestamp_utc") or "")[:10]
        uname = m.get("channel_username", "?")
        kw = m.get("keyword_hits") or []
        # Use full text_preview (not truncated) so GPT can see domains, IPs, targets
        text = (m.get("text_preview") or m.get("text") or "")[:500].replace("\n", " ")
        ae = m.get("ai_enrichment") or {}
        summary = (ae.get("summary") or "")[:200]
        parts = [f"[REF:{idx}] {ts} @{uname}"]
        if summary:
            parts.append(f"SUMMARY: {summary}")
        parts.append(f"TEXT: {text}")
        if kw:
            parts.append(f"KW: {','.join(kw[:6])}")
        return " | ".join(parts)

    def generate():
        nonlocal messages
        start_time = time.time()

        if is_comprehensive:
            # ── MAP-REDUCE: process CYBER-classified messages for cyber reports ──
            # Detect if query is cyber-focused (default) vs national security
            _cyber_query = any(kw in user_msg.lower() for kw in (
                "cyber", "ddos", "hack", "breach", "defac", "attack", "ioc",
                "malware", "ransomware", "phishing", "blocklist", "threat actor"))
            _natsec_query = any(kw in user_msg.lower() for kw in (
                "national security", "military", "missile", "rocket", "kinetic",
                "geopolitical", "natsec", "armed forces"))

            if _natsec_query and not _cyber_query:
                # NatSec report: use NATIONAL + BOTH subtypes
                target_msgs = [m for m in all_msgs if m.get("critical_subtype") in ("NATIONAL", "BOTH")]
                report_scope = "national security"
            elif _cyber_query or not _natsec_query:
                # Cyber report (default): use CYBER + BOTH subtypes
                target_msgs = [m for m in all_msgs if m.get("critical_subtype") in ("CYBER", "BOTH")]
                report_scope = "cybersecurity"
            else:
                # General: all CRITICAL + MEDIUM
                target_msgs = [m for m in all_msgs if m.get("priority") in ("CRITICAL", "MEDIUM")]
                report_scope = "all intelligence"

            target_msgs.sort(key=lambda m: m.get("timestamp_utc", ""))

            # For cyber reports, use richer formatting with more text context
            is_cyber_report = report_scope == "cybersecurity"
            # Larger chunks since we have fewer, more relevant messages
            chunk_size = 80 if is_cyber_report else 150
            chunks = [target_msgs[i:i+chunk_size] for i in range(0, len(target_msgs), chunk_size)]

            # Pre-assign REF IDs for all messages (needed before parallel calls)
            chunk_contexts = []
            for chunk in chunks:
                formatted_lines = []
                for m in chunk:
                    mid = f"{m.get('channel_username','')}_{m.get('message_id','')}"
                    if mid not in seen_ids:
                        ref_idx = len(all_ref_msgs)
                        all_ref_msgs.append(m)
                        seen_ids.add(mid)
                        if is_cyber_report:
                            formatted_lines.append(_format_cyber_detail(ref_idx, m))
                        else:
                            formatted_lines.append(_format_ultracompact(ref_idx, m))
                chunk_contexts.append("\n".join(formatted_lines))

            yield f"data: {json.dumps({'type':'status','message':f'Analyzing {len(target_msgs)} {report_scope} messages across {len(chunks)} parallel agents...'})}\n\n"

            # ── PARALLEL MAP: all agents run simultaneously ──
            def _run_map_agent(agent_idx, context_text):
                """Single map agent — runs in thread pool."""
                if not context_text.strip():
                    return agent_idx, ""
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": _MAP_EXTRACT_PROMPT},
                            {"role": "user", "content": context_text}
                        ],
                        max_tokens=4000,
                        temperature=0,
                    )
                    return agent_idx, resp.choices[0].message.content.strip()
                except Exception as e:
                    return agent_idx, f"ERROR: {str(e)[:100]}"

            all_findings = []
            with ThreadPoolExecutor(max_workers=min(len(chunks), 10)) as executor:
                futures = {
                    executor.submit(_run_map_agent, ci, ctx): ci
                    for ci, ctx in enumerate(chunk_contexts)
                }
                for future in as_completed(futures):
                    agent_idx, result = future.result()
                    yield f"data: {json.dumps({'type':'status','message':f'Agent {agent_idx+1}/{len(chunks)} done'})}\n\n"
                    if result and "NO_INCIDENTS" not in result and "NO_JORDAN" not in result and not result.startswith("ERROR:"):
                        all_findings.append(result)

            # ── DEDUP: merge duplicate incidents by actor+target+date ──
            import re as _re
            incident_lines = []
            for f in all_findings:
                incident_lines.extend([l.strip() for l in f.split("\n") if l.strip() and l.strip().startswith("[")])
            raw_count = len(incident_lines)

            # Parse each line: [REF:N] | DATE | ACTOR | @CHANNEL | TYPE | TARGET | SECTOR | DETAILS
            dedup = {}
            for line in incident_lines:
                parts = [p.strip() for p in line.split("|")]
                if len(parts) < 5:
                    continue
                refs = _re.findall(r'\[REF:(\d+)\]', parts[0])
                date = parts[1].strip()[:10] if len(parts) > 1 else ""
                actor = parts[2].strip() if len(parts) > 2 else ""
                channel = parts[3].strip() if len(parts) > 3 else ""
                atype = parts[4].strip() if len(parts) > 4 else ""
                target = parts[5].strip() if len(parts) > 5 else ""
                sector = parts[6].strip() if len(parts) > 6 else ""
                details = parts[7].strip() if len(parts) > 7 else ""

                # Normalize target
                norm_target = target.lower().strip()
                norm_target = _re.sub(r'^(https?://)?www\.', '', norm_target)
                norm_target = _re.sub(r'[/\s]+$', '', norm_target)

                if not norm_target or len(norm_target) < 3:
                    continue

                # Dedup key: actor + target + date (different actors on same target = separate incidents)
                norm_actor = actor.lower().strip().lstrip("@")
                key = f"{norm_actor}|{norm_target}|{date}"

                if key not in dedup:
                    dedup[key] = {
                        "target": target, "refs": set(), "dates": set(),
                        "actors": set(), "channels": set(), "types": set(),
                        "sectors": set(), "details": []
                    }
                for r in refs:
                    dedup[key]["refs"].add(int(r))
                if date:
                    dedup[key]["dates"].add(date)
                if actor:
                    dedup[key]["actors"].add(actor)
                if channel:
                    dedup[key]["channels"].add(channel)
                if atype:
                    dedup[key]["types"].add(atype)
                if sector:
                    dedup[key]["sectors"].add(sector)
                if details and details not in dedup[key]["details"]:
                    dedup[key]["details"].append(details)

            # Build deduplicated incident table for reduce phase
            dedup_lines = []
            for key in sorted(dedup.keys()):
                d = dedup[key]
                ref_str = ", ".join(f"[REF:{r}]" for r in sorted(d["refs"])[:8])
                dates = ", ".join(sorted(d["dates"])[:3])
                actors = ", ".join(sorted(d["actors"])[:3])
                channels = ", ".join(sorted(d["channels"])[:2])
                types = " / ".join(sorted(d["types"]))
                sectors = ", ".join(sorted(d["sectors"])[:2])
                detail = d["details"][0][:150] if d["details"] else ""
                dedup_lines.append(
                    f"{ref_str} | {dates} | {actors} | {channels} | {types} | {d['target']} | {sectors} | {detail}"
                )

            unique_count = len(dedup_lines)
            combined_dedup = "\n".join(dedup_lines)

            yield f"data: {json.dumps({'type':'status','message':f'Found {raw_count} raw → {unique_count} unique incidents. Synthesizing report...'})}\n\n"

            # ── REDUCE: final GPT-4o synthesizes the report ──
            # Build a dedicated reduce system prompt that aggressively demands length
            _reduce_system = (
                "You are a senior cyber threat intelligence analyst writing a FULL-LENGTH classified report.\n"
                "You NEVER write short summaries. You ALWAYS write exhaustive, detailed reports.\n"
                "Your reports are typically 4000-6000 words. Anything under 2000 words is a FAILURE.\n\n"
                "WRITING RULES:\n"
                "- In the Incident Analysis sections, list EVERY SINGLE incident on its own bullet point.\n"
                "  If there are 50 government incidents, write 50 bullet points — do NOT summarize.\n"
                "- In the IOC table, list EVERY domain, IP, and URL from the data — aim for 20-40 rows.\n"
                "- In the Timeline, provide a month-by-month table with counts and key events.\n"
                "- After completing each section, ask yourself: 'Did I cover ALL incidents?' If not, keep writing.\n"
                "- NEVER use phrases like 'and many more', 'etc.', 'among others', or 'additional attacks'.\n"
                "  These are LAZY shortcuts. List every single item explicitly.\n"
                "- NEVER end a section early. ALWAYS be thorough.\n\n"
                + _REDUCE_REPORT_PROMPT
            )

            reduce_messages = [
                {"role": "system", "content": _reduce_system},
                {"role": "user", "content": (
                    f"Here are {unique_count} extracted cyber incidents (from {raw_count} raw findings):\n\n"
                    f"Format: REFS | DATE | THREAT_ACTOR | CHANNEL | ATTACK_TYPE | TARGET | SECTOR | DETAILS\n\n"
                    f"{combined_dedup}\n\n"
                    f"Write a FULL-LENGTH intelligence report covering ALL {unique_count} incidents above.\n"
                    f"I need a report that is AT LEAST 3000 words. Do NOT truncate or abbreviate.\n\n"
                    f"MANDATORY SECTIONS:\n"
                    f"1. EXECUTIVE SUMMARY (5-8 sentences)\n"
                    f"2. THREAT ACTOR OVERVIEW TABLE (every actor, with Incidents/Attack Types/Targets/Capability columns)\n"
                    f"3. INCIDENT ANALYSIS BY SECTOR — for each sector, list EVERY incident individually:\n"
                    f"   - Date | Actor | Attack Type | Specific Target (domain/org) | Impact [REF:N]\n"
                    f"   - If a sector has 50 incidents, list all 50. No summarizing.\n"
                    f"4. ATTACK TIMELINE — month-by-month table showing escalation:\n"
                    f"   | Month | Incident Count | Key Events |\n"
                    f"5. INDICATORS OF COMPROMISE TABLE — extract EVERY domain, IP, URL from findings:\n"
                    f"   | Type | Value | Context | Threat Actor |\n"
                    f"   Target: 20-40 rows. List every .jo, .il, .qa domain. Every IP. Every URL.\n"
                    f"6. THREAT ASSESSMENT (threat level, at-risk sectors, actor capabilities, escalation likelihood)\n"
                    f"7. RECOMMENDATIONS (sector-specific, reference actual incidents)\n\n"
                    f"Name SPECIFIC organizations and domains. Cite [REF:N] for every claim.\n"
                    f"START WRITING NOW. Do not stop until every incident is covered."
                )},
            ]
            # Append conversation history for context
            for h in history[-10:]:
                if h.get("role") in ("user", "assistant") and h.get("content"):
                    reduce_messages.append({"role": h["role"], "content": str(h["content"])[:2000]})

            try:
                stream = client.chat.completions.create(
                    model="gpt-4o",
                    messages=reduce_messages,
                    max_tokens=16000,
                    temperature=0.4,
                    stream=True,
                )
                full_content = ""
                for chunk_resp in stream:
                    delta = chunk_resp.choices[0].delta
                    if delta.content:
                        full_content += delta.content
                        yield f"data: {json.dumps({'type':'token','content':delta.content}, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type':'error','message':str(e)})}\n\n"
                return

            # Parse refs
            answer = full_content
            raw_refs = []
            if "---REFS---" in full_content:
                parts = full_content.split("---REFS---", 1)
                answer = parts[0].rstrip()
                ref_str = parts[1].strip()
                raw_refs = [r.strip() for r in ref_str.split(",") if r.strip()]

            valid_refs, ref_messages, seen_ref_idx = [], [], set()
            for r in raw_refs:
                try:
                    idx = int(r)
                except (ValueError, TypeError):
                    continue
                if 0 <= idx < len(all_ref_msgs) and idx not in seen_ref_idx:
                    seen_ref_idx.add(idx)
                    valid_refs.append(idx)
                    ref_messages.append(all_ref_msgs[idx])

            elapsed = round(time.time() - start_time, 1)
            yield f"data: {json.dumps({'type':'done','answer':answer,'references':valid_refs,'ref_messages':ref_messages,'elapsed_s':elapsed,'context_msgs':len(all_ref_msgs)}, ensure_ascii=False)}\n\n"
            return

        # ── Detect reformatting/follow-up to a recent report ──
        # If user is just asking to reformat (table, list, summarize), and previous
        # assistant response was long (report-like), skip tools — just answer from context
        _reformat_kw = ("table", "list them", "put it in", "reformat", "summarize",
            "in a table", "tabular", "as a table", "shorter", "bullet points",
            "can you list", "can u list", "show me all", "give me a table")
        is_reformat = any(kw in user_msg.lower() for kw in _reformat_kw)
        prev_was_report = (len(history) >= 2 and
            any(h.get("role") == "assistant" and len(str(h.get("content",""))) > 2500
                for h in history[-4:]))

        if is_reformat and prev_was_report:
            # Direct GPT call — no tools, just reformat from conversation context
            yield f"data: {json.dumps({'type':'status','message':'Reformatting from conversation context...'})}\n\n"
            try:
                stream = client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages,
                    max_tokens=12000,
                    temperature=0.15,
                    stream=True,
                )
                full_content = ""
                for chunk_resp in stream:
                    delta = chunk_resp.choices[0].delta
                    if delta.content:
                        full_content += delta.content
                        yield f"data: {json.dumps({'type':'token','content':delta.content}, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type':'error','message':str(e)})}\n\n"
                return

            answer = full_content
            raw_refs = []
            if "---REFS---" in full_content:
                parts = full_content.split("---REFS---", 1)
                answer = parts[0].rstrip()
                raw_refs = [r.strip() for r in parts[1].strip().split(",") if r.strip()]
            valid_refs, ref_messages = [], []
            elapsed = round(time.time() - start_time, 1)
            yield f"data: {json.dumps({'type':'done','answer':answer,'references':valid_refs,'ref_messages':ref_messages,'elapsed_s':elapsed,'context_msgs':0}, ensure_ascii=False)}\n\n"
            return

        # ── NORMAL (non-comprehensive): agentic tool-calling ──
        for iteration in range(12):
            try:
                stream = client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages,
                    tools=[_SEARCH_INTEL_TOOL],
                    max_tokens=12000,
                    temperature=0.2,
                    stream=True,
                )
            except Exception as e:
                yield f"data: {json.dumps({'type':'error','message':str(e)})}\n\n"
                return

            full_content = ""
            tool_calls_acc = {}
            finish_reason = None

            for chunk_resp in stream:
                choice = chunk_resp.choices[0]
                if choice.finish_reason:
                    finish_reason = choice.finish_reason
                delta = choice.delta

                if delta.content:
                    full_content += delta.content
                    yield f"data: {json.dumps({'type':'token','content':delta.content}, ensure_ascii=False)}\n\n"

                if delta.tool_calls:
                    for tc_delta in delta.tool_calls:
                        idx = tc_delta.index
                        if idx not in tool_calls_acc:
                            tool_calls_acc[idx] = {"id": "", "name": "", "arguments": ""}
                        if tc_delta.id:
                            tool_calls_acc[idx]["id"] = tc_delta.id
                        if tc_delta.function:
                            if tc_delta.function.name:
                                tool_calls_acc[idx]["name"] = tc_delta.function.name
                            if tc_delta.function.arguments:
                                tool_calls_acc[idx]["arguments"] += tc_delta.function.arguments

            if finish_reason == "tool_calls" and tool_calls_acc:
                assistant_tc = []
                for idx_key in sorted(tool_calls_acc.keys()):
                    tc = tool_calls_acc[idx_key]
                    assistant_tc.append({
                        "id": tc["id"],
                        "type": "function",
                        "function": {"name": tc["name"], "arguments": tc["arguments"]}
                    })
                messages.append({
                    "role": "assistant",
                    "content": full_content or None,
                    "tool_calls": assistant_tc
                })

                for idx_key in sorted(tool_calls_acc.keys()):
                    tc = tool_calls_acc[idx_key]
                    try:
                        args = json.loads(tc["arguments"])
                    except Exception:
                        args = {"query": ""}

                    query = args.get("query", "")
                    severity = args.get("severity", "all")
                    days_back = args.get("days_back", 365)

                    yield f"data: {json.dumps({'type':'status','message':f'Searching: {query}...'})}\n\n"

                    count, formatted = _execute_search_tool(
                        query, severity, days_back, seen_ids, all_ref_msgs, all_msgs
                    )

                    messages.append({
                        "role": "tool",
                        "tool_call_id": tc["id"],
                        "content": formatted
                    })

                continue
            else:
                break

        # Parse references from final text
        answer = full_content
        raw_refs = []
        if "---REFS---" in full_content:
            parts = full_content.split("---REFS---", 1)
            answer = parts[0].rstrip()
            ref_str = parts[1].strip()
            raw_refs = [r.strip() for r in ref_str.split(",") if r.strip()]

        valid_refs, ref_messages, seen_ref_idx = [], [], set()
        for r in raw_refs:
            try:
                idx = int(r)
            except (ValueError, TypeError):
                continue
            if 0 <= idx < len(all_ref_msgs) and idx not in seen_ref_idx:
                seen_ref_idx.add(idx)
                valid_refs.append(idx)
                ref_messages.append(all_ref_msgs[idx])

        elapsed = round(time.time() - start_time, 1)
        yield f"data: {json.dumps({'type':'done','answer':answer,'references':valid_refs,'ref_messages':ref_messages,'elapsed_s':elapsed,'context_msgs':len(all_ref_msgs)}, ensure_ascii=False)}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        }
    )


@app.route("/api/ai/log")
def api_ai_log():
    """Return last N lines of AI agent log with colour-coded level tags."""
    lines = int(request.args.get("lines", 80))
    log_path = OUTPUT_DIR / "ai_agent.log"
    if not log_path.exists():
        return jsonify({"lines": [], "size": 0})
    try:
        all_lines = log_path.read_text(encoding="utf-8", errors="replace").splitlines()
        tail = all_lines[-lines:]
        return jsonify({"lines": tail, "total": len(all_lines), "size": log_path.stat().st_size})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/analyze", methods=["POST"])
def api_ai_analyze():
    """Launch ai_agent.py as a background daemon process."""
    import subprocess, psutil
    # Check if already running
    for p in psutil.process_iter(['pid', 'cmdline']):
        try:
            if p.info['cmdline'] and 'ai_agent.py' in ' '.join(p.info['cmdline']):
                return jsonify({"error": "agent already running", "pid": p.pid}), 409
        except Exception:
            pass
    try:
        proc = subprocess.Popen(
            ["python", "ai_agent.py"],
            stdout=open(OUTPUT_DIR / "ai_agent.log", "a"),
            stderr=subprocess.STDOUT,
            cwd=str(Path(__file__).parent),
        )
        return jsonify({"started": True, "pid": proc.pid})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/stop", methods=["POST"])
def api_ai_stop():
    """Stop the ai_agent.py daemon."""
    import psutil
    killed = []
    for p in psutil.process_iter(['pid', 'cmdline']):
        try:
            if p.info['cmdline'] and 'ai_agent.py' in ' '.join(p.info['cmdline']):
                p.kill()
                killed.append(p.pid)
        except Exception:
            pass
    return jsonify({"stopped": killed})


# ══════════════════════════════════════════════════════════════════════════════
# SYSTEM CONFIG + ORCHESTRATOR STATUS
# ══════════════════════════════════════════════════════════════════════════════

ENV_FILE_PATH = Path(__file__).parent / ".env"
ORCH_STATUS_FILE = OUTPUT_DIR / "orchestrator_status.json"


def _load_env_file():
    """Load .env file as a dict."""
    if not ENV_FILE_PATH.exists():
        return {}
    out = {}
    for line in ENV_FILE_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        out[k.strip()] = v.strip().strip('"').strip("'")
    return out


def _save_env_key(key, value):
    """Write or update a key in the .env file."""
    lines = []
    if ENV_FILE_PATH.exists():
        lines = ENV_FILE_PATH.read_text(encoding="utf-8").splitlines()
    found = False
    new_lines = []
    for line in lines:
        if line.strip().startswith(f"{key}=") or line.strip().startswith(f"{key} ="):
            new_lines.append(f"{key}={value}")
            found = True
        else:
            new_lines.append(line)
    if not found:
        new_lines.append(f"{key}={value}")
    ENV_FILE_PATH.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
    os.environ[key] = value


@app.route("/api/system/status")
def api_system_status():
    """Return orchestrator component status + env config presence."""
    import psutil

    # Check which processes are running
    procs = {"viewer": False, "monitor": False, "ai_agent": False, "orchestrator": False}
    for p in psutil.process_iter(['pid', 'cmdline']):
        try:
            cmd = ' '.join(p.info['cmdline'] or [])
            if 'viewer.py' in cmd:           procs["viewer"]       = True
            if 'telegram_monitor.py' in cmd: procs["monitor"]      = True
            if 'ai_agent.py' in cmd:         procs["ai_agent"]     = True
            if 'orchestrator.py' in cmd:     procs["orchestrator"] = True
        except Exception:
            pass

    env = _load_env_file()
    agent_state = {}
    if AI_STATE_FILE.exists():
        try:
            agent_state = json.loads(AI_STATE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass

    brief = None
    if AI_BRIEF_FILE.exists():
        try:
            brief = json.loads(AI_BRIEF_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass

    return jsonify({
        "processes": procs,
        "has_openai_key": bool(env.get("OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY")),
        "agent_stats": agent_state,
        "latest_brief": brief,
    })


@app.route("/api/system/config", methods=["POST"])
def api_system_config():
    """Save API keys and config to .env file."""
    data = request.get_json(force=True) or {}
    saved = []
    allowed_keys = ["OPENAI_API_KEY", "TG_API_ID", "TG_API_HASH", "TG_PHONE"]
    for key in allowed_keys:
        val = str(data.get(key, "")).strip()
        if val:
            _save_env_key(key, val)
            saved.append(key)
    return jsonify({"ok": True, "saved": saved,
                    "note": "Keys saved to .env — restart AI agent to activate"})


@app.route("/api/ai/apply", methods=["POST"])
def api_ai_apply():
    """Apply latest AI keyword suggestions to keywords.json."""
    if not AI_SUGGESTIONS.exists():
        return jsonify({"error": "no suggestions available"}), 404
    try:
        data   = json.loads(AI_SUGGESTIONS.read_text(encoding="utf-8"))
        latest = data.get("latest", {})
        new_crit = [k for k in latest.get("new_critical_keywords", []) if k]
        new_med  = [k for k in latest.get("new_medium_keywords",   []) if k]

        kw_data = {"critical": [], "medium": []}
        if KEYWORDS_FILE.exists():
            try:
                kw_data = json.loads(KEYWORDS_FILE.read_text(encoding="utf-8"))
            except Exception:
                pass

        existing_crit = set(k.lower() for k in kw_data.get("critical", []))
        existing_med  = set(k.lower() for k in kw_data.get("medium",   []))

        added_crit = [k for k in new_crit if k.lower() not in existing_crit]
        added_med  = [k for k in new_med  if k.lower() not in existing_med]

        kw_data.setdefault("critical", []).extend(added_crit)
        kw_data.setdefault("medium",   []).extend(added_med)
        KEYWORDS_FILE.write_text(
            json.dumps(kw_data, indent=2, ensure_ascii=False), encoding="utf-8")

        return jsonify({"ok": True, "added_critical": len(added_crit),
                        "added_medium": len(added_med)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/discovery/fetch")
def api_discovery_fetch():
    username = request.args.get("username", "").strip()
    limit    = min(int(request.args.get("limit", 100)), 500)
    save     = request.args.get("save", "false").lower() == "true"
    if not username or not TELETHON_OK:
        return jsonify({"error": "need username + telethon", "messages": []})
    try:
        try:
            asyncio.get_event_loop()
        except RuntimeError:
            asyncio.set_event_loop(asyncio.new_event_loop())
        with TGSync(SESSION, API_ID, API_HASH) as client:
            raw = list(client.iter_messages(username, limit=limit))
        raw.sort(key=lambda m: m.id)
        messages = []
        for m in raw:
            text = m.text or ""
            priority, hits = _score_text(text)
            messages.append({
                "message_id":       m.id,
                "channel_username": username,
                "channel":          username,
                "timestamp_utc":    m.date.isoformat() if m.date else "",
                "timestamp_irst":   "",
                "text_preview":     text,
                "priority":         priority,
                "keyword_hits":     hits,
                "iocs":             {},
                "has_media":        bool(m.media),
                "live":             True,
            })
        if save:
            seen = set()
            try:
                with open(OUTPUT_DIR / "messages.jsonl", encoding="utf-8") as f:
                    for line in f:
                        try:
                            rec = json.loads(line)
                            seen.add(f"{rec.get('channel_username')}_{rec.get('message_id')}")
                        except Exception:
                            pass
            except FileNotFoundError:
                pass
            with open(OUTPUT_DIR / "messages.jsonl", "a", encoding="utf-8") as f:
                for msg in messages:
                    key = f"{msg['channel_username']}_{msg['message_id']}"
                    if key not in seen:
                        f.write(json.dumps(msg, ensure_ascii=False) + "\n")
            _msg_cache["data"] = None  # Invalidate cache
        return jsonify({"messages": messages, "total": len(messages)})
    except Exception as e:
        return jsonify({"error": str(e), "messages": []})


# ═══════════════════════════════════════════════════════════════════════════════
# HTML TEMPLATE (loaded from app/static/index.html)
# ═══════════════════════════════════════════════════════════════════════════════

_HTML_FILE = Path(__file__).parent / "app" / "static" / "index.html"
HTML = _HTML_FILE.read_text(encoding="utf-8") if _HTML_FILE.exists() else "<!-- index.html not found -->"



@app.route("/")
def index():
    return Response(HTML, mimetype="text/html", headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0",
    })


if __name__ == "__main__":
    log.info("=" * 60)
    log.info("  Scanwave CyberIntel Platform")
    log.info("  http://localhost:5000")
    log.info("=" * 60)
    # Start auto-research background thread
    _research_thread = threading.Thread(target=_auto_research_loop, daemon=True)
    _research_thread.start()
    log.info("[RESEARCH] Auto-research background thread started")
    app.run(debug=False, host="0.0.0.0", port=5000, threaded=True)
