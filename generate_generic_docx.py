"""
Clone the original DOCX and replace all specific names with generic equivalents.
Preserves ALL formatting, logos, headers, footers, tables, etc.
"""

from docx import Document
import copy
import re
import os

INPUT = r"c:\Users\BisharahQraitem\Downloads\cybernews\ScanWave_BAE_IOC_Combined_v3.docx"
OUTPUT = r"c:\Users\BisharahQraitem\Downloads\cybernews\ScanWave_Generic_IOC_Intelligence_Report.docx"

# ── Replacement map (order matters — longer/more specific first) ──
REPLACEMENTS = [
    # Arabic name
    ("بنك الاتحاد", ""),
    ("بنك االتحاد", ""),

    # Full entity names with context
    ("Bank al Etihad (Bank al Etihad)", "the targeted financial institution"),
    ("Bank al Etihad (بنك الاتحاد) INTELLIGENCE REPORT", "Threat Intelligence Report — Financial & Energy Sector"),
    ("Bank al Etihad (بنك الاتحاد)", "a major financial institution in the Middle East"),
    ("Bank al Etihad ()", "a major financial institution"),
    # After Arabic removal, the subtitle becomes "X  INTELLIGENCE REPORT"
    ("a major financial institution in the Middle East INTELLIGENCE REPORT", "Threat Intelligence Report — Financial & Energy Sector"),
    ("a major financial institution INTELLIGENCE REPORT", "Threat Intelligence Report — Financial & Energy Sector"),
    ("Bank al Etihad", "the targeted financial institution"),
    ("bank al Etihad", "the targeted financial institution"),
    ("BankAlEtihad", "TargetOrganization"),
    ("bankaletihad", "target-organization"),

    # BAE abbreviation (careful — must not match random text)
    ("BAE/BELECTRIC", "the targeted organization/BELECTRIC"),
    ("BELECTRIC/BAE", "BELECTRIC/the targeted organization"),
    ("BAE's S3 bucket (prod-bae-blog-frontend)", "the organization's cloud storage bucket"),
    ("prod-bae-blog-frontend", "org-cloud-storage"),
    ("BAE web assets", "the organization's web assets"),
    ("BAE employee", "organizational employee"),
    ("BAE assets", "organizational assets"),
    ("BAE infrastructure", "organizational infrastructure"),
    ("BAE ", "the targeted organization's "),
    (" BAE ", " the targeted organization "),
    ("/BAE", "/targeted-org"),
    ("UNK-IR-ETIHAD-0307", "UNK-IR-SECTOR-0307"),

    # BELECTRIC — KEEP as-is (user requested)
    # No BELECTRIC replacements

    # Specific victim names in tables
    ("Israel Opportunity Energy", "Energy Company — Middle East"),
    ("Aramco (Saudi)", "Oil & Gas Corporation — Gulf Region"),
    ("Sharjah National Oil Corporation", "National Oil Corporation — Gulf Region"),
    ("INSS (Israel National Security Studies)", "Government Research Institute"),
    ("Atlas Insurances Ltd", "Regional Insurance Provider"),
    ("Clalit Healthcare", "Healthcare Organization"),

    # Jordan-specific
    ("Jordan CERT (JO-CERT)", "National CERT"),
    ("JO-CERT", "National CERT"),
    ("Jordan Telecom / Orange Jordan", "the national telecom provider"),
    ("Orange Jordan", "the national telecom provider"),
    ("Jordanian financial institutions", "Middle Eastern financial institutions"),
    ("Jordanian financial infrastructure", "Middle Eastern financial infrastructure"),
    ("Jordanian financial and energy sectors", "regional financial and energy sectors"),
    ("Jordanian critical infrastructure", "regional critical infrastructure"),
    ("Jordanian targets", "regional targets"),
    ("Jordan, Israel, and the Gulf region", "the Middle East region"),
    ("South Amman Solar Plant", "The largest solar installation"),
    ("Aqaba Special Economic Zone", "a national economic zone"),
    ("the Aqaba solar monitoring infrastructure", "the solar energy monitoring infrastructure"),
    ("in Jordan", "in the target country"),
    ("of Jordan", "in the Middle East"),
    ("within Jordan", "within the target country"),
    ("deployed in Jordan", "deployed in the target country"),
    ("across Jordan", "across the region"),
    ("Jordan's leading commercial banks", "the region's leading financial institutions"),
    ("Jordan's", "the target country's"),
    ("Jordan ", "the target country "),

    # Specific hostnames / infrastructure
    ("Primary-NGFW.bankaletihad.com (Serial: 013201006040)", "Primary-NGFW.[target-org].com"),
    ("Primary-NGFW.bankaletihad.com", "Primary-NGFW.[target-org].com"),
    ("Batch-ADP-PROD.BankAlEtihad.Com (10.12.11.233)", "Batch-Server.[target-org].com"),
    ("Batch-ADP-PROD.BankAlEtihad.Com", "Batch-Server.[target-org].com"),
    ("Batch-ADP-PROD\\Admin", "BATCH-SERVER\\Admin"),
    ("BATCH-ADP-PROD\\Admin", "BATCH-SERVER\\Admin"),
    ("examproctor.bankaletihad.com", "[subdomain].[target-org].com"),
    ("m.alkhateeb@bankaletihad.com", "[employee]@[target-org].com"),
    ("bankaletihad.com", "[target-org].com"),
    ("829210555596", "[REDACTED]"),

    # Internal IPs (specific to the org)
    ("10.0.6.69, 10.0.6.71", "internal branch devices"),
    ("10.0.6.69", "10.x.x.69"),
    ("10.0.6.71", "10.x.x.71"),
    ("172.17.30.1", "10.x.x.1"),
    ("10.12.11.233", "10.x.x.233"),
    ("94.249.51.64", "[REDACTED]"),
    ("13.41.61.138 and 23.21.102.205", "externally-reported IPs"),
    ("13.41.61.138", "[REDACTED]"),
    ("23.21.102.205", "[REDACTED]"),

    # Screenshots references
    ("Screenshots provided as proof-of-compromise include live server monitoring dashboards", "Evidence provided as proof-of-compromise includes live server monitoring dashboards"),
    ("Evidence from the posted screenshots indicates compromise of a", "Evidence from the posted materials indicates compromise of a"),
    ("from the posted screenshots", "from the posted materials"),
    ("screenshots", "evidence materials"),
    ("screenshot", "evidence"),
    ("Screenshot capture", "Capture"),

    # Smart/curly quote variants (DOCX uses \u2019 for apostrophe)
    ("the bank\u2019s core server infrastructure", "the organization\u2019s core server infrastructure"),
    ("the bank\u2019s core servers", "the organization\u2019s core servers"),
    ("the bank\u2019s technical infrastructure", "the organization\u2019s technical infrastructure"),
    ("the bank\u2019s security team\u2019s", "the organization\u2019s security team\u2019s"),
    ("the bank\u2019s security team", "the organization\u2019s security team"),
    ("the bank\u2019s", "the organization\u2019s"),
    ("this bank\u2019s", "this organization\u2019s"),
    ("the Middle East\u2019s leading", "the region\u2019s leading"),
    ("commercial banks", "commercial financial institutions"),
    ("commercial banks in the Middle East", "financial institutions in the region"),

    # Banking-specific references in findings
    ("banking core systems", "core organizational systems"),
    ("banking core servers", "core servers"),
    ("banking infrastructure", "organizational infrastructure"),
    ("banking servers", "core servers"),
    ("Banking Core Server Dashboard", "Core Server Dashboard"),
    ("the bank's core server infrastructure", "the organization's core server infrastructure"),
    ("the bank's core servers", "the organization's core servers"),
    ("the bank's technical infrastructure", "the organization's technical infrastructure"),
    ("the bank's security team", "the organization's security team"),
    ("the bank's security team's", "the organization's security team's"),
    ("the bank's", "the organization's"),
    ("the bank ", "the organization "),
    ("Bank ", "Organization "),
    ("core banking servers", "core servers"),
    ("across banking infrastructure", "across organizational infrastructure"),
    ("financial infrastructure", "critical infrastructure"),

    # Section-specific rewrites
    ("Probes against the organization's cloud storage bucket (org-cloud-storage), all returned AccessDenied:", "Probes against the organization's cloud storage bucket, all returned AccessDenied:"),
    ("Device: Primary-NGFW.[target-org].com | 4 events detected:", "4 high-severity events detected on the primary next-generation firewall:"),
    ("Host: Batch-Server.[target-org].com:", "Database audit events from the batch processing server:"),
    ("AWS Account: [REDACTED] (eu-west-2 / London)", "Cloud platform console access review:"),
    ("[REDACTED] resolves to the national telecom provider. Single legitimate console login from a organizational employee. No failed login attempts or suspicious IAM activity detected.", "A single legitimate console login was detected from an authorized employee originating from the national telecom provider's IP range. No failed login attempts or suspicious IAM activity detected."),
    ("Both source IPs (internal branch devices) are internal branch network devices targeting a VoIP/SIP gateway at 10.x.x.1.", "Source IPs are internal branch network devices targeting a VoIP/SIP gateway."),

    # Specific subsection titles
    ("6.3 S3 Bucket Reconnaissance Probes", "6.3 Cloud Storage Reconnaissance Probes"),
    ("6.4 Palo Alto NGFW — High Severity Threat Events", "6.4 NGFW — High Severity Threat Events"),
    ("6.5 Oracle Database Audit Events", "6.5 Database Audit Events"),
    ("6.6 AWS CloudTrail — Console Access", "6.6 Cloud Console — Access Review"),
    # These cleanup rules no longer needed since BELECTRIC is kept as-is

    # Aqaba standalone (in table cells etc.)
    ("Aqaba Special Economic Zone", "a national economic zone"),
    ("Aqaba SEZ", "the national economic zone"),
    ("Aqaba solar monitoring infrastructure", "solar energy monitoring infrastructure"),
    ("Aqaba", "the national economic zone"),

    # examproctor standalone
    ("examproctor.bankaletihad.com", "[subdomain].[target-org].com"),
    ("examproctor.target-organization.com", "[subdomain].[target-org].com"),
    ("examproctor", "[subdomain]"),

    # alkhateeb standalone (table cells)
    ("m.alkhateeb@bankaletihad.com", "[employee]@[target-org].com"),
    ("m.alkhateeb@target-organization.com", "[employee]@[target-org].com"),
    ("m.alkhateeb", "[employee]"),
    ("alkhateeb", "[employee]"),

    # Israel references
    ("Israel", "the region"),
    ("Israeli targets", "regional targets"),

    # Cleanup any double spaces or awkward phrasing from replacements
    ("one of the region's leading financial institutions", "a major financial institution in the region"),
    ("the targeted financial institution (), ", "the targeted financial institution, "),
    ("one in the Middle East\u2019s leading commercial financial institutions", "a major financial institution in the Middle East region"),
    ("one in the Middle East's leading commercial financial institutions", "a major financial institution in the Middle East region"),
    ("the banking and energy sectors", "the financial services and energy sectors"),
    ("a an industrial", "an industrial"),
    ("the a national", "a national"),
    ("one in the region\u2019s leading commercial financial institutions", "one of the region\u2019s leading financial institutions"),
    ("a major financial institution, one of the region\u2019s leading financial institutions", "a major financial institution in the Middle East region"),
    ("a major financial institution in the Middle East, a major financial institution in the region", "a major financial institution in the Middle East region"),
    # BELECTRIC kept, so these cleanups no longer needed
    (" ()", ""),
    ("  ", " "),
]


def replace_in_text(text):
    """Apply all replacements to a text string."""
    if not text:
        return text
    result = text
    for old, new in REPLACEMENTS:
        result = result.replace(old, new)
    return result


def replace_in_runs(paragraph):
    """Replace text in runs while preserving formatting."""
    full_text = paragraph.text
    if not full_text.strip():
        return

    new_text = replace_in_text(full_text)
    if new_text == full_text:
        return

    # Strategy: if single run, simple replace
    runs = paragraph.runs
    if len(runs) == 1:
        runs[0].text = replace_in_text(runs[0].text)
        return

    if len(runs) == 0:
        return

    # For multi-run paragraphs, try run-by-run replacement first
    changed = False
    for run in runs:
        new_run_text = replace_in_text(run.text)
        if new_run_text != run.text:
            run.text = new_run_text
            changed = True

    # Check if full text is now correct
    if changed:
        current = "".join(r.text for r in runs)
        target = replace_in_text(full_text)
        if current == target:
            return

    # If run-by-run didn't work perfectly (replacements span runs),
    # put all text in first run and clear others
    new_full = replace_in_text(full_text)
    if new_full != "".join(r.text for r in runs):
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""


def replace_in_table_cell(cell):
    """Replace text in a table cell."""
    for paragraph in cell.paragraphs:
        replace_in_runs(paragraph)


def process_document():
    doc = Document(INPUT)

    # Process all paragraphs
    for p in doc.paragraphs:
        replace_in_runs(p)

    # Process all tables (including section heading tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_table_cell(cell)

    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for p in header.paragraphs:
                    replace_in_runs(p)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_table_cell(cell)

        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    replace_in_runs(p)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_table_cell(cell)

    # Add disclaimer at the bottom of the last page
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add spacing before disclaimer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(24)

    # Add horizontal rule
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(6)
    rule_run = rule.add_run("_" * 95)
    rule_run.font.size = Pt(7)
    rule_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add DISCLAIMER heading
    disc_heading = doc.add_paragraph()
    disc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    disc_heading.paragraph_format.space_before = Pt(8)
    disc_heading.paragraph_format.space_after = Pt(4)
    dh_run = disc_heading.add_run("DISCLAIMER")
    dh_run.font.size = Pt(9)
    dh_run.font.bold = True
    dh_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Add disclaimer text
    disclaimer_text = (
        "All information contained in this report was collected exclusively through publicly available "
        "Open Source Intelligence (OSINT), social media monitoring, and third-party threat intelligence feeds. "
        "No unauthorized access to any systems was performed during the preparation of this report. "
        "While this report contains security recommendations, ScanWave Cybersecurity does not guarantee the "
        "accuracy, completeness, or reliability of any information presented herein. This report is provided "
        'on an "as-is" basis for informational and defensive security purposes only and does not constitute '
        "legal or professional advice. ScanWave Cybersecurity, its employees, and its affiliates shall not be "
        "held liable for any damages, losses, or consequences \u2014 whether direct, indirect, incidental, or "
        "consequential \u2014 arising from the use of, reliance on, or actions taken based on the information "
        "contained in this report. Recipients assume full responsibility for any decisions made based on this material."
    )
    disc_para = doc.add_paragraph()
    disc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disc_para.paragraph_format.space_after = Pt(0)
    d_run = disc_para.add_run(disclaimer_text)
    d_run.font.size = Pt(7.5)
    d_run.font.italic = True
    d_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

    # Verify no banned terms remain
    verify_doc = Document(OUTPUT)
    all_text = ""
    for p in verify_doc.paragraphs:
        all_text += p.text + "\n"
    for table in verify_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
    for section in verify_doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                all_text += p.text + "\n"

    banned = [
        "Bank al Etihad", "bankaletihad", "BankAlEtihad",
        # BELECTRIC kept intentionally
        "Aramco", "Clalit", "Atlas Insurances", "INSS",
        "Sharjah", "Jordan", "Aqaba",
        "alkhateeb", "screenshot",
        "بنك", "الاتحاد",
        "829210555596", "examproctor",
        "prod-bae-blog",
    ]

    found_issues = False
    for term in banned:
        # Case-sensitive search
        count = all_text.count(term)
        if count > 0:
            # Find context
            idx = all_text.find(term)
            context = all_text[max(0, idx-30):idx+len(term)+30].replace('\n', ' ')
            print(f"  WARNING: \"{term}\" found {count}x — context: ...{context}...")
            found_issues = True

    if not found_issues:
        print("  ALL CLEAN — no banned terms found!")


if __name__ == "__main__":
    process_document()
