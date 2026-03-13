#!/usr/bin/env python3
"""
Generate ScanWave CyberIntel Platform — API Reference Documentation
Outputs: DOCX + PDF
"""
import os, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx..."); os.system(f"{sys.executable} -m pip install python-docx");
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

# ── Colors ──────────────────────────────────────────────────────────────
DARK_BG     = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_CYAN = RGBColor(0x00, 0xD4, 0xFF)
ACCENT_RED  = RGBColor(0xFF, 0x00, 0x44)
ACCENT_GRN  = RGBColor(0x00, 0xFF, 0x88)
ACCENT_YLW  = RGBColor(0xFF, 0xD7, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LT     = RGBColor(0xBB, 0xBB, 0xBB)
GRAY_DK     = RGBColor(0x66, 0x66, 0x66)
BLACK       = RGBColor(0x00, 0x00, 0x00)
TBL_HEADER  = "1A1A2E"
TBL_ROW_A   = "F7F9FC"
TBL_ROW_B   = "EBEEF5"
BORDER_CLR  = "CCCCCC"

HTTP_COLORS = {
    "GET":    RGBColor(0x00, 0xAA, 0x55),  # green
    "POST":   RGBColor(0x00, 0x77, 0xDD),  # blue
    "DELETE": RGBColor(0xDD, 0x33, 0x33),  # red
    "PUT":    RGBColor(0xEE, 0x99, 0x00),  # orange
    "PATCH":  RGBColor(0xAA, 0x55, 0xDD),  # purple
}

# ── Helper Functions ────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{val}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_para(doc, text, style_name, size=11, bold=False, color=BLACK, alignment=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(1)
    # light gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F2F5" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    return p

def make_endpoint_table(doc, rows_data):
    """Create a styled parameter/response table."""
    table = doc.add_table(rows=1, cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    # Header row
    hdr = table.rows[0]
    for i, txt in enumerate(rows_data[0]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        set_cell_shading(cell, TBL_HEADER)
    # Data rows
    for ri, row in enumerate(rows_data[1:]):
        r = table.add_row()
        bg = TBL_ROW_A if ri % 2 == 0 else TBL_ROW_B
        for ci, txt in enumerate(row):
            cell = r.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(txt))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            set_cell_shading(cell, bg)
    # Set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            for border_edge in ["top", "bottom", "start", "end"]:
                set_cell_border(cell, **{border_edge: BORDER_CLR})
    doc.add_paragraph()  # spacer
    return table

def add_method_badge(p, method):
    run = p.add_run(f" {method} ")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = HTTP_COLORS.get(method, BLACK)
    run.font.name = "Consolas"

def add_endpoint_section(doc, method, path, description, params=None, body=None, response=None, response_fields=None, notes=None, example_response=None):
    """Add a complete endpoint documentation block."""
    # Method + Path header
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)

    add_method_badge(p, method)
    run = p.add_run(f"  {path}")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Description
    add_styled_para(doc, description, "Normal", size=10, color=RGBColor(0x44, 0x44, 0x44), space_after=6)

    # Parameters table
    if params:
        add_styled_para(doc, "Query / URL Parameters", "Normal", size=10, bold=True, color=RGBColor(0x33, 0x33, 0x55), space_before=4, space_after=2)
        header = ["Parameter", "Type", "Required", "Default", "Description"]
        make_endpoint_table(doc, [header] + params)

    # Request body
    if body:
        add_styled_para(doc, "Request Body (JSON)", "Normal", size=10, bold=True, color=RGBColor(0x33, 0x33, 0x55), space_before=4, space_after=2)
        header = ["Field", "Type", "Required", "Description"]
        make_endpoint_table(doc, [header] + body)

    # Response fields
    if response_fields:
        add_styled_para(doc, "Response Fields", "Normal", size=10, bold=True, color=RGBColor(0x33, 0x33, 0x55), space_before=4, space_after=2)
        header = ["Field", "Type", "Description"]
        make_endpoint_table(doc, [header] + response_fields)

    # Response format note
    if response:
        add_styled_para(doc, f"Response: {response}", "Normal", size=9, color=GRAY_DK, space_after=4)

    # Example response
    if example_response:
        add_styled_para(doc, "Example Response", "Normal", size=10, bold=True, color=RGBColor(0x33, 0x33, 0x55), space_before=4, space_after=2)
        add_code_block(doc, example_response)

    # Notes
    if notes:
        for note in notes:
            add_styled_para(doc, f"  {note}", "Normal", size=9, color=RGBColor(0x88, 0x66, 0x00), space_after=2)

    # Separator line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("_" * 95)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


def add_section_header(doc, number, title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(8)
    run = p.add_run(f"{number}  ")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = ACCENT_CYAN if ACCENT_CYAN else RGBColor(0x00, 0x77, 0xCC)
    run.font.name = "Calibri"
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "Calibri"
    # underline
    p2 = doc.add_paragraph()
    run2 = p2.add_run("━" * 80)
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)
    p2.paragraph_format.space_after = Pt(8)

def add_subsection_header(doc, number, title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)
    run = p.add_run(f"{number}  {title}")
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x4E)
    run.font.name = "Calibri"


# ══════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT GENERATION
# ══════════════════════════════════════════════════════════════════════════
def generate():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── COVER PAGE ──────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SCANWAVE CYBERSECURITY")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)
    run.font.name = "Calibri"
    run.font.letter_spacing = Pt(3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("CyberIntel Platform")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Backend API Reference Documentation")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    run.font.name = "Calibri"

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run("━" * 50)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(f"Version 2.4  |  {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_DK
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("CONFIDENTIAL — FOR AUTHORIZED PERSONNEL ONLY")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = ACCENT_RED
    run.font.name = "Calibri"

    # Meta info
    for _ in range(4):
        doc.add_paragraph()

    meta_items = [
        ("Document Classification:", "CONFIDENTIAL"),
        ("Prepared By:", "ScanWave Cybersecurity — Engineering Division"),
        ("Platform:", "ScanWave CyberIntel Telegram Monitoring Platform"),
        ("Base URL:", "https://<deployment-host>:<port>"),
        ("Content-Type:", "application/json (unless otherwise noted)"),
        ("Authentication:", "Session-based (internal deployment)"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + "  ")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Calibri"
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = "Calibri"

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "Calibri"

    toc_items = [
        ("1", "Introduction & Overview"),
        ("2", "Authentication & Common Conventions"),
        ("3", "Message & Channel Intelligence APIs"),
        ("  3.1", "Channel Listing & Statistics"),
        ("  3.2", "Message Retrieval & Filtering"),
        ("  3.3", "Real-Time Polling"),
        ("  3.4", "Message Context"),
        ("  3.5", "Export Endpoints"),
        ("4", "Dashboard & Analytics APIs"),
        ("  4.1", "Dashboard Aggregation"),
        ("  4.2", "Trend Analysis"),
        ("  4.3", "Intelligence Briefing"),
        ("  4.4", "Status Summary"),
        ("  4.5", "Threat Matrix"),
        ("5", "APT Tracker APIs"),
        ("  5.1", "APT Profiles"),
        ("  5.2", "APT Network Map"),
        ("  5.3", "APT Detail"),
        ("  5.4", "IOC Lookup & Research"),
        ("  5.5", "AI-Powered IOC Extraction"),
        ("6", "Blocklist & IOC Management APIs"),
        ("  6.1", "Central Blocklist"),
        ("  6.2", "Blocklist Export"),
        ("  6.3", "Report Generation"),
        ("7", "Discovery Engine APIs"),
        ("  7.1", "Discovery Listing"),
        ("  7.2", "Scan Trigger & Status"),
        ("  7.3", "On-Demand Channel Fetch"),
        ("8", "Dark Web & Threat Intelligence APIs"),
        ("  8.1", "Dark Intel Feed"),
        ("  8.2", "Dark Collector Statistics"),
        ("  8.3", "Daily Digest"),
        ("  8.4", "Domain Squatting"),
        ("  8.5", "Behavioral Threat Level"),
        ("9", "AI Agent & Enrichment APIs"),
        ("  9.1", "Agent Status & Control"),
        ("  9.2", "Keyword Suggestions"),
        ("  9.3", "Threat Briefs"),
        ("  9.4", "Enriched Alerts"),
        ("  9.5", "Translation"),
        ("10", "Escalation & Threat Hunting APIs"),
        ("11", "Chat & Intelligence Query APIs"),
        ("  11.1", "Standard Chat"),
        ("  11.2", "Streaming Chat (SSE)"),
        ("12", "Administration APIs"),
        ("  12.1", "System Status & Monitoring"),
        ("  12.2", "Keyword Management"),
        ("  12.3", "Channel Management"),
        ("  12.4", "Backfill & Maintenance"),
        ("  12.5", "Discovery Management"),
        ("  12.6", "System Configuration"),
        ("13", "Media & File Serving APIs"),
        ("14", "Error Handling & Status Codes"),
        ("15", "Rate Limiting & Performance Notes"),
        ("A", "Appendix: Complete Endpoint Index"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        indent = "  " in num
        run = p.add_run(f"{num.strip()}")
        run.font.size = Pt(10 if not indent else 9)
        run.font.bold = not indent
        run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC) if not indent else RGBColor(0x44, 0x44, 0x66)
        run.font.name = "Calibri"
        run = p.add_run(f"    {title}")
        run.font.size = Pt(10 if not indent else 9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22) if not indent else RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Calibri"
        pf = p.paragraph_format
        pf.space_before = Pt(2 if indent else 6)
        pf.space_after = Pt(1)
        if indent:
            pf.left_indent = Cm(1.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 1 — INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "1", "Introduction & Overview")

    add_styled_para(doc,
        "The ScanWave CyberIntel Platform is a comprehensive Telegram-based cyber threat intelligence "
        "monitoring system designed for tracking Iranian-aligned hacktivist groups and Advanced Persistent "
        "Threat (APT) actors targeting critical infrastructure in the Middle East region. The platform "
        "provides real-time monitoring of 100+ Telegram channels, automated IOC extraction, AI-powered "
        "enrichment, dark web intelligence collection, and full threat lifecycle management.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=8)

    add_styled_para(doc,
        "This document provides a complete technical reference for all backend REST API endpoints exposed "
        "by the platform's Flask-based web server (viewer.py). It is intended for security engineers, "
        "SOC analysts, and integration developers who need to interact with the platform programmatically "
        "or understand its data model for custom tooling and reporting.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_styled_para(doc, "Platform Capabilities", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)

    capabilities = [
        "Real-time Telegram channel monitoring with priority-based message classification (CRITICAL / MEDIUM / LOW)",
        "Automated IOC extraction: IPv4 addresses, domains, SHA256/MD5 hashes, URLs, email addresses",
        "APT group profiling with sector targeting analysis and collaboration network mapping",
        "AI-powered message enrichment using GPT-4o-mini for threat context and attack categorization",
        "Dark web intelligence collection across 8 source categories (paste sites, breach databases, certificate transparency, domain squatting, GitHub dorks, ransomware leak sites, Tor hidden services, threat intel feeds)",
        "Behavioral threat level computation with 7 signal dimensions",
        "Automated channel discovery engine with relevance scoring and analyst review workflow",
        "Interactive intelligence chatbot with map-reduce query processing and SSE streaming",
        "Multi-format export: CSV, DOCX, PDF report generation with legal disclaimers",
        "Escalation detection and proactive threat hunting lead generation",
    ]
    for cap in capabilities:
        p = doc.add_paragraph()
        run = p.add_run("  \u2022  ")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)
        run = p.add_run(cap)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(3)

    add_styled_para(doc, "Architecture Overview", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=14, space_after=6)

    arch_items = [
        ("Web Server (viewer.py)", "Flask application serving the UI and all API endpoints on port 5000. Maintains an in-memory message cache loaded from SQLite at startup for high-performance queries."),
        ("Telegram Monitor (telegram_monitor.py)", "Async Telethon client running in live monitoring mode. Handles real-time message ingestion, media downloading, backfill operations, and channel discovery."),
        ("AI Agent (ai_agent.py)", "GPT-powered analysis engine running 6 automated loops: message enrichment, keyword suggestion, channel vetting, threat briefing, hunting leads, and escalation detection."),
        ("Dark Collector (dark_collector.py)", "Multi-source OSINT collector with 8 parallel loops monitoring paste sites, breach databases, certificate transparency logs, domain lookalikes, GitHub, ransomware leak sites, Tor hidden services, and threat intelligence feeds."),
        ("Orchestrator (orchestrator.py)", "Process manager that starts all components and auto-restarts crashed services with exponential backoff."),
    ]
    for title, desc in arch_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)
        run.font.name = "Calibri"
        run = p.add_run(desc)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 2 — CONVENTIONS
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "2", "Authentication & Common Conventions")

    add_styled_para(doc, "Authentication", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=4, space_after=6)
    add_styled_para(doc,
        "The platform operates as an internal deployment with session-based authentication. All API endpoints "
        "are accessible without bearer tokens in the default configuration. For production deployments, it is "
        "recommended to place the application behind a reverse proxy (nginx/Caddy) with TLS termination and "
        "implement appropriate access controls at the network level.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_styled_para(doc, "Common Response Format", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=4, space_after=6)
    add_styled_para(doc,
        "All endpoints return JSON (Content-Type: application/json) unless explicitly noted otherwise "
        "(CSV exports, DOCX/PDF reports, media files). Successful responses return HTTP 200. Error responses "
        "include an 'error' field with a descriptive message.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=6)

    add_code_block(doc, '// Success\n{"total": 1523, "critical": 89, "data": [...]}\n\n// Error\n{"error": "Channel not found", "code": 404}')

    add_styled_para(doc, "Timestamp Format", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)
    add_styled_para(doc,
        "All timestamps are returned in ISO 8601 format (YYYY-MM-DDTHH:MM:SS) in UTC timezone. "
        "Some endpoints also provide IRST (Iran Standard Time, UTC+3:30) timestamps for analyst convenience. "
        "Query parameters accept ISO 8601 format for date filtering.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=6)

    add_styled_para(doc, "Priority Classification", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)

    priority_data = [
        ["Priority", "Level", "Description"],
        ["CRITICAL", "High-impact threat indicators matching critical keyword patterns: attack claims, data leaks, zero-day exploits, Jordan-targeted operations, APT tool sharing", "Red"],
        ["MEDIUM", "Moderate relevance: geopolitical commentary, reconnaissance indicators, general hacktivism, tool discussions", "Amber"],
        ["LOW", "Minimal direct threat: general chatter, off-topic posts, automated bot messages, channel administration", "Gray"],
    ]
    make_endpoint_table(doc, [
        ["Priority", "Description", "Color Code"],
        ["CRITICAL", "High-impact: attack claims, data leaks, zero-days, Jordan-targeted operations, APT tool sharing", "Red"],
        ["MEDIUM", "Moderate: geopolitical commentary, reconnaissance, general hacktivism, tool discussions", "Amber"],
        ["LOW", "Minimal: general chatter, off-topic posts, bot messages, channel administration", "Gray"],
    ])

    add_styled_para(doc, "Critical Subtypes", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)
    make_endpoint_table(doc, [
        ["Subtype", "Description"],
        ["CYBER", "Technical cyber attacks: DDoS, defacement, malware deployment, data exfiltration"],
        ["NATIONAL", "National security threats: military targeting, government infrastructure, critical services"],
        ["BOTH", "Messages classified under both CYBER and NATIONAL categories"],
        ["GENERAL", "Critical messages not fitting specific subcategories"],
    ])

    add_styled_para(doc, "Pagination", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)
    add_styled_para(doc,
        "Endpoints supporting pagination accept 'page' and 'per_page' query parameters. When pagination is active, "
        "the response wraps data in an object with a 'pagination' field containing page metadata. When pagination "
        "parameters are omitted, the response returns a flat array (up to the endpoint's limit).",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=6)

    add_code_block(doc, '// Paginated response\n{\n  "data": [...],\n  "pagination": {\n    "page": 1,\n    "per_page": 50,\n    "total": 1523,\n    "pages": 31\n  }\n}')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 3 — MESSAGE & CHANNEL APIs
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "3", "Message & Channel Intelligence APIs")

    add_styled_para(doc,
        "These endpoints provide access to the core intelligence data: monitored Telegram channels and their "
        "ingested messages. The platform maintains an in-memory cache of all messages loaded from SQLite at "
        "startup, enabling sub-millisecond query performance for filtering and aggregation operations.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    # 3.1 Channel Listing
    add_subsection_header(doc, "3.1", "Channel Listing & Statistics")

    add_endpoint_section(doc, "GET", "/api/channels",
        "Returns a list of all monitored Telegram channels with aggregated statistics, activity sparklines, "
        "and threat classification metadata. This is the primary endpoint for populating the channel sidebar "
        "and monitoring dashboard.",
        response_fields=[
            ["channel", "string", "Display name of the Telegram channel"],
            ["channel_username", "string", "Telegram username identifier (used as key in other API calls)"],
            ["count", "integer", "Total number of ingested messages from this channel"],
            ["critical", "integer", "Count of CRITICAL priority messages"],
            ["medium", "integer", "Count of MEDIUM priority messages"],
            ["last_date", "string", "ISO timestamp of the most recent message"],
            ["last_critical_date", "string", "ISO timestamp of the most recent CRITICAL message"],
            ["spark_7d", "array[int]", "7-element array of daily message counts for sparkline visualization"],
            ["tier", "integer", "Channel tier ranking (1=highest threat)"],
            ["tier_label", "string", "Human-readable tier classification"],
            ["threat_level", "string", "Threat level assessment (e.g., 'Critical', 'High', 'Medium')"],
            ["status", "string", "Channel status: 'active', 'banned', 'dormant'"],
            ["days_silent", "integer", "Number of days since last message (0 = active today)"],
        ],
        example_response='[\n  {\n    "channel": "Handala Hack Team",\n    "channel_username": "handal_a",\n    "count": 2847,\n    "critical": 156,\n    "medium": 891,\n    "last_date": "2026-03-11T14:22:00",\n    "spark_7d": [12, 8, 15, 23, 11, 19, 7],\n    "tier": 1,\n    "tier_label": "Tier 1 - Primary Threat",\n    "threat_level": "Critical",\n    "status": "active",\n    "days_silent": 0\n  }\n]'
    )

    add_endpoint_section(doc, "GET", "/api/channel/<channel_username>/iocs",
        "Returns aggregated Indicators of Compromise (IOCs) extracted from a specific channel's messages. "
        "IOCs include IPv4 addresses, domains, SHA256/MD5 hashes, URLs, and email addresses. Results are "
        "ranked by frequency (top 50).",
        params=[
            ["channel_username", "string", "Yes (URL)", "—", "Telegram channel username"],
        ],
        response_fields=[
            ["channel", "string", "Channel username"],
            ["msg_count", "integer", "Total messages scanned in this channel"],
            ["iocs", "array", "Top 50 IOCs sorted by frequency"],
            ["iocs[].type", "string", "IOC type: 'ipv4', 'domain', 'sha256', 'md5', 'url', 'email'"],
            ["iocs[].value", "string", "The IOC value"],
            ["iocs[].count", "integer", "Number of messages containing this IOC"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/channel/<channel_username>/trend",
        "Returns daily message volume statistics for a specific channel over the past 30 days. "
        "Used for channel-level trend visualization and activity pattern analysis.",
        params=[
            ["channel_username", "string", "Yes (URL)", "—", "Telegram channel username"],
        ],
        response_fields=[
            ["date", "string", "Date in YYYY-MM-DD format"],
            ["total", "integer", "Total messages on this date"],
            ["critical", "integer", "CRITICAL messages on this date"],
            ["medium", "integer", "MEDIUM messages on this date"],
        ],
        response="Array of 30 daily objects"
    )

    doc.add_page_break()

    # 3.2 Message Retrieval
    add_subsection_header(doc, "3.2", "Message Retrieval & Filtering")

    add_endpoint_section(doc, "GET", "/api/messages/<channel_username>",
        "Retrieves messages from a specific channel with comprehensive filtering support. Returns messages "
        "sorted by timestamp (newest first). Supports priority filtering, date range selection, critical "
        "subtype classification, and full-text search.",
        params=[
            ["channel_username", "string", "Yes (URL)", "—", "Telegram channel username"],
            ["priority", "string", "No", "ALL", "Filter: ALL, CRITICAL, MEDIUM, LOW"],
            ["critical_subtype", "string", "No", "ALL", "Filter: ALL, CYBER, NATIONAL, BOTH, GENERAL"],
            ["since", "string", "No", "—", "ISO timestamp — return messages after this time"],
            ["until", "string", "No", "—", "ISO timestamp — return messages before this time"],
            ["search", "string", "No", "—", "Full-text search within message content"],
        ],
        response_fields=[
            ["message_id", "integer", "Telegram message ID"],
            ["channel", "string", "Channel display name"],
            ["channel_username", "string", "Channel username"],
            ["timestamp_utc", "string", "ISO timestamp in UTC"],
            ["text_preview", "string", "Message text (may be truncated)"],
            ["priority", "string", "CRITICAL, MEDIUM, or LOW"],
            ["keyword_hits", "array[string]", "Keywords that triggered the priority classification"],
            ["iocs", "object", "Extracted IOCs: {ipv4: [], domains: [], sha256: [], ...}"],
            ["critical_subtype", "string", "CYBER, NATIONAL, BOTH, or GENERAL (CRITICAL messages only)"],
            ["has_media", "boolean", "Whether the message contains media attachments"],
            ["media_path", "string", "Relative path to downloaded media file (if available)"],
            ["ai_enrichment", "object", "AI-generated enrichment data (if processed)"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/messages/all",
        "Unified timeline endpoint returning messages across all monitored channels. This is the primary "
        "endpoint for the platform's Timeline view. Supports all filtering options plus pagination, "
        "media/document filtering, and per-keyword/per-channel scoping.",
        params=[
            ["priority", "string", "No", "ALL", "Priority filter: ALL, CRITICAL, MEDIUM, LOW"],
            ["critical_subtype", "string", "No", "ALL", "Subtype filter: ALL, CYBER, NATIONAL, BOTH, GENERAL"],
            ["since", "string", "No", "—", "Start of date range (ISO 8601)"],
            ["until", "string", "No", "—", "End of date range (ISO 8601)"],
            ["search", "string", "No", "—", "Full-text search across all message content"],
            ["keyword", "string", "No", "—", "Filter by specific keyword hit"],
            ["channel", "string", "No", "—", "Filter by channel username"],
            ["has_media", "string", "No", "—", "Set to '1' to return only messages with media (images, videos, files)"],
            ["has_docs", "string", "No", "—", "Set to '1' to return only messages with document files (excludes images, videos, audio)"],
            ["limit", "integer", "No", "1000", "Maximum messages to return (max: 5000). Ignored when pagination is active."],
            ["page", "integer", "No", "—", "Page number (activates pagination mode)"],
            ["per_page", "integer", "No", "50", "Results per page (max: 200). Only used with pagination."],
        ],
        response_fields=[
            ["(flat mode)", "array", "Array of message objects when pagination is not active"],
            ["data", "array", "Message objects (pagination mode)"],
            ["pagination.page", "integer", "Current page number"],
            ["pagination.per_page", "integer", "Results per page"],
            ["pagination.total", "integer", "Total matching messages"],
            ["pagination.pages", "integer", "Total number of pages"],
        ],
        notes=[
            "Note: has_media and has_docs are mutually exclusive — only one should be set per request.",
            "Note: Document filter (has_docs) excludes image extensions (jpg, png, gif, webp, bmp, svg, ico), video extensions (mp4, webm, mov, avi, mkv, flv, wmv, m4v), and audio extensions (mp3, ogg, wav, flac, aac, m4a, wma, opus).",
            "Note: Only messages with actual downloaded files (media_path present) are returned when has_docs=1."
        ]
    )

    doc.add_page_break()

    # 3.3 Polling
    add_subsection_header(doc, "3.3", "Real-Time Polling")

    add_endpoint_section(doc, "GET", "/api/messages/poll",
        "Lightweight polling endpoint for real-time message updates. Uses database-level WHERE clause "
        "filtering for optimal performance. Designed to be called every 10-30 seconds by the frontend "
        "to check for new messages without reloading the full dataset.",
        params=[
            ["after", "string", "Yes", "—", "ISO timestamp — only return messages newer than this"],
            ["priority", "string", "No", "ALL", "Priority filter"],
            ["limit", "integer", "No", "200", "Maximum messages to return (max: 500)"],
        ],
        response_fields=[
            ["new_count", "integer", "Number of new messages since the 'after' timestamp"],
            ["newest_ts", "string", "Timestamp of the most recent message (use as next 'after' value)"],
            ["messages", "array", "Array of new message objects"],
        ],
        example_response='{\n  "new_count": 3,\n  "newest_ts": "2026-03-11T14:22:00",\n  "messages": [\n    {"message_id": 4521, "channel": "...", "priority": "CRITICAL", ...}\n  ]\n}'
    )

    add_endpoint_section(doc, "GET", "/api/messages/count",
        "Ultra-lightweight endpoint returning only the total message count and newest timestamp. "
        "Used by the UI to determine if a full data refresh is needed before making heavier API calls.",
        params=[
            ["priority", "string", "No", "ALL", "Priority filter"],
        ],
        response_fields=[
            ["count", "integer", "Total message count matching the filter"],
            ["newest_ts", "string", "Timestamp of the most recent matching message"],
        ],
    )

    # 3.4 Context
    add_subsection_header(doc, "3.4", "Message Context")

    add_endpoint_section(doc, "GET", "/api/messages/<channel_username>/<message_id>/context",
        "Retrieves surrounding messages for context analysis. Returns N messages before and after the "
        "target message from the same channel. Attempts live Telegram fetch first, falls back to stored "
        "database messages if the live fetch fails.",
        params=[
            ["channel_username", "string", "Yes (URL)", "—", "Channel username"],
            ["message_id", "string", "Yes (URL)", "—", "Target message ID"],
            ["before", "integer", "No", "5", "Number of messages before target (1-50)"],
            ["after", "integer", "No", "5", "Number of messages after target (1-50)"],
        ],
        response_fields=[
            ["messages", "array", "Array of context messages in chronological order"],
            ["target_idx", "integer", "Index of the target message within the array"],
            ["total", "integer", "Total messages in the channel"],
            ["source", "string", "'live' (fetched from Telegram) or 'stored' (from database)"],
        ],
    )

    # 3.5 Exports
    add_subsection_header(doc, "3.5", "Export Endpoints")

    add_endpoint_section(doc, "GET", "/api/messages/export",
        "Exports filtered messages as a downloadable CSV file. Default filter is CRITICAL priority. "
        "CSV columns: timestamp_utc, timestamp_irst, channel, channel_username, priority, keyword_hits, "
        "text_preview, iocs.",
        params=[
            ["priority", "string", "No", "CRITICAL", "Priority filter for export"],
            ["since", "string", "No", "—", "Start date filter"],
            ["search", "string", "No", "—", "Text search filter"],
        ],
        response="CSV file download (Content-Type: text/csv)",
    )

    add_endpoint_section(doc, "GET", "/api/iocs/export",
        "Exports all aggregated IOCs across all channels as a CSV file. Includes IOC type, value, "
        "occurrence count, associated channels, and last seen timestamp. Useful for integration with "
        "SIEM/SOAR platforms and firewall rule generation.",
        response="CSV file download — columns: type, value, count, channels, last_seen",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 4 — DASHBOARD & ANALYTICS
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "4", "Dashboard & Analytics APIs")

    add_styled_para(doc,
        "These endpoints power the platform's analytical dashboards, providing aggregated views of threat "
        "intelligence data including activity heatmaps, keyword frequency analysis, IOC rankings, coordinated "
        "campaign detection, and trend analysis over configurable time periods.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "4.1", "Dashboard Aggregation")

    add_endpoint_section(doc, "GET", "/api/dashboard",
        "Returns comprehensive aggregated intelligence data for the main dashboard view. Includes total "
        "counts, keyword frequency analysis with weighted scoring, IOC rankings, hourly activity heatmap "
        "(7x24 matrix), coordinated campaign detection, and channel threat rankings.",
        response_fields=[
            ["total", "integer", "Total messages across all channels"],
            ["critical", "integer", "Total CRITICAL messages"],
            ["medium", "integer", "Total MEDIUM messages"],
            ["ioc_count", "integer", "Total unique IOCs extracted"],
            ["channel_count", "integer", "Number of active monitored channels"],
            ["total_configured", "integer", "Total channels in configuration"],
            ["banned_count", "integer", "Number of banned/suspended channels"],
            ["keywords", "array", "Top 100 keywords by weighted score: {keyword, total, critical, medium, weight}"],
            ["iocs", "array", "Top 500 IOCs: {type, value, count, channels, last_seen}"],
            ["activity_matrix", "object", "7-day x 24-hour heatmap: {weekday_name: [24 hourly counts]}"],
            ["campaigns", "array", "Top 30 coordinated campaigns: {keyword, channels, date, count}"],
            ["ch_ranking", "array", "Top 20 channels by CRITICAL volume: {channel, count, last_critical}"],
        ],
        notes=["Note: Keyword weight formula: (critical_count * 3) + medium_count + log2(total_count)"]
    )

    add_subsection_header(doc, "4.2", "Trend Analysis")

    add_endpoint_section(doc, "GET", "/api/trend",
        "Returns daily message volume breakdown by priority level over a configurable time period. "
        "Used for the main trend chart on the dashboard.",
        params=[
            ["days", "integer", "No", "60", "Number of days to include in the trend (max: 365)"],
        ],
        response_fields=[
            ["date", "string", "Date in YYYY-MM-DD format"],
            ["CRITICAL", "integer", "CRITICAL message count for this date"],
            ["MEDIUM", "integer", "MEDIUM message count for this date"],
            ["LOW", "integer", "LOW message count for this date"],
        ],
        response="Array of daily objects"
    )

    add_subsection_header(doc, "4.3", "Intelligence Briefing")

    add_endpoint_section(doc, "GET", "/api/briefing",
        "Generates a 24-hour intelligence briefing with contextual 7-day comparison. Provides executive-level "
        "summary of threat activity including top targeted entities, active channels, fresh IOCs, and the "
        "most recent critical messages.",
        response_fields=[
            ["generated_at", "string", "Briefing generation timestamp"],
            ["period_hours", "integer", "Reporting period (24)"],
            ["summary", "object", "Counts: {total_24h, critical_24h, medium_24h, channels_active}"],
            ["top_targeted_entities", "array", "Most mentioned targets: [{keyword, count}]"],
            ["active_channels", "array", "Channels with recent activity: [{channel, critical, medium, total}]"],
            ["fresh_iocs", "object", "New IOCs by type: {ipv4: [...], domain: [...], sha256: [...]}"],
            ["newest_critical", "array", "Last 5 CRITICAL messages with full details"],
            ["weekly_critical", "integer", "7-day CRITICAL count for trend comparison"],
        ],
    )

    add_subsection_header(doc, "4.4", "Status Summary")

    add_endpoint_section(doc, "GET", "/api/stats/summary",
        "Ultra-lightweight endpoint providing real-time statistics for the UI status bar. Uses direct SQL "
        "queries for minimal latency.",
        response_fields=[
            ["total", "integer", "Total messages in database"],
            ["critical", "integer", "Total CRITICAL messages"],
            ["medium", "integer", "Total MEDIUM messages"],
            ["critical_24h", "integer", "CRITICAL messages in last 24 hours"],
            ["critical_1h", "integer", "CRITICAL messages in last hour"],
            ["channels", "integer", "Number of monitored channels"],
            ["iocs", "integer", "Total unique IOCs"],
            ["generated_at", "string", "Timestamp of this summary"],
        ],
    )

    add_subsection_header(doc, "4.5", "Threat Matrix")

    add_endpoint_section(doc, "GET", "/api/threat_matrix",
        "Returns a matrix visualization dataset mapping threat actors (channels) against Jordan target "
        "categories. Only considers CRITICAL messages for sector classification. Categories include "
        "Banking, Telecom, Government/Military, ISP/Network, Media, Infrastructure, and General.",
        response_fields=[
            ["categories", "array[string]", "Target sector names: ['Banking', 'Telecom', 'Gov/Mil', ...]"],
            ["actors", "array", "Threat actors with per-sector counts"],
            ["actors[].channel", "string", "Channel username"],
            ["actors[].label", "string", "Channel display label"],
            ["actors[].tier", "integer", "Threat tier"],
            ["actors[].threat", "string", "Threat level"],
            ["actors[].total", "integer", "Total CRITICAL messages"],
            ["actors[].<Category>", "integer", "Count of messages targeting this sector"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 5 — APT TRACKER
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "5", "APT Tracker APIs")

    add_styled_para(doc,
        "The APT Tracker module provides advanced threat actor profiling, collaboration network analysis, "
        "and deep IOC research capabilities. APT groups are identified by grouping related Telegram channels "
        "based on naming patterns, shared infrastructure, and coordinated posting behavior.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "5.1", "APT Profiles")

    add_endpoint_section(doc, "GET", "/api/apt/profiles",
        "Returns aggregated profiles for all identified APT groups. Each profile includes channel associations, "
        "message statistics, IOC counts, sector targeting analysis, attack type classification, and temporal "
        "activity boundaries.",
        response_fields=[
            ["name", "string", "APT group name (derived from channel naming patterns)"],
            ["tier", "integer", "Threat tier (1=highest)"],
            ["threat", "string", "Threat level classification"],
            ["channels", "array[string]", "Associated Telegram channel usernames"],
            ["status", "string", "Operational status"],
            ["total_msgs", "integer", "Total messages across all associated channels"],
            ["critical_count", "integer", "CRITICAL message count"],
            ["medium_count", "integer", "MEDIUM message count"],
            ["ioc_count", "integer", "Total extracted IOCs"],
            ["ioc_malicious", "integer", "IOCs verified as malicious via external research"],
            ["ioc_suspicious", "integer", "IOCs flagged as suspicious"],
            ["sectors", "object", "Top 10 targeted sectors: {sector_name: count}"],
            ["attack_types", "object", "Top 10 attack types: {type_name: count}"],
            ["first_seen", "string", "Earliest message timestamp"],
            ["last_seen", "string", "Most recent message timestamp"],
            ["jordan_attacks", "integer", "Messages specifically targeting Jordan entities"],
        ],
    )

    add_subsection_header(doc, "5.2", "APT Network Map")

    add_endpoint_section(doc, "GET", "/api/apt/network",
        "Computes and returns a collaboration network graph between APT groups based on five distinct "
        "relationship metrics. This endpoint powers the D3.js force-directed graph visualization in the "
        "APT Tracker's Network Map sub-tab.",
        response_fields=[
            ["generated_at", "string", "Computation timestamp"],
            ["nodes", "array", "APT group nodes with metadata"],
            ["nodes[].id", "string", "APT group name"],
            ["nodes[].tier", "integer", "Threat tier"],
            ["nodes[].threat", "string", "Threat level"],
            ["nodes[].msg_count", "integer", "Total messages"],
            ["nodes[].connections", "integer", "Number of connected APTs"],
            ["edges", "array", "Collaboration edges between APTs"],
            ["edges[].source", "string", "Source APT name"],
            ["edges[].target", "string", "Target APT name"],
            ["edges[].types", "array[string]", "Relationship types: mention, parallel, infrastructure, sector, temporal"],
            ["edges[].weight", "float", "Combined edge weight (higher = stronger collaboration signal)"],
            ["edges[].examples", "array", "Evidence examples for each relationship type"],
            ["metrics", "object", "Network summary statistics"],
            ["metrics.total_apts", "integer", "Total APT groups in network"],
            ["metrics.total_edges", "integer", "Total collaboration edges"],
            ["metrics.mention_edges", "integer", "Edges from cross-group mentions"],
            ["metrics.parallel_edges", "integer", "Edges from parallel target claims"],
        ],
        notes=[
            "Note: Mention edges are detected via word-boundary regex matching of APT names in message text.",
            "Note: Parallel edges require different APTs claiming same target within 48-hour window.",
            "Note: Infrastructure edges are based on shared IPs, domains, or file hashes across APT channels.",
            "Note: Sector edges use weighted Jaccard similarity of target sector distributions.",
            "Note: Temporal edges identify co-active peak days where both APTs exceed 1.5x average activity."
        ]
    )

    add_subsection_header(doc, "5.3", "APT Detail")

    add_endpoint_section(doc, "GET", "/api/apt/<name>/detail",
        "Returns comprehensive detail for a specific APT group including attack timeline, sector distribution, "
        "IOC analysis, and recent critical messages. The name parameter supports URL-encoded values for APT "
        "names containing spaces or special characters.",
        params=[
            ["name", "string", "Yes (URL path)", "—", "APT group name (URL-encoded if necessary)"],
        ],
        response_fields=[
            ["name", "string", "APT group name"],
            ["total_msgs", "integer", "Total message count"],
            ["critical_count", "integer", "CRITICAL message count"],
            ["medium_count", "integer", "MEDIUM message count"],
            ["first_seen", "string", "Earliest activity timestamp"],
            ["last_seen", "string", "Most recent activity timestamp"],
            ["sectors", "object", "Target sector distribution"],
            ["attack_types", "object", "Attack type distribution"],
            ["attacks", "array", "Cataloged attacks: [{date, target, type, channel, msg_id, summary}]"],
            ["recent_critical", "array", "Last 20 CRITICAL messages with full text"],
            ["timeline", "object", "Monthly activity: {YYYY-MM: {critical, medium, low}}"],
        ],
    )

    doc.add_page_break()

    add_subsection_header(doc, "5.4", "IOC Lookup & Research")

    add_endpoint_section(doc, "POST", "/api/apt/ioc/lookup",
        "Multi-source IOC lookup combining local database search with external threat intelligence enrichment. "
        "Supports automatic IOC type detection for IPv4, domains, SHA256, MD5, URLs, and email addresses. "
        "External enrichment includes AbuseIPDB scoring for IP addresses.",
        body=[
            ["value", "string", "Yes", "The IOC value to look up (IP, domain, hash, URL, or email)"],
            ["type", "string", "No", "IOC type hint: 'auto' (default), 'ipv4', 'domain', 'sha256', 'md5', 'url', 'email'"],
        ],
        response_fields=[
            ["value", "string", "The queried IOC"],
            ["type", "string", "Detected IOC type"],
            ["local.found", "boolean", "Whether the IOC exists in local database"],
            ["local.count", "integer", "Number of messages containing this IOC"],
            ["local.apts", "array[string]", "APT groups associated with this IOC"],
            ["local.channels", "array[string]", "Channels where this IOC appeared"],
            ["local.first_seen", "string", "Earliest occurrence"],
            ["local.last_seen", "string", "Most recent occurrence"],
            ["local.messages", "array", "Message references: [{msg_id, channel, timestamp, summary_snippet}]"],
            ["abuseipdb", "object", "AbuseIPDB enrichment (IPs only)"],
            ["abuseipdb.abuseConfidenceScore", "integer", "Confidence score 0-100"],
            ["abuseipdb.countryCode", "string", "Country code"],
            ["abuseipdb.isp", "string", "Internet Service Provider"],
            ["abuseipdb.totalReports", "integer", "Number of abuse reports"],
            ["verdict", "string", "Computed verdict: MALICIOUS, SUSPICIOUS, or CLEAN"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/apt/<name>/research",
        "Returns external threat intelligence research results for a specific APT group's IOCs. Research "
        "data includes results from ThreatFox, URLhaus, MalwareBazaar, OTX AlienVault, CISA KEV, and "
        "RSS threat feeds. Results are cached to avoid redundant API calls.",
        params=[
            ["name", "string", "Yes (URL path)", "—", "APT group name"],
        ],
        response_fields=[
            ["name", "string", "APT group name"],
            ["researched_at", "string", "Research timestamp"],
            ["cached", "boolean", "Whether results were served from cache"],
            ["stats.total", "integer", "Total IOCs researched"],
            ["stats.malicious", "integer", "IOCs confirmed malicious"],
            ["stats.suspicious", "integer", "IOCs flagged suspicious"],
            ["stats.unverified", "integer", "IOCs with no external data"],
            ["iocs", "array", "Research results: [{type, value, abuse_verdict, abuse_score, context, sources}]"],
        ],
    )

    add_subsection_header(doc, "5.5", "AI-Powered IOC Extraction")

    add_endpoint_section(doc, "POST", "/api/apt/ioc/ai-extract",
        "Triggers AI-powered deep IOC extraction for a specific APT group. Uses GPT-4o-mini to analyze "
        "the full message corpus and identify IOCs that may have been missed by regex-based extraction "
        "(e.g., obfuscated IPs, encoded domains, contextual file references).",
        body=[
            ["apt_name", "string", "Yes", "APT group name to analyze"],
        ],
        response_fields=[
            ["apt_name", "string", "APT group analyzed"],
            ["messages_scanned", "integer", "Number of messages processed"],
            ["chunks_processed", "integer", "Number of GPT API chunks submitted"],
            ["new_iocs", "array", "Newly discovered IOCs: [{type, value, context, sources}]"],
            ["total_found", "integer", "Total IOCs identified by AI"],
            ["already_known", "integer", "IOCs already in database (not duplicated)"],
        ],
        notes=["Note: Requires OPENAI_API_KEY environment variable. Cost: ~$0.02-0.05 per extraction run."]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 6 — BLOCKLIST
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "6", "Blocklist & IOC Management APIs")

    add_styled_para(doc,
        "The blocklist module provides a centralized, deduplicated repository of all researched IOCs across "
        "all APT groups. It supports filtering, export, and automated report generation for SOC advisory "
        "distribution.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "6.1", "Central Blocklist")

    add_endpoint_section(doc, "GET", "/api/blocklist",
        "Returns the deduplicated central blocklist aggregating IOCs from all researched APT groups. "
        "Supports filtering by APT, IOC type, verdict, and free-text search. Pagination is available "
        "for large result sets.",
        params=[
            ["apt", "string", "No", "—", "Filter by APT group name"],
            ["type", "string", "No", "—", "Filter by IOC type: ipv4, domain, sha256, md5, url"],
            ["verdict", "string", "No", "—", "Filter by verdict: MALICIOUS, SUSPICIOUS, CLEAN"],
            ["q", "string", "No", "—", "Free-text search across IOC values and context"],
            ["page", "integer", "No", "—", "Page number (activates pagination)"],
            ["per_page", "integer", "No", "100", "Results per page (max: 500)"],
        ],
        response_fields=[
            ["total", "integer", "Total IOCs matching filters"],
            ["apts_researched", "integer", "Number of APT groups with research data"],
            ["malicious", "integer", "Count of MALICIOUS verdicts"],
            ["suspicious", "integer", "Count of SUSPICIOUS verdicts"],
            ["clean", "integer", "Count of CLEAN verdicts"],
            ["apt_summary", "object", "Per-APT statistics: {apt_name: {total, malicious, suspicious}}"],
            ["iocs", "array", "IOC records: [{type, value, apt, abuse_verdict, abuse_score, context}]"],
            ["pagination", "object", "Present when page param is set: {page, per_page, total, pages}"],
        ],
    )

    add_subsection_header(doc, "6.2", "Blocklist Export")

    add_endpoint_section(doc, "GET", "/api/blocklist/export",
        "Exports the blocklist as a CSV file for integration with SIEM platforms, firewalls, and other "
        "security tooling. CSV columns include APT attribution, IOC metadata, AbuseIPDB enrichment, "
        "and contextual intelligence.",
        params=[
            ["verdict", "string", "No", "—", "Filter export by verdict (MALICIOUS, SUSPICIOUS, CLEAN)"],
        ],
        response="CSV file download — columns: APT Group, IOC Type, Value, Verdict, AbuseIPDB Score, Country, ISP, Reports, Source, Context, Confidence",
    )

    add_subsection_header(doc, "6.3", "Report Generation")

    add_endpoint_section(doc, "GET", "/api/blocklist/report",
        "Generates a professional ScanWave SOC Client Advisory report from the current blocklist data. "
        "The report includes executive summary, IOC tables categorized by type and verdict, APT attribution, "
        "recommended actions, and legal disclaimer. Output format is DOCX with automatic PDF conversion "
        "via LibreOffice when available.",
        response="DOCX or PDF file download (Content-Disposition: attachment)",
        notes=[
            "Note: Uses scanwave_report_template.docx as the branded template.",
            "Note: PDF conversion requires LibreOffice installed (available in Docker deployment).",
            "Note: Report includes mandatory legal disclaimer regarding OSINT-only collection.",
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 7 — DISCOVERY ENGINE
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "7", "Discovery Engine APIs")

    add_styled_para(doc,
        "The Discovery Engine automatically identifies new Telegram channels potentially associated with "
        "monitored threat actors. It uses forward network analysis, username similarity matching, and "
        "keyword-based content scoring to surface channels for analyst review.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "7.1", "Discovery Listing")

    add_endpoint_section(doc, "GET", "/api/discovery/list",
        "Returns all auto-discovered channels with their discovery metadata, confidence scores, and "
        "current review status.",
        response_fields=[
            ["username", "string", "Discovered channel username"],
            ["reason", "string", "Discovery reason (e.g., 'forward_from: handal_a', 'keyword_match')"],
            ["confidence", "float", "Relevance confidence score (0.0 - 1.0)"],
            ["added_date", "string", "Discovery timestamp"],
            ["status", "string", "Review status: 'pending_review', 'approved', 'dismissed'"],
        ],
    )

    add_subsection_header(doc, "7.2", "Scan Trigger & Status")

    add_endpoint_section(doc, "POST", "/api/discovery/scan",
        "Triggers a channel discovery scan as a background process. The scan analyzes forward networks, "
        "message references, and username patterns to identify new potentially relevant channels.",
        body=[
            ["quick", "boolean", "No", "If true, runs a faster scan with reduced depth"],
        ],
        response_fields=[
            ["started", "boolean", "Whether the scan was successfully initiated"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/discovery/status",
        "Returns the current status of the discovery scan process.",
        response_fields=[
            ["running", "boolean", "Whether a scan is currently active"],
            ["pid", "integer", "Process ID of the running scan (if active)"],
            ["last_run", "string", "Timestamp of the most recent completed scan"],
        ],
    )

    add_subsection_header(doc, "7.3", "On-Demand Channel Fetch")

    add_endpoint_section(doc, "GET", "/api/discovery/fetch",
        "Fetches recent messages from a specified channel on-demand via Telethon. Useful for analyst "
        "preview of discovered channels before approving them for monitoring.",
        params=[
            ["username", "string", "Yes", "—", "Telegram channel username to fetch from"],
            ["limit", "integer", "No", "100", "Maximum messages to fetch (max: 500)"],
            ["save", "string", "No", "false", "Set to 'true' to persist fetched messages to the database"],
        ],
        response_fields=[
            ["messages", "array", "Fetched message objects with full text and metadata"],
            ["total", "integer", "Number of messages returned"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 8 — DARK WEB INTELLIGENCE
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "8", "Dark Web & Threat Intelligence APIs")

    add_styled_para(doc,
        "These endpoints expose findings from the dark_collector.py module which operates 8 parallel "
        "collection loops monitoring paste sites, breach databases, certificate transparency logs, "
        "domain lookalikes, GitHub repositories, ransomware leak sites, Tor hidden services, and "
        "threat intelligence feeds (ThreatFox, URLhaus, MalwareBazaar, CISA KEV, OTX AlienVault, "
        "16 RSS feeds).",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "8.1", "Dark Intel Feed")

    add_endpoint_section(doc, "GET", "/api/dark/feed",
        "Returns paginated dark web intelligence findings with filtering by severity and source. "
        "Findings include paste site mentions, breach entries, certificate anomalies, domain lookalikes, "
        "GitHub exposure, ransomware listings, Tor hidden service content, and threat feed matches.",
        params=[
            ["limit", "integer", "No", "100", "Maximum findings to return (max: 500)"],
            ["severity", "string", "No", "—", "Filter by severity level"],
            ["source", "string", "No", "—", "Filter by source: ThreatFox, URLhaus, MalwareBazaar, RSS, CISA_KEV, OTX, Ransomware, Iranian_APT, paste, breach, certstream, dnstwist, github, dark_web"],
            ["after", "string", "No", "—", "Only return findings after this timestamp"],
        ],
        response_fields=[
            ["count", "integer", "Number of findings returned"],
            ["findings", "array", "Finding objects"],
            ["findings[].timestamp", "string", "Discovery timestamp"],
            ["findings[].source", "string", "Collection source identifier"],
            ["findings[].severity", "string", "Severity assessment"],
            ["findings[].value", "string", "Primary finding value (IOC, domain, paste URL, etc.)"],
            ["findings[].context", "string", "Contextual description"],
        ],
    )

    add_subsection_header(doc, "8.2", "Dark Collector Statistics")

    add_endpoint_section(doc, "GET", "/api/dark/stats",
        "Returns operational statistics for the dark collector including finding counts by source "
        "and severity, false positive tracking, Tor circuit rotation count, and last scan timestamps "
        "for each collection loop.",
        response_fields=[
            ["total_findings", "integer", "Total findings collected"],
            ["false_positives", "integer", "Findings marked as false positives"],
            ["tor_circuits_rotated", "integer", "Number of Tor circuit rotations performed"],
            ["by_source", "object", "Findings count per source: {paste: N, breach: N, ...}"],
            ["by_severity", "object", "Findings count per severity: {critical: N, high: N, ...}"],
            ["last_scans", "object", "Last scan timestamps per loop: {paste, breach, certstream, dnstwist, github, ransomware, dark_web, intel_feeds, daily_digest}"],
        ],
    )

    add_subsection_header(doc, "8.3", "Daily Digest")

    add_endpoint_section(doc, "GET", "/api/dark/digest",
        "Returns the latest AI-generated daily digest summarizing dark web intelligence findings. "
        "The digest is generated by GPT-4o analyzing the day's findings and producing an executive "
        "summary with key takeaways and recommended actions.",
        response="JSON object with digest content and generation timestamp",
    )

    add_subsection_header(doc, "8.4", "Domain Squatting")

    add_endpoint_section(doc, "GET", "/api/dark/domains",
        "Returns domain squatting detection results from dnstwist scans of monitored .jo (Jordan) "
        "domains. Identifies potential phishing domains, typosquatting, and homoglyph attacks.",
        response_fields=[
            ["targets_scanned", "array", "List of legitimate domains that were scanned"],
            ["known_lookalikes", "array", "Previously identified lookalike domains"],
            ["total_lookalikes", "integer", "Total count of detected domain lookalikes"],
        ],
    )

    add_subsection_header(doc, "8.5", "Behavioral Threat Level")

    add_endpoint_section(doc, "GET", "/api/dark/threat_level",
        "Computes and returns a composite behavioral threat level based on 7 signal dimensions derived "
        "from message volume anomalies, target convergence patterns, Jordan-specific targeting, tool "
        "sharing indicators, forward network activation, member surge detection, and message deletion "
        "spike analysis.",
        response_fields=[
            ["current_level", "string", "Threat level: CRITICAL, HIGH, ELEVATED, GUARDED, or LOW"],
            ["score", "float", "Composite score from 0.0 (lowest) to 1.0 (highest)"],
            ["updated_at", "string", "Last computation timestamp"],
            ["signals.volume_anomaly", "float", "Message volume deviation from baseline"],
            ["signals.target_convergence", "float", "Multiple APTs converging on same targets"],
            ["signals.jordan_targeting", "float", "Jordan-specific attack keyword density"],
            ["signals.tool_sharing", "float", "Cross-group tool/exploit sharing indicators"],
            ["signals.forward_activation", "float", "Forward network activity spike"],
            ["signals.member_surge", "float", "Unusual member count increases"],
            ["signals.deletion_spike", "float", "Message deletion rate anomaly"],
            ["triggered_alerts", "array", "Specific threshold-crossing alerts"],
            ["raw", "object", "Raw computation data for debugging"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 9 — AI AGENT
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "9", "AI Agent & Enrichment APIs")

    add_styled_para(doc,
        "The AI Agent module (ai_agent.py) runs 6 automated analysis loops powered by OpenAI GPT models. "
        "These endpoints provide status monitoring, control, and access to AI-generated intelligence products "
        "including enriched alerts, keyword suggestions, threat briefs, and on-demand translation.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "9.1", "Agent Status & Control")

    add_endpoint_section(doc, "GET", "/api/ai/status",
        "Returns comprehensive AI agent status including process health, enrichment statistics, keyword "
        "suggestion counts, channel vetting results, and brief generation history.",
        response_fields=[
            ["agent_running", "boolean", "Whether ai_agent.py process is active"],
            ["enrichments_done", "integer", "Total messages enriched"],
            ["keywords_added", "integer", "Keywords suggested and applied"],
            ["channels_autoapproved", "integer", "Channels automatically approved by AI"],
            ["channels_autodismissed", "integer", "Channels automatically dismissed by AI"],
            ["briefs_generated", "integer", "Threat briefs generated"],
            ["last_kw_run", "string", "Last keyword analysis run timestamp"],
            ["last_brief_run", "string", "Last brief generation timestamp"],
            ["latest_brief", "object", "Most recent threat brief content"],
        ],
    )

    add_endpoint_section(doc, "POST", "/api/ai/analyze",
        "Launches the AI agent as a background daemon process. If the agent is already running, "
        "returns the existing process ID.",
        response_fields=[
            ["started", "boolean", "Whether the agent was successfully started"],
            ["pid", "integer", "Process ID of the agent"],
        ],
        notes=["Note: Requires OPENAI_API_KEY environment variable to be set."]
    )

    add_endpoint_section(doc, "POST", "/api/ai/stop",
        "Terminates all running AI agent processes.",
        response_fields=[
            ["stopped", "array[int]", "Process IDs that were terminated"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/ai/log",
        "Returns the tail of the AI agent log file for debugging and monitoring.",
        params=[
            ["lines", "integer", "No", "80", "Number of log lines to return"],
        ],
        response_fields=[
            ["lines", "array[string]", "Log file lines"],
            ["total", "integer", "Total lines in log file"],
            ["size", "integer", "Log file size in bytes"],
        ],
    )

    add_subsection_header(doc, "9.2", "Keyword Suggestions")

    add_endpoint_section(doc, "GET", "/api/ai/suggestions",
        "Returns AI-generated keyword suggestions derived from analysis of recent messages. The AI "
        "identifies emerging threat terminology, new attack patterns, and previously uncategorized "
        "keywords that should be added to the monitoring filters.",
        response_fields=[
            ["runs", "array", "History of suggestion runs with timestamps"],
            ["latest", "object", "Most recent suggestions"],
            ["latest.new_critical_keywords", "array[string]", "Suggested CRITICAL keywords"],
            ["latest.new_medium_keywords", "array[string]", "Suggested MEDIUM keywords"],
        ],
    )

    add_endpoint_section(doc, "POST", "/api/ai/apply",
        "Applies the latest AI keyword suggestions to the active keyword filter lists (keywords.json). "
        "Merges suggested keywords without removing existing ones.",
        response_fields=[
            ["ok", "boolean", "Whether the operation succeeded"],
            ["added_critical", "integer", "Number of new CRITICAL keywords added"],
            ["added_medium", "integer", "Number of new MEDIUM keywords added"],
        ],
    )

    add_subsection_header(doc, "9.3", "Threat Briefs")

    add_endpoint_section(doc, "GET", "/api/ai/brief",
        "Returns the latest AI-generated threat intelligence brief. Briefs are automatically generated "
        "every 6 hours by the AI agent's LOOP 1 and provide executive-level summaries of threat actor "
        "activity, emerging patterns, and recommended defensive actions.",
        response="JSON object with brief content, generation timestamp, and analysis period",
    )

    add_subsection_header(doc, "9.4", "Enriched Alerts")

    add_endpoint_section(doc, "GET", "/api/ai/enriched",
        "Returns recent AI-enriched CRITICAL alerts. Enrichment adds contextual intelligence including "
        "APT group attribution, target sector classification, attack type categorization, confidence "
        "scoring, and recommended response actions.",
        params=[
            ["limit", "integer", "No", "20", "Number of enriched alerts to return (max: 100)"],
        ],
        response="Array of enriched alert objects from enriched_alerts.jsonl",
    )

    add_subsection_header(doc, "9.5", "Translation")

    add_endpoint_section(doc, "POST", "/api/translate",
        "On-demand translation of Arabic and Farsi text to English using GPT-4o-mini. Designed for "
        "analyst use when reviewing messages in non-English languages.",
        body=[
            ["text", "string", "Yes", "Text to translate (max 2000 characters)"],
        ],
        response_fields=[
            ["translation", "string", "English translation of the input text"],
        ],
        notes=["Note: Requires OPENAI_API_KEY. Cost: ~$0.001 per translation."]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 10 — ESCALATION & HUNTING
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "10", "Escalation & Threat Hunting APIs")

    add_styled_para(doc,
        "These endpoints provide proactive threat detection capabilities. The escalation detector monitors "
        "for sudden threat level changes, while the hunting module generates investigative leads for "
        "SOC analysts based on pattern analysis and behavioral anomalies.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_endpoint_section(doc, "GET", "/api/escalation/status",
        "Returns the current escalation state. The escalation detector (AI LOOP 6) runs every 15 minutes "
        "and evaluates multiple threat signals to determine if an escalation condition exists.",
        response_fields=[
            ["escalation_detected", "boolean", "Whether an active escalation is in progress"],
            ["urgency", "string", "Urgency level: IMMEDIATE, HIGH, MODERATE, LOW"],
            ["signals", "array", "Contributing threat signals that triggered the escalation"],
            ["summary", "string", "Human-readable escalation summary"],
            ["recommended_action", "string", "Suggested response action for SOC analysts"],
            ["checked_at", "string", "Last evaluation timestamp"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/hunting/leads",
        "Returns threat hunting leads generated by AI LOOP 5. Leads are categorized into group-level "
        "(APT behavior patterns) and operation-level (specific campaign indicators) investigations, "
        "ranked by confidence score.",
        response_fields=[
            ["group_leads", "array", "Top 100 group-level hunting leads: [{name, confidence, evidence}]"],
            ["operation_leads", "array", "Top 50 operation-level leads: [{name, confidence, evidence}]"],
            ["runs", "array", "History of hunting analysis runs"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/network/graph",
        "Returns a channel relationship graph based on forward network analysis. Identifies unknown "
        "channels that frequently appear in message forwards from monitored channels, scores them "
        "for relevance, and queues high-confidence channels for discovery review.",
        response_fields=[
            ["generated_at", "string", "Graph generation timestamp"],
            ["monitored_channels", "integer", "Number of monitored channels in the graph"],
            ["unknown_channels_scored", "integer", "Unknown channels that were scored"],
            ["newly_queued", "integer", "Channels newly added to discovery queue"],
            ["top_unknown", "array", "Top-ranked unknown channels with scores"],
            ["edges", "array", "Top 50 relationship edges: [{source, target, type, count}]"],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 11 — CHAT
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "11", "Chat & Intelligence Query APIs")

    add_styled_para(doc,
        "The platform includes an interactive AI-powered chatbot that allows analysts to query the "
        "intelligence database using natural language. The chatbot supports both synchronous and "
        "streaming response modes, with an agentic tool-calling loop that can search messages, "
        "look up IOCs, and cross-reference threat data.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "11.1", "Standard Chat")

    add_endpoint_section(doc, "POST", "/api/chat",
        "Synchronous chat endpoint that processes the query and returns a complete response with "
        "message references. For simple queries, uses direct database lookup. For complex queries, "
        "uses an agentic tool-calling loop with up to 12 iterations.",
        body=[
            ["message", "string", "Yes", "Natural language query about the intelligence data"],
            ["history", "array", "No", "Previous conversation messages for context: [{role, content}]"],
        ],
        response_fields=[
            ["answer", "string", "AI-generated response text"],
            ["references", "array[int]", "Indices of referenced messages"],
            ["ref_messages", "array", "Full message objects for citations"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/chat/history",
        "Returns the current session's chat history. Note: History is session-only and managed "
        "primarily on the frontend via the _chatHistory JavaScript array.",
        response_fields=[
            ["messages", "array", "Chat history messages (typically empty on backend)"],
        ],
    )

    add_endpoint_section(doc, "POST", "/api/chat/reset",
        "Resets the chat session state.",
        response_fields=[
            ["ok", "boolean", "Confirmation"],
        ],
    )

    add_subsection_header(doc, "11.2", "Streaming Chat (SSE)")

    add_endpoint_section(doc, "POST", "/api/chat/stream",
        "Server-Sent Events (SSE) streaming endpoint for real-time chat responses. Uses an agentic "
        "tool-calling loop that iteratively searches the intelligence database. For comprehensive "
        "queries, employs a map-reduce strategy: splits data into chunks, processes each with "
        "GPT-4o-mini, then synthesizes results with GPT-4o.",
        body=[
            ["message", "string", "Yes", "Natural language intelligence query"],
            ["history", "array", "No", "Previous conversation context"],
        ],
        response="Server-Sent Events stream (Content-Type: text/event-stream)",
        notes=[
            "Note: SSE events include 'data: {chunk}' for text tokens and 'data: [DONE]' for completion.",
            "Note: Comprehensive queries cost approximately $0.04 per query (map-reduce with GPT-4o).",
            "Note: The chatbot always responds in English regardless of the input language.",
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 12 — ADMIN
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "12", "Administration APIs")

    add_styled_para(doc,
        "Administration endpoints provide system monitoring, configuration management, keyword editing, "
        "channel management, database maintenance, and discovery workflow controls. These endpoints are "
        "intended for platform administrators and senior SOC analysts.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_subsection_header(doc, "12.1", "System Status & Monitoring")

    add_endpoint_section(doc, "GET", "/api/admin/status",
        "Returns comprehensive system health information including monitor process status, database "
        "statistics, recent log entries, and backfill queue state.",
        response_fields=[
            ["monitor_running", "boolean", "Whether telegram_monitor.py is active"],
            ["cursor", "object", "Monitor's current position in channel history"],
            ["db.total", "integer", "Total messages in SQLite database"],
            ["db.critical", "integer", "Total CRITICAL messages"],
            ["db.medium", "integer", "Total MEDIUM messages"],
            ["db.channels", "integer", "Number of channels with messages"],
            ["db.iocs", "integer", "Total unique IOCs"],
            ["db.last_message", "string", "Timestamp of most recent message"],
            ["log_tail", "array[string]", "Last 30 lines of monitor log"],
            ["backfill_queue", "object", "Current backfill queue contents"],
        ],
    )

    add_endpoint_section(doc, "GET", "/api/system/status",
        "Returns orchestrator-level component status showing which platform processes are running.",
        response_fields=[
            ["processes.viewer", "boolean", "Web server status"],
            ["processes.monitor", "boolean", "Telegram monitor status"],
            ["processes.ai_agent", "boolean", "AI agent status"],
            ["processes.orchestrator", "boolean", "Orchestrator status"],
            ["has_openai_key", "boolean", "Whether OPENAI_API_KEY is configured"],
            ["agent_stats", "object", "AI agent statistics summary"],
            ["latest_brief", "object", "Most recent threat brief"],
        ],
    )

    add_subsection_header(doc, "12.2", "Keyword Management")

    add_endpoint_section(doc, "GET", "/api/admin/keywords",
        "Retrieves the current keyword filter configuration used for message priority classification.",
        response_fields=[
            ["critical", "array[string]", "Keywords that trigger CRITICAL classification"],
            ["medium", "array[string]", "Keywords that trigger MEDIUM classification"],
        ],
    )

    add_endpoint_section(doc, "POST", "/api/admin/keywords",
        "Updates the keyword filter lists. Merges provided keywords with existing lists and syncs "
        "changes to keywords.json on disk. The monitor process picks up keyword changes within 60 seconds.",
        body=[
            ["critical", "array[string]", "No", "CRITICAL keywords to add/update"],
            ["medium", "array[string]", "No", "MEDIUM keywords to add/update"],
        ],
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["critical", "integer", "Total CRITICAL keywords after update"],
            ["medium", "integer", "Total MEDIUM keywords after update"],
            ["note", "string", "Status message"],
        ],
    )

    add_subsection_header(doc, "12.3", "Channel Management")

    add_endpoint_section(doc, "GET", "/api/admin/channels",
        "Lists all channel configurations including monitoring metadata.",
        response_fields=[
            ["username", "string", "Channel username"],
            ["tier", "integer", "Threat tier"],
            ["label", "string", "Display label"],
            ["threat", "string", "Threat level"],
            ["status", "string", "Channel status"],
        ],
        response="Array of channel configuration objects"
    )

    add_endpoint_section(doc, "POST", "/api/admin/channels",
        "Adds a new channel or updates an existing channel's configuration.",
        body=[
            ["username", "string", "Yes", "Telegram channel username"],
            ["tier", "integer", "No", "Threat tier ranking (1=highest)"],
            ["label", "string", "No", "Display label"],
            ["threat", "string", "No", "Threat level classification"],
            ["status", "string", "No", "Channel status (active, banned, dormant)"],
        ],
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["channel", "string", "Channel username"],
            ["config", "object", "Full channel configuration after update"],
        ],
    )

    add_endpoint_section(doc, "DELETE", "/api/admin/channels/<username>",
        "Removes a channel from the monitoring configuration. Does not delete existing messages.",
        params=[
            ["username", "string", "Yes (URL)", "—", "Channel username to remove"],
        ],
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["removed", "string", "Username of removed channel"],
        ],
    )

    add_subsection_header(doc, "12.4", "Backfill & Maintenance")

    add_endpoint_section(doc, "POST", "/api/admin/backfill",
        "Queues a historical message backfill job for a specific channel. The backfill is picked up "
        "by the live monitor process within 60 seconds and fetches historical messages from Telegram.",
        body=[
            ["channel", "string", "Yes", "Channel username to backfill"],
            ["limit", "integer", "No", "Maximum messages to fetch (default: 500, max: 2000)"],
            ["since", "string", "No", "Only fetch messages after this timestamp"],
        ],
        response_fields=[
            ["ok", "boolean", "Whether the backfill was queued"],
            ["queued", "object", "Backfill job details"],
            ["note", "string", "Status message"],
        ],
    )

    add_endpoint_section(doc, "POST", "/api/admin/compact",
        "Runs database maintenance: deduplicates messages by (channel, message_id), removes exact "
        "duplicates, and performs SQLite VACUUM to reclaim disk space.",
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["unique", "integer", "Remaining unique messages after deduplication"],
            ["deleted", "integer", "Number of duplicate rows removed"],
            ["critical", "integer", "CRITICAL messages remaining"],
        ],
    )

    add_subsection_header(doc, "12.5", "Discovery Management")

    add_endpoint_section(doc, "GET", "/api/admin/discovered",
        "Lists discovered channels with pagination support for the admin review workflow.",
        params=[
            ["page", "integer", "No", "—", "Page number (activates pagination)"],
            ["per_page", "integer", "No", "50", "Results per page (max: 200)"],
        ],
        response="Array of discovered channel objects (with pagination wrapper when page is specified)",
    )

    add_endpoint_section(doc, "POST", "/api/admin/discovered/<action>/<username>",
        "Performs a review action on a discovered channel. Available actions: approve (adds to monitoring), "
        "dismiss (rejects permanently), ignore (skips for now).",
        params=[
            ["action", "string", "Yes (URL)", "—", "Review action: 'approve', 'dismiss', or 'ignore'"],
            ["username", "string", "Yes (URL)", "—", "Discovered channel username"],
        ],
        body=[
            ["tier", "integer", "No (approve only)", "Threat tier for approved channel"],
            ["threat", "string", "No (approve only)", "Threat level for approved channel"],
            ["label", "string", "No (approve only)", "Display label for approved channel"],
        ],
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["action", "string", "Action performed"],
            ["username", "string", "Channel username"],
        ],
    )

    add_subsection_header(doc, "12.6", "System Configuration")

    add_endpoint_section(doc, "POST", "/api/system/config",
        "Updates environment configuration by writing to the .env file. Only whitelisted keys are "
        "accepted for security. Changes take effect on next service restart.",
        body=[
            ["OPENAI_API_KEY", "string", "No", "OpenAI API key for AI agent and chat"],
            ["TG_API_ID", "string", "No", "Telegram API ID"],
            ["TG_API_HASH", "string", "No", "Telegram API hash"],
            ["TG_PHONE", "string", "No", "Telegram phone number for session"],
        ],
        response_fields=[
            ["ok", "boolean", "Operation success"],
            ["saved", "array[string]", "List of keys that were saved"],
            ["note", "string", "Status message"],
        ],
        notes=["Note: Only the listed keys are accepted. Arbitrary key-value pairs are rejected for security."]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 13 — MEDIA
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "13", "Media & File Serving APIs")

    add_styled_para(doc,
        "These endpoints serve media files (images, videos, documents) downloaded from Telegram messages. "
        "Media is selectively downloaded based on message priority: only CRITICAL and MEDIUM messages "
        "have their media automatically retrieved. Files are stored in the telegram_intel/media/ directory.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    add_endpoint_section(doc, "GET", "/api/media/<filepath>",
        "Serves a media file from the telegram_intel/media/ directory. Includes path traversal protection "
        "to prevent directory escape attacks. Returns appropriate Content-Type based on file extension.",
        params=[
            ["filepath", "string", "Yes (URL path)", "—", "Relative path within the media directory"],
        ],
        response="Binary file with appropriate Content-Type header",
        notes=["Note: Path traversal attempts (../) are blocked and return HTTP 403."]
    )

    add_endpoint_section(doc, "GET", "/api/media/lookup/<channel>/<message_id>",
        "Looks up all media files associated with a specific message. Returns file metadata including "
        "URL, type classification, filename, and size. Used by the frontend to render inline media "
        "and attachment cards.",
        params=[
            ["channel", "string", "Yes (URL)", "—", "Channel username"],
            ["message_id", "integer", "Yes (URL)", "—", "Telegram message ID"],
        ],
        response_fields=[
            ["files", "array", "Media file descriptors"],
            ["files[].url", "string", "Download URL path"],
            ["files[].type", "string", "File type: 'image', 'video', or 'file'"],
            ["files[].name", "string", "Original filename"],
            ["files[].size", "integer", "File size in bytes"],
        ],
        notes=[
            "Note: Type classification is based on file extension:",
            "  - image: jpg, jpeg, png, gif, webp, bmp, svg, ico",
            "  - video: mp4, webm, mov, avi, mkv, flv, wmv, m4v",
            "  - file: all other extensions (pdf, zip, rar, xlsx, docx, txt, csv, etc.)",
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 14 — ERROR HANDLING
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "14", "Error Handling & Status Codes")

    add_styled_para(doc,
        "The platform uses standard HTTP status codes with JSON error responses. All error responses "
        "include an 'error' field with a human-readable description of the problem.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    make_endpoint_table(doc, [
        ["Status Code", "Meaning", "Common Causes"],
        ["200 OK", "Request succeeded", "Standard success response for all endpoints"],
        ["400 Bad Request", "Invalid parameters", "Missing required fields, invalid date format, out-of-range values"],
        ["403 Forbidden", "Access denied", "Path traversal attempt on media endpoint"],
        ["404 Not Found", "Resource not found", "Invalid channel username, non-existent APT group"],
        ["500 Internal Server Error", "Server error", "Database connection failure, external API timeout, unhandled exception"],
    ])

    add_styled_para(doc, "Error Response Format", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=12, space_after=6)
    add_code_block(doc, '// HTTP 400\n{\n  "error": "Missing required field: apt_name"\n}\n\n// HTTP 500\n{\n  "error": "Database connection failed",\n  "detail": "sqlite3.OperationalError: database is locked"\n}')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  SECTION 15 — PERFORMANCE
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "15", "Rate Limiting & Performance Notes")

    add_styled_para(doc, "In-Memory Cache", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=4, space_after=6)
    add_styled_para(doc,
        "The platform loads all messages from SQLite into an in-memory cache at startup. This enables "
        "sub-millisecond filtering and aggregation for most endpoints. The cache is updated incrementally "
        "as new messages arrive. For databases with 100K+ messages, initial startup may take 30-60 seconds.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=8)

    add_styled_para(doc, "Rate Limiting", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=4, space_after=6)
    add_styled_para(doc,
        "No rate limiting is applied by default. For production deployments, implement rate limiting "
        "at the reverse proxy level. Recommended limits:",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=6)

    make_endpoint_table(doc, [
        ["Endpoint Category", "Recommended Rate Limit", "Rationale"],
        ["Polling endpoints (/poll, /count)", "1 req/10s per client", "High-frequency calls by design"],
        ["Dashboard & analytics", "1 req/5s", "Moderate computation cost"],
        ["APT research & IOC lookup", "1 req/2s", "External API calls (AbuseIPDB)"],
        ["AI endpoints (/chat, /translate)", "1 req/5s", "OpenAI API cost and latency"],
        ["Export & report generation", "1 req/30s", "Heavy I/O and computation"],
        ["Admin endpoints", "1 req/s", "State-modifying operations"],
    ])

    add_styled_para(doc, "Performance Characteristics", "Normal", size=12, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E), space_before=10, space_after=6)

    make_endpoint_table(doc, [
        ["Endpoint", "Typical Latency", "Notes"],
        ["/api/messages/count", "<5ms", "Direct SQL COUNT query"],
        ["/api/messages/poll", "<10ms", "SQL WHERE with index"],
        ["/api/channels", "<20ms", "In-memory aggregation"],
        ["/api/messages/all", "20-100ms", "In-memory filter + sort (depends on result size)"],
        ["/api/dashboard", "50-200ms", "Full dataset aggregation with keyword scoring"],
        ["/api/apt/network", "200-500ms", "Cross-APT graph computation"],
        ["/api/apt/ioc/lookup", "500-3000ms", "Includes AbuseIPDB API call"],
        ["/api/chat/stream", "2-30s", "GPT-4o streaming with tool calls"],
        ["/api/blocklist/report", "3-10s", "DOCX generation + optional PDF conversion"],
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  APPENDIX A — ENDPOINT INDEX
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "A", "Appendix: Complete Endpoint Index")

    add_styled_para(doc,
        "Quick-reference table listing all API endpoints sorted by category and HTTP method.",
        "Normal", size=10, color=RGBColor(0x33, 0x33, 0x33), space_after=10)

    all_endpoints = [
        ["GET", "/", "Serve frontend UI"],
        ["GET", "/api/channels", "List monitored channels with statistics"],
        ["GET", "/api/channel/<username>/iocs", "Per-channel IOC aggregation"],
        ["GET", "/api/channel/<username>/trend", "30-day channel activity trend"],
        ["GET", "/api/messages/<username>", "Channel-specific messages with filtering"],
        ["GET", "/api/messages/all", "Unified timeline with advanced filtering"],
        ["GET", "/api/messages/poll", "Real-time polling for new messages"],
        ["GET", "/api/messages/count", "Lightweight message count"],
        ["GET", "/api/messages/<ch>/<id>/context", "Surrounding message context"],
        ["GET", "/api/messages/export", "Export messages as CSV"],
        ["GET", "/api/iocs/export", "Export all IOCs as CSV"],
        ["GET", "/api/dashboard", "Dashboard aggregation"],
        ["GET", "/api/trend", "Multi-day volume trend"],
        ["GET", "/api/briefing", "24-hour intelligence briefing"],
        ["GET", "/api/stats/summary", "Real-time status bar stats"],
        ["GET", "/api/threat_matrix", "Actor vs. sector matrix"],
        ["GET", "/api/apt/profiles", "APT group profiles"],
        ["GET", "/api/apt/network", "APT collaboration network graph"],
        ["GET", "/api/apt/<name>/detail", "APT group detailed profile"],
        ["GET", "/api/apt/<name>/research", "APT IOC research results"],
        ["POST", "/api/apt/ioc/lookup", "Multi-source IOC lookup"],
        ["POST", "/api/apt/ioc/ai-extract", "AI-powered IOC extraction"],
        ["GET", "/api/blocklist", "Central IOC blocklist"],
        ["GET", "/api/blocklist/export", "Export blocklist as CSV"],
        ["GET", "/api/blocklist/report", "Generate SOC advisory report"],
        ["GET", "/api/discovery/list", "List discovered channels"],
        ["POST", "/api/discovery/scan", "Trigger discovery scan"],
        ["GET", "/api/discovery/status", "Discovery scan status"],
        ["GET", "/api/discovery/fetch", "On-demand channel fetch"],
        ["GET", "/api/dark/feed", "Dark web intel feed"],
        ["GET", "/api/dark/stats", "Dark collector statistics"],
        ["GET", "/api/dark/digest", "AI daily digest"],
        ["GET", "/api/dark/domains", "Domain squatting results"],
        ["GET", "/api/dark/threat_level", "Behavioral threat level"],
        ["GET", "/api/ai/status", "AI agent status"],
        ["GET", "/api/ai/suggestions", "AI keyword suggestions"],
        ["GET", "/api/ai/brief", "AI threat brief"],
        ["GET", "/api/ai/enriched", "AI-enriched alerts"],
        ["GET", "/api/ai/log", "AI agent log tail"],
        ["POST", "/api/ai/analyze", "Start AI agent"],
        ["POST", "/api/ai/stop", "Stop AI agent"],
        ["POST", "/api/ai/apply", "Apply AI keyword suggestions"],
        ["POST", "/api/translate", "Arabic/Farsi to English translation"],
        ["GET", "/api/escalation/status", "Escalation detection status"],
        ["GET", "/api/hunting/leads", "Threat hunting leads"],
        ["GET", "/api/network/graph", "Channel relationship graph"],
        ["POST", "/api/chat", "Synchronous chat query"],
        ["GET", "/api/chat/history", "Chat session history"],
        ["POST", "/api/chat/reset", "Reset chat session"],
        ["POST", "/api/chat/stream", "Streaming chat (SSE)"],
        ["GET", "/api/admin/status", "System health status"],
        ["GET", "/api/admin/keywords", "Get keyword filters"],
        ["POST", "/api/admin/keywords", "Update keyword filters"],
        ["GET", "/api/admin/channels", "List channel configs"],
        ["POST", "/api/admin/channels", "Add/update channel"],
        ["DELETE", "/api/admin/channels/<username>", "Remove channel"],
        ["POST", "/api/admin/backfill", "Queue message backfill"],
        ["POST", "/api/admin/compact", "Database deduplication"],
        ["GET", "/api/admin/discovered", "List discovered channels (admin)"],
        ["POST", "/api/admin/discovered/<action>/<username>", "Review discovered channel"],
        ["GET", "/api/system/status", "Orchestrator status"],
        ["POST", "/api/system/config", "Update environment config"],
        ["GET", "/api/media/<filepath>", "Serve media file"],
        ["GET", "/api/media/lookup/<ch>/<id>", "Lookup message media"],
    ]

    make_endpoint_table(doc, [["Method", "Endpoint", "Description"]] + all_endpoints)

    doc.add_page_break()

    # ── DISCLAIMER ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("DISCLAIMER")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ACCENT_RED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        "This document is CONFIDENTIAL and intended solely for authorized personnel of ScanWave Cybersecurity "
        "and its designated clients. The information contained herein describes a cyber threat intelligence "
        "platform that operates exclusively through open-source intelligence (OSINT) collection methods. "
        "No unauthorized access to any systems, networks, or accounts has been performed or is implied. "
        "All Telegram channel monitoring is conducted through publicly accessible channels using standard "
        "Telegram API access.\n\n"
        "The API endpoints described in this document provide access to intelligence data that may contain "
        "sensitive threat indicators, IOCs, and attack attribution. Access to these endpoints should be "
        "restricted to authorized security analysts and system administrators. ScanWave Cybersecurity "
        "makes no warranties regarding the completeness or accuracy of threat intelligence data and "
        "assumes no liability for actions taken based on the information provided through these APIs.\n\n"
        "Unauthorized reproduction, distribution, or disclosure of this document or its contents is "
        "strictly prohibited."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Calibri"

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run(f"\u00A9 {datetime.now().year} ScanWave Cybersecurity. All rights reserved.")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x77, 0xCC)
    run.font.name = "Calibri"

    # ── SAVE ────────────────────────────────────────────────────────────
    base = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(base, "ScanWave_CyberIntel_API_Reference.docx")
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")

    # Try PDF conversion
    pdf_path = os.path.join(base, "ScanWave_CyberIntel_API_Reference.pdf")

    # Try LibreOffice
    lo_cmd = None
    for lo in ["/usr/bin/libreoffice", "/usr/bin/soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice", "libreoffice", "soffice"]:
        if os.path.isfile(lo) or os.system(f"which {lo} > /dev/null 2>&1") == 0:
            lo_cmd = lo
            break

    if lo_cmd:
        import subprocess
        result = subprocess.run([
            lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", base, docx_path
        ], capture_output=True, text=True, timeout=120)
        if os.path.isfile(pdf_path):
            print(f"PDF saved: {pdf_path}")
        else:
            print(f"PDF conversion attempted but file not found. LibreOffice output: {result.stderr}")
    else:
        # Try docx2pdf
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            print(f"PDF saved: {pdf_path}")
        except Exception as e:
            print(f"PDF conversion not available locally ({e}). DOCX created successfully — convert manually or deploy to OCI for PDF generation.")


if __name__ == "__main__":
    generate()
