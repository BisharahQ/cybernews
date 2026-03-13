"""
Build combined BAE IOC + SIEM Investigation Report as DOCX
Styling: ScanWave SOC Client Advisory (white cover, dark bar headings, Calibri, green accent)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE = r"C:\Users\BisharahQraitem\Downloads\cybernews"
LOGO_PATH = os.path.join(BASE, "logo_xref9.jpeg")

# ── ScanWave Advisory PDF color palette ──
DARK_HEADING = RGBColor(0x19, 0x21, 0x28)   # #192128
BODY_TEXT = RGBColor(0x1f, 0x2a, 0x37)       # #1f2a37
SUBTITLE = RGBColor(0x4a, 0x55, 0x68)        # #4a5568
RED_CONF = RGBColor(0xb7, 0x1c, 0x1c)        # #b71c1c
GREEN_ACC = "5ed292"                           # green accent line
WHITE = RGBColor(0xff, 0xff, 0xff)
SECTION_BAR = "1a2332"                         # dark section heading bar
TBL_HEADER = "192128"                          # dark table header
ALT_ROW = "f8fafb"                             # alternating row
FONT = "Calibri"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10)
style.font.color.rgb = BODY_TEXT
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

# ── Helpers ──

def shd(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def borders(tbl, style_str):
    el = tbl._tbl
    pr = el.tblPr if el.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    pr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}>{style_str}</w:tblBorders>'))

NONE_BORDERS = (
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>')

LIGHT_BORDERS = (
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="e5e7eb"/>')

def R(p, text, size=10, bold=False, color=BODY_TEXT, font=FONT):
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    return r

def section_heading(text):
    """Dark bar with white text — matches advisory PDF"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); shd(c, SECTION_BAR)
    t.columns[0].width = Inches(6.5)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    R(p, f"  {text}", size=11, bold=True, color=WHITE)
    borders(t, NONE_BORDERS)
    doc.add_paragraph()

def subsection(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    R(p, text, size=10, bold=True, color=DARK_HEADING)

def body(text):
    p = doc.add_paragraph(); R(p, text, size=10, color=BODY_TEXT); return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.cell(0, i); shd(c, TBL_HEADER)
        R(c.paragraphs[0], h, size=8, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.cell(ri+1, ci)
            if ri % 2 == 1: shd(c, ALT_ROW)
            R(c.paragraphs[0], str(v), size=8, color=BODY_TEXT)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    borders(t, LIGHT_BORDERS)
    doc.add_paragraph()
    return t

def callout(label, text, bcolor="e65c00", bg="fff3e0", lcolor=None):
    if lcolor is None: lcolor = RGBColor(0xe6, 0x5c, 0x00)
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0); shd(c, bg); t.columns[0].width = Inches(6.5)
    p = c.paragraphs[0]
    R(p, label, size=10, bold=True, color=lcolor)
    R(p, text, size=10, color=BODY_TEXT)
    borders(t,
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{bcolor}"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{bcolor}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{bcolor}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{bcolor}"/>')
    doc.add_paragraph()

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); R(p, text, size=10, color=BODY_TEXT)

def code_block(title, lines):
    subsection(title)
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        R(p, ln, size=8, color=BODY_TEXT, font="Consolas")


# ═══════════════════════════════════════════════════════════════
# COVER PAGE  (white bg, centered, matching ScanWave Advisory)
# ═══════════════════════════════════════════════════════════════

for _ in range(4): doc.add_paragraph()

if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO_PATH, width=Inches(2.0))

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "ScanWave", size=36, bold=True, color=DARK_HEADING)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "IOC INTELLIGENCE REPORT", size=22, bold=True, color=DARK_HEADING)

# Green accent line
t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
shd(t.cell(0,0), GREEN_ACC); t.columns[0].width = Inches(4)
R(t.cell(0,0).paragraphs[0], " ", size=1); borders(t, NONE_BORDERS)
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "Bank al Etihad (\u0628\u0646\u0643 \u0627\u0644\u0627\u062a\u062d\u0627\u062f) Compromise", size=12, color=SUBTITLE)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "Aqaba Special Economic Zone Solar Infrastructure Attack", size=12, color=SUBTITLE)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "March 2026", size=12, color=RGBColor(0x2d, 0x3a, 0x4a))

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "CONFIDENTIAL", size=11, bold=True, color=RED_CONF)

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "Prepared by ScanWave Security Operations Center", size=10, color=SUBTITLE)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(p, "platform.scanwave.io  \u00b7  soc@scanwave.io", size=8, color=SUBTITLE)

doc.add_page_break()

# ── Header & Footer (advisory style) ──
sec = doc.sections[0]
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
R(hp, "  ScanWave", size=8, bold=True, color=DARK_HEADING)
R(hp, "  |  IOC Intelligence Report  |  March 2026", size=7, color=SUBTITLE)
R(hp, "\t\t\tCONFIDENTIAL", size=6, bold=True, color=RED_CONF)

fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
R(fp, "\u00a9 2026 ScanWave Cybersecurity  |  platform.scanwave.io  |  soc@scanwave.io", size=6, color=SUBTITLE)


# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

section_heading("1. EXECUTIVE SUMMARY")

body(
    "A threat actor with assessed Iranian-nexus affiliations has claimed successful compromise of "
    "Bank al Etihad (\u0628\u0646\u0643 \u0627\u0644\u0627\u062a\u062d\u0627\u062f), one of Jordan\u2019s leading commercial banks. "
    "The threat communique, published in Farsi, details a multi-stage intrusion that leveraged a compromised "
    "third-party management system deployed within Jordan to gain persistent backdoor access to the bank\u2019s "
    "core server infrastructure.")

body(
    "The actor claims persistent access to the bank\u2019s technical infrastructure, asserting that defensive "
    "measures undertaken by the bank\u2019s security team \u2014 including port blocking and access restriction \u2014 "
    "were ineffective. Screenshots provided as proof-of-compromise include live server monitoring dashboards "
    "and a BELECTRIC solar energy SCADA/ICS panel showing grid-disconnected inverters in the Aqaba Special "
    "Economic Zone.")

body(
    "The initial access vector is attributed to a third-party contractor responsible for deploying a management "
    "system in Jordan, which inadvertently introduced multiple backdoor access points to Bank al Etihad\u2019s "
    "core servers. This represents a classic supply chain compromise with significant implications for both "
    "the banking and energy sectors.")

body(
    "Threat intelligence correlation links this activity to the Handala Hack Team (MOIS-linked), an Iranian "
    "state-sponsored group conducting Operation HamsaUpdate \u2014 a destructive wiper campaign targeting energy "
    "and financial sectors across Jordan, Israel, and the Gulf region. BELECTRIC Gulf Ltd was previously "
    "breached by Handala (listed as \"BLEnergy\" on 2024-07-23), establishing a pattern of repeated targeting.")

callout(
    "INVESTIGATION FINDING: ",
    "A comprehensive investigation was conducted across the Bank al Etihad environment. "
    "No evidence of compromise via the claimed SCADA/solar attack vector was found. However, active external "
    "reconnaissance against BAE web assets, SIP brute-force attacks against branch VoIP infrastructure, and a "
    "complete telemetry blind spot on the alleged attack surface (solar/SCADA systems) were identified.",
    bcolor="b71c1c", bg="fce4ec", lcolor=RED_CONF)


# ═══════════════════════════════════════════════════════════════
# 2. THREAT ACTOR PROFILE
# ═══════════════════════════════════════════════════════════════

section_heading("2. THREAT ACTOR PROFILE")

table(["Attribute", "Detail"], [
    ["Designation", "Handala Hack Team (a.k.a. Handala, Hanzala)"],
    ["Tracked As", "UNK-IR-ETIHAD-0307 / Handala Hack Team"],
    ["State Affiliation", "Iran \u2014 Ministry of Intelligence and Security (MOIS)"],
    ["Assessed Origin", "Iran (Farsi-language communique)"],
    ["Motivation", "Geopolitical / Hacktivism / Destructive (wiper deployment)"],
    ["Sophistication", "Moderate-to-High (supply chain exploitation, wiper malware, SCADA access)"],
    ["Target Profile", "Jordanian financial institutions, energy/solar infrastructure, Israeli targets"],
    ["Active Operations", "Operation HamsaUpdate (wiper campaign), RedAlert fake app, BELECTRIC solar breach"],
    ["Language", "Farsi (Persian) \u2014 operational communique"],
    ["Leak Sites", "handala-hack.to (clearnet, DDoS-Guard) / .onion Tor site (NGINX)"],
    ["Total Known Victims", "137 (as of 2026-03-07) | Last victim posted: 2026-03-07"],
    ["Status", "ACTIVE \u2014 Claims ongoing access"],
], [2.0, 4.5])

body(
    "The threat actor\u2019s communique was authored entirely in Farsi, which combined with the targeting pattern "
    "(Jordanian financial infrastructure) and the parallel SCADA/ICS compromise is consistent with Iranian-nexus "
    "cyber operations. The targeting of solar energy infrastructure in the Aqaba Special Economic Zone aligns "
    "with known Iranian interest in disrupting Gulf-state and Jordanian critical infrastructure.")

subsection("2.1 Handala Recent Victims (March 2026 Campaign)")
body("The following victims were posted by Handala in the same campaign wave as the BAE/BELECTRIC attack:")

table(["Victim", "Date", "Sector", "Relevance"], [
    ["BELECTRIC Gulf Ltd (as \"BLEnergy\")", "2024-07-23", "Energy / Solar", "PRIOR breach of same target"],
    ["Israel Opportunity Energy", "2026-03-02", "Energy", "Same campaign wave"],
    ["Aramco (Saudi)", "2026-03-03", "Energy / Oil & Gas", "Same campaign wave"],
    ["Sharjah National Oil Corporation", "2026-03-03", "Energy / Oil & Gas", "Same campaign wave"],
    ["INSS (Israel National Security Studies)", "2026-03-04", "Government / Think Tank", "Same campaign wave"],
    ["Atlas Insurances Ltd", "2026-03-05", "Financial / Insurance", "Same campaign wave"],
    ["Clalit Healthcare", "2026-02-25", "Healthcare", "Same campaign wave"],
], [2.5, 1.0, 1.5, 1.5])

subsection("2.2 Handala Online Presence")
body("Known Telegram channels, forums, and communication platforms:")

table(["Platform", "Identifier", "Status"], [
    ["Telegram (Primary)", "@HANDALA_HPR2", "Active \u2014 latest known primary channel"],
    ["Telegram", "@handala_hack26", "Active (1553 members)"],
    ["Telegram", "@handala24", "Active news platform"],
    ["Telegram", "@Handalahackteam", "Active (939 members)"],
    ["Telegram", "@Handala_channal", "Active (689 members)"],
    ["Telegram", "@Handala_Hack_Team", "Active (544 members)"],
    ["Telegram", "@Handala_hack_iranian", "Active (242 members)"],
    ["Twitter/X", "@Handala_news", "Active"],
    ["Forum", "BreachForums", "Active account"],
    ["Forum", "Ramp / Exploit", "Active accounts"],
    ["Comms", "Tox / ProtonMail", "Secure communication channels"],
], [1.2, 2.5, 2.8])


# ═══════════════════════════════════════════════════════════════
# 3. ATTACK NARRATIVE & TIMELINE
# ═══════════════════════════════════════════════════════════════

section_heading("3. ATTACK NARRATIVE & TIMELINE")

for bold_part, txt in [
    ("Phase 1 \u2014 Initial Access (Supply Chain Compromise): ",
     "The threat actor exploited a management system deployed by a third-party contractor operating within Jordan. "
     "This system contained inherent vulnerabilities that created multiple unintended backdoor access points to "
     "Bank al Etihad\u2019s core server infrastructure. Handala\u2019s known TTPs include spear-phishing with fake F5 "
     "BIG-IP update notifications and RedAlert app updates, delivering .NET loaders (F5UPDATER.exe)."),
    ("Phase 2 \u2014 Persistence & Lateral Movement: ",
     "Upon gaining initial access, the actor established persistent footholds across multiple segments of the "
     "bank\u2019s technical infrastructure. The communique explicitly states that the bank\u2019s security team\u2019s "
     "efforts to block ports and restrict access were unsuccessful in dislodging the adversary. Handala\u2019s "
     "Operation HamsaUpdate deploys second-stage loaders (Handala.exe) and uses AutoIt-based shellcode injection."),
    ("Phase 3 \u2014 Active Surveillance & Exfiltration: ",
     "The actor conducted live monitoring of the bank\u2019s primary servers, capturing real-time screenshots of "
     "server dashboards as evidence of continued access. Handala\u2019s known exfiltration methods include "
     "Telegram bot C2 (token: 6428401585:AAG...), AWS S3, and Storj decentralized storage."),
    ("Phase 4 \u2014 Parallel Infrastructure Attack (SCADA/ICS): ",
     "Evidence from the provided screenshots indicates compromise of a BELECTRIC solar energy monitoring system "
     "in the Aqaba Special Economic Zone. All five solar inverters (INV/1 through INV/5) were observed in an OFF "
     "state with \u201cGrid Loss\u201d status, suggesting deliberate grid disconnection or manipulation of the solar "
     "energy control systems. BELECTRIC Gulf Ltd operates 7 solar plants in Jordan (~100 MW total)."),
    ("Phase 5 \u2014 Wiper Deployment (Assessed): ",
     "Based on Handala\u2019s known TTPs from Operation HamsaUpdate, destructive wiper payloads (Hatef.exe for Windows, "
     "update.sh/Hamsa for Linux) may have been staged or deployed. These wipers perform disk structure wipe "
     "(T1561.002) and use BYOVD privilege escalation via vulnerable drivers (ListOpenedFileDrv_32.sys)."),
]:
    p = doc.add_paragraph()
    R(p, bold_part, size=10, bold=True, color=BODY_TEXT)
    R(p, txt, size=10, color=BODY_TEXT)

body(
    "Observed Timestamp: 2024-01-28 08:47:09 UTC (from BELECTRIC dashboard). Environmental data anomalies "
    "(ambient temperature reading of \u221240\u00b0C) suggest possible sensor manipulation or system compromise "
    "affecting telemetry integrity.")


# ═══════════════════════════════════════════════════════════════
# 4. RISK ASSESSMENT MATRIX
# ═══════════════════════════════════════════════════════════════

section_heading("4. RISK ASSESSMENT MATRIX")

table(["Category", "Rating", "Justification"], [
    ["Confidentiality Impact", "HIGH", "Potential access to banking core systems, customer data exposure"],
    ["Integrity Impact", "CRITICAL", "Wiper malware (Hatef.exe) designed to destroy disk structures; server manipulation"],
    ["Availability Impact", "CRITICAL", "Solar grid disconnection confirmed; wiper deployment destroys systems permanently"],
    ["Overall Risk Rating", "CRITICAL", "Active nation-state/hacktivist intrusion with destructive capability into financial infrastructure"],
], [1.8, 0.8, 3.9])


# ═══════════════════════════════════════════════════════════════
# 5. INDICATORS OF COMPROMISE (IOCs)
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
section_heading("5. INDICATORS OF COMPROMISE (IOCs)")

body(
    "The following IOCs have been extracted from the threat actor\u2019s communique, attached screenshots, "
    "associated infrastructure analysis, and Handala Hack Team threat intelligence. All network indicators "
    "are presented in defanged format.")

subsection("5.1 Network Indicators")
table(["Type", "Indicator", "Context", "Confidence"], [
    ["IP (C2)", "31.192.237[.]207", "Handala C2 server (port 2515/HTTPS) \u2014 Chelyabinsk, Russia", "CRITICAL"],
    ["IP", "38.180.239[.]161", "Associated threat infrastructure", "HIGH"],
    ["IP", "209.74.87[.]100", "Associated threat infrastructure", "HIGH"],
    ["IP", "157.20.182[.]49", "Associated threat infrastructure", "HIGH"],
    ["IP", "185.236.25[.]11", "Associated threat infrastructure", "HIGH"],
    ["IP", "92.243.65[.]243", "Associated threat infrastructure", "HIGH"],
    ["IP Range", "185.173.35[.]0/24", "Suspected C2 infrastructure", "HIGH"],
    ["IP Range", "91.243.44[.]0/24", "SCADA/ICS callback range", "MEDIUM"],
    ["Port", "TCP/2515", "Handala C2 communication port (HTTPS)", "CRITICAL"],
    ["Port", "TCP/4443", "Non-standard HTTPS backdoor", "HIGH"],
    ["Port", "TCP/8443", "Management interface backdoor", "HIGH"],
    ["Port", "TCP/9090", "Web management console", "MEDIUM"],
], [1.0, 1.8, 2.2, 1.0])

subsection("5.2 Infrastructure & System Indicators")
table(["Type", "Indicator", "Context", "Confidence"], [
    ["Platform", "BELECTRIC Solar Monitoring", "Compromised SCADA/ICS panel", "CONFIRMED"],
    ["System", "Solar Inverters INV/1\u2013INV/5", "Grid-disconnected (OFF state)", "CONFIRMED"],
    ["Timestamp", "2024-01-28T08:47:09Z", "Observed compromise timestamp", "CONFIRMED"],
    ["Location", "Aqaba Special Economic Zone", "Solar farm geographic target", "CONFIRMED"],
    ["System", "Banking Core Server Dashboard", "Live server monitoring claimed", "HIGH"],
    ["Vector", "Third-party management system", "Initial access vector (contractor)", "HIGH"],
    ["Prior Breach", "BELECTRIC Gulf Ltd (\"BLEnergy\")", "Handala victim 2024-07-23", "CONFIRMED"],
], [1.0, 1.8, 2.2, 1.0])

subsection("5.3 Behavioral Indicators (TTPs)")
table(["Type", "Indicator", "Context", "Confidence"], [
    ["TTP", "Contractor supply chain compromise", "Third-party management platform", "HIGH"],
    ["TTP", "Persistent backdoor implantation", "Multiple access points on core servers", "HIGH"],
    ["TTP", "Live server screenshot capture", "Real-time surveillance of target", "CONFIRMED"],
    ["TTP", "Firewall/port restriction bypass", "Evading defensive countermeasures", "HIGH"],
    ["TTP", "SCADA/ICS manipulation", "Solar inverter grid disconnect", "CONFIRMED"],
    ["TTP", "Farsi-language threat communique", "Iranian-nexus attribution indicator", "MEDIUM"],
    ["TTP", "BYOVD privilege escalation", "ListOpenedFileDrv_32.sys vulnerable driver", "HIGH"],
    ["TTP", "Telegram bot C2 exfiltration", "Bot token + chat ID for data theft", "HIGH"],
], [1.0, 2.0, 2.0, 1.0])

subsection("5.4 Malicious URLs & Payloads")
table(["Type", "Indicator", "Context", "Confidence"], [
    ["URL", "hxxps://www[.]shirideitch[.]com/wp-content/uploads/2022/06/RedAlert.apk", "Malicious Android APK (RedAlert trojan)", "HIGH"],
    ["URL", "hxxps://api[.]ra-backup[.]com/analytics/submit.php", "C2 callback / data exfiltration endpoint", "HIGH"],
    ["URL", "hxxps://bit[.]ly/4tWJhQh", "Shortened URL \u2014 suspected phishing/malware delivery", "HIGH"],
    ["URL", "hxxps://sjc1[.]vultrobjects[.]com/f5update/update.sh", "Linux wiper (Hamsa) download URL", "CRITICAL"],
    ["URL", "hxxp://38[.]180[.]239[.]161", "Threat infrastructure \u2014 HTTP endpoint", "HIGH"],
    ["URL", "hxxps://web27[.]info/0d894cce589ccd44", "Suspected malware staging / phishing", "HIGH"],
    ["URL", "hxxps://web27[.]info/b4cac8a7128d6507", "Suspected malware staging / phishing", "HIGH"],
    ["URL", "hxxps://web27[.]info/772b49c3c513b961", "Suspected malware staging / phishing", "HIGH"],
    ["URL", "hxxps://web27[.]info/b4f44ec4db01887aHOR", "Suspected malware staging / phishing", "HIGH"],
    ["URL", "hxxps://web27[.]info/ec714d155863dcd6", "Suspected malware staging / phishing", "HIGH"],
], [0.6, 3.2, 1.7, 0.8])

subsection("5.5 Malicious Domains")
table(["Domain", "Context", "Action", "Confidence"], [
    ["api[.]ra-backup[.]com", "C2 callback domain \u2014 RedAlert fake app analytics endpoint", "BLOCK", "CRITICAL"],
    ["shirideitch[.]com", "Phishing delivery domain \u2014 hosts malicious RedAlert APK", "BLOCK", "HIGH"],
    ["sjc1[.]vultrobjects[.]com", "Vultr object storage \u2014 hosts Linux wiper payload", "BLOCK", "CRITICAL"],
    ["web27[.]info", "Suspected malware staging / phishing infrastructure", "BLOCK", "HIGH"],
    ["handala[.]to", "Handala clearnet leak site (unavailable since 2025-06-26)", "BLOCK", "HIGH"],
    ["handala-hack[.]to", "Handala clearnet leak site (active, DDoS-Guard hosted)", "BLOCK", "HIGH"],
    ["belectric[.]com", "BELECTRIC main domain \u2014 monitor for credential abuse/phishing", "MONITOR", "HIGH"],
], [1.8, 2.8, 0.8, 0.8])

subsection("5.6 Handala C2 Infrastructure")
table(["Type", "Indicator", "Context", "Severity"], [
    ["C2 Server", "31.192.237[.]207:2515/HTTPS", "Primary wiper C2 \u2014 Chelyabinsk, Russia", "CRITICAL"],
    ["Telegram Bot Token", "6428401585:AAGE6SbwtVJxOpLjdMcrL45gb18H9UV7tQA", "C2 channel for wiper exfiltration", "CRITICAL"],
    ["Telegram Chat ID", "6932028002", "Receiving exfiltrated data from wiper", "CRITICAL"],
    ["Tor Leak Site", "vmjfieomxhnfjba57sd6jjws2ogvowjgxhhfglsikqvvrnrajbmpxqqd.onion", "Handala Tor leak site (NGINX)", "HIGH"],
    ["Exfil Service", "AWS S3 / Storj", "Decentralized storage for data exfiltration staging", "HIGH"],
    ["Exfil Service", "api.telegram.org", "Telegram API used as C2 channel", "HIGH"],
], [1.2, 2.8, 1.8, 0.8])

doc.add_page_break()
subsection("5.7 Malware Hashes (SHA256)")
body("Windows wiper campaign, Operation HamsaUpdate, and RedAlert fake app malware samples:")

table(["SHA256", "Filename / Description", "Type", "Severity"], [
    ["96dec6e07229...c634f0", "update.zip \u2014 Phishing delivery archive", "Dropper", "CRITICAL"],
    ["19001dd441e5...c634f0", "Phishing attachment PDF \u2014 initial lure", "Lure", "CRITICAL"],
    ["8316065c4536...e4d67", "OpenFileFinder.dll \u2014 wiper component", "Wiper DLL", "CRITICAL"],
    ["9e519211947c...e4d67", "Wiper variant sample", "Wiper", "CRITICAL"],
    ["fe07dca68f28...e30bd2", "F5UPDATER.exe \u2014 .NET loader variant 1 (signed)", "Loader", "CRITICAL"],
    ["ca9bf13897af...e0d74a", "F5UPDATER.exe \u2014 .NET loader variant 2 (signed)", "Loader", "CRITICAL"],
    ["e28085e8d64b...3af35", "Hatef.exe \u2014 Windows wiper malware", "Wiper", "CRITICAL"],
    ["454e6d3782f2...b840ad", "Handala.exe \u2014 Delphi second-stage loader", "Loader", "CRITICAL"],
    ["6f79c0e0e1aa...b840ad", "update.sh \u2014 Linux wiper (Hamsa) encrypted", "Wiper", "CRITICAL"],
    ["ad66251d9e87...166e8a", "ZIP archive \u2014 loader + Hatef wiper", "Archive", "CRITICAL"],
    ["64c5fd791ee3...5428c", "ZIP archive \u2014 loader + Handala loader", "Archive", "CRITICAL"],
    ["336167b8c5cf...b2767", "FlashDevelop \u2014 third-stage payload", "Payload", "CRITICAL"],
    ["f58d3a4b2f3f...efbd3", "Naples.pif \u2014 renamed AutoIt interpreter", "Evasion", "HIGH"],
    ["aae989743ddd...8acd4", "Obfuscated AutoIt script \u2014 shellcode injector", "Injector", "CRITICAL"],
    ["83651b058966...578b72", "RedAlert.apk \u2014 malicious Android app", "Mobile Malware", "HIGH"],
    ["7d9fb236607e...e964c9", "Associated malware sample", "Malware", "HIGH"],
], [1.8, 2.5, 0.8, 0.8])

body("Full SHA256 hashes for EDR/SIEM integration:")

# Full hashes in a compact table
table(["#", "Full SHA256 Hash"], [
    ["1", "96dec6e07229201a02f538310815c695cf6147c548ff1c6a0def2fe38f3dcbc8"],
    ["2", "19001dd441e50233d7f0addb4fcd405a70ac3d5e310ff20b331d6f1a29c634f0"],
    ["3", "8316065c4536384611cbe7b6ba6a5f12f10db09949e66cb608c92ae8b69e4d67"],
    ["4", "9e519211947c63d9bf6f4a51bc161f5b9ace596c2935a8eedfce4057f747b961"],
    ["5", "fe07dca68f288a4f6d7cbd34d79bb70bc309635876298d4fde33c25277e30bd2"],
    ["6", "ca9bf13897af109cb354f2629c10803966eb757ee4b2e468abc04e7681d0d74a"],
    ["7", "e28085e8d64bb737721b1a1d494f177e571c47aab7c9507dba38253f6183af35"],
    ["8", "454e6d3782f23455875a5db64e1a8cd8eb743400d8c6dadb1cd8fd2ffc2f9567"],
    ["9", "6f79c0e0e1aab63c3aba0b781e0e46c95b5798b2d4f7b6ecac474b5c40b840ad"],
    ["10", "ad66251d9e8792cf4963b0c97f7ab44c8b68101e36b79abc501bee1807166e8a"],
    ["11", "64c5fd791ee369082273b685f724d5916bd4cad756750a5fe953c4005bb5428c"],
    ["12", "336167b8c5cfc5cd330502e7aa515cc133656e12cbedb4b41ebbf847347b2767"],
    ["13", "f58d3a4b2f3f7f10815c24586fae91964eeed830369e7e0701b43895b0cefbd3"],
    ["14", "aae989743dddc84adef90622c657e45e23386488fa79d7fe7cf0863043b8acd4"],
    ["15", "83651b0589665b112687f0858bfe2832ca317ba75e700c91ac34025ee6578b72"],
    ["16", "7d9fb236607e7fe5e0921b06f45d9cb69acbdb923e2877d87a99720b8dc964c9"],
], [0.3, 6.0])

subsection("5.8 Malicious Filenames (EDR/SIEM Hunting)")
body("The following filenames should be hunted across all endpoints and monitored in EDR solutions:")

table(["Filename", "Description", "Severity", "Action"], [
    ["F5UPDATER.exe", "Handala loader masquerading as F5 BIG-IP update", "CRITICAL", "HUNT"],
    ["Hatef.exe", "Windows wiper malware", "CRITICAL", "HUNT"],
    ["Handala.exe", "Delphi second-stage loader", "CRITICAL", "HUNT"],
    ["senvarservice-DC.exe", "Data exfiltration component (AWS S3/Storj/Telegram)", "CRITICAL", "HUNT"],
    ["Naples.pif", "Renamed AutoIt interpreter (.pif extension evasion)", "HIGH", "HUNT"],
    ["OpenFileFinder.dll", "Wiper DLL component", "CRITICAL", "HUNT"],
    ["ListOpenedFileDrv_32.sys", "Vulnerable driver used for BYOVD privilege escalation", "CRITICAL", "HUNT"],
    ["Carroll.cmd", "Obfuscated batch script for wiper execution", "HIGH", "HUNT"],
    ["RedAlert.apk", "Fake RedAlert app \u2014 Android malware", "HIGH", "HUNT"],
    ["update.sh", "Linux wiper payload script (Hamsa)", "CRITICAL", "HUNT"],
], [2.0, 2.5, 0.8, 0.6])


# ═══════════════════════════════════════════════════════════════
# 6. SIEM INVESTIGATION — CORRELATED FINDINGS
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
section_heading("6. INVESTIGATION \u2014 CORRELATED FINDINGS")

body(
    "A comprehensive search was conducted across the Bank al Etihad environment to correlate "
    "the threat actor\u2019s claimed IOCs against actual telemetry. The following sections present all confirmed "
    "findings.")

subsection("6.1 Threat Actor IOC Correlation (Negative)")
body("The following claimed indicators were searched and returned ZERO results:")

table(["Indicator", "Context", "Result"], [
    ["13.41.61.138", "Shodan report \u2014 BAE infrastructure", "NO HITS"],
    ["23.21.102.205", "Shodan report \u2014 BAE infrastructure", "NO HITS"],
    ["CVE-2024-3400", "Palo Alto PAN-OS RCE", "NO HITS"],
    ["CVE-2024-21887", "Ivanti Connect Secure", "NO HITS"],
    ["CVE-2023-46805", "Ivanti Connect Secure", "NO HITS"],
    ["BELECTRIC", "Solar management vendor", "NO HITS"],
    ["SCADA / ICS / OT keywords", "Industrial control system data", "NO HITS"],
    ["examproctor.bankaletihad.com", "Subdomain with critical vulnerability", "NO HITS"],
    ["185.173.35.0/24", "Suspected C2 infrastructure", "NO HITS"],
    ["91.243.44.0/24", "SCADA/ICS callback range", "NO HITS"],
    ["31.192.237.207", "Handala C2 server", "NO HITS"],
], [2.2, 2.5, 1.3])

subsection("6.2 WAF \u2014 External Attacker IPs Detected")
body(
    "While the threat actor\u2019s specific IOCs were not found, the following external attackers were detected "
    "probing Bank al Etihad web assets:")

table(["Source IP", "Country", "Attack Type", "Target"], [
    ["178.79.148.111", "United Kingdom", "db.sql file probe (database dump attempt)", "bankaletihad.com"],
    ["165.154.24.96", "Hong Kong", "Phishing reconnaissance probe", "bankaletihad.com"],
    ["194.238.28.128", "Unknown", "HackerOne-style security researcher recon", "bankaletihad.com"],
], [1.3, 1.0, 2.7, 1.3])

subsection("6.3 S3 Bucket Reconnaissance Probes")
body("Probes against BAE\u2019s S3 bucket (prod-bae-blog-frontend), all returned AccessDenied:")

table(["URI Probed", "Intent", "Result"], [
    ["/blog-cdn.7z", "Archive exfiltration attempt", "AccessDenied"],
    ["/blog-cdn/actuator/prometheus", "Spring Boot actuator endpoint discovery", "AccessDenied"],
    ["/blog-cdn-api/api/v1/users", "API user enumeration attempt", "AccessDenied"],
], [2.0, 2.5, 1.5])

subsection("6.4 Palo Alto NGFW \u2014 High Severity Threat Events")
body("Device: Primary-NGFW.bankaletihad.com (Serial: 013201006040) | 4 events detected:")

table(["Timestamp", "Source IP", "Dest IP", "Port", "Threat", "Action"], [
    ["2026-03-04 16:11:32", "10.0.6.69", "172.17.30.1", "5060/UDP", "SIP Register Brute Force (40023)", "DROP"],
    ["2026-03-04 16:11:36", "10.0.6.69", "172.17.30.1", "5060/UDP", "SIP Register Brute Force (40023)", "DROP"],
    ["2026-03-04 16:15:38", "10.0.6.71", "172.17.30.1", "5060/UDP", "SIP Register Brute Force (40023)", "DROP"],
    ["2026-03-04 16:15:41", "10.0.6.71", "172.17.30.1", "5060/UDP", "SIP Register Brute Force (40023)", "DROP"],
], [1.2, 0.7, 0.8, 0.6, 2.0, 0.5])

body(
    "Both source IPs (10.0.6.69, 10.0.6.71) are internal branch network devices targeting a VoIP/SIP "
    "gateway at 172.17.30.1. Threat category: brute-force. Application flags include: used-by-malware, "
    "has-known-vulnerability. All attacks were DROPPED by the NGFW.")

subsection("6.5 Oracle Database Audit Events")
body("Host: Batch-ADP-PROD.BankAlEtihad.Com (10.12.11.233):")

table(["Timestamp", "Session", "User", "Table Accessed", "OS User", "RC"], [
    ["2026-03-04 08:46:04Z", "2330981", "BATCH_ADAPTER", "XX_ACCOUNT_MAPPING", "BATCH-ADP-PROD\\Admin", "0"],
    ["2026-03-05 06:10:28Z", "2332571", "BATCH_ADAPTER", "XX_ACCOUNT_MAPPING", "BATCH-ADP-PROD\\Admin", "0"],
], [1.2, 0.7, 1.0, 1.3, 1.3, 0.3])

body(
    "These are Oracle DB audit trail entries (EventID 34) showing the BATCH_ADAPTER service account "
    "performing routine SELECT queries on XX_ACCOUNT_MAPPING. Running under Administrator context. "
    "Assessment: normal batch processing activity.")

subsection("6.6 AWS CloudTrail \u2014 Console Access")
body("AWS Account: 829210555596 (eu-west-2 / London)")

table(["Timestamp", "User", "Source IP", "Result"], [
    ["2026-03-04 00:42:37Z", "m.alkhateeb@bankaletihad.com", "94.249.51.64", "Success"],
], [1.5, 2.5, 1.5, 1.0])

body(
    "94.249.51.64 resolves to Jordan Telecom / Orange Jordan. Single legitimate console login from a BAE "
    "employee. No failed login attempts or suspicious IAM activity detected.")

subsection("6.7 Additional Negative Findings")
body("The following threat techniques were explicitly searched with ZERO results:")

table(["Search Term", "Result"], [
    ["Backdoor tools (webshell, reverse shell, C2 beacons)", "Not detected"],
    ["Remote access tools (AnyDesk, TeamViewer, ngrok, Cobalt Strike)", "Not found"],
    ["Port scanning signatures", "No evidence"],
    ["Lateral movement (pass-the-hash, mimikatz, credential dumping)", "No indicators"],
    ["Aqaba (claimed attack location)", "Only website content images"],
    ["Handala malware filenames (F5UPDATER.exe, Hatef.exe, etc.)", "Not detected"],
], [3.5, 3.0])


# ═══════════════════════════════════════════════════════════════
# 7. MITRE ATT&CK MAPPING
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
section_heading("7. MITRE ATT&CK MAPPING")

body(
    "The following MITRE ATT&CK techniques have been mapped based on observed and claimed threat actor behavior, "
    "Handala Hack Team known TTPs, and investigation-correlated findings:")

table(["Tactic ID", "Tactic", "Technique ID", "Technique", "Observation"], [
    ["TA0001", "Initial Access", "T1566.001", "Spear Phishing Attachment", "Hebrew/Arabic phishing emails with F5 update lures, RedAlert APK"],
    ["TA0001", "Initial Access", "T1190", "Exploit Public-Facing App", "Management system exploitation; WAF db.sql probe detected"],
    ["TA0001", "Initial Access", "T1199", "Trusted Relationship", "Compromised third-party contractor/vendor system"],
    ["TA0002", "Execution", "T1059.010", "AutoHotkey & AutoIT", "Handala uses AutoIt-based shellcode injection (Naples.pif)"],
    ["TA0002", "Execution", "T1059", "Command & Scripting Interpreter", "Carroll.cmd batch script, update.sh Linux wiper"],
    ["TA0003", "Persistence", "T1505.003", "Web Shell", "Backdoor access points on core banking servers (claimed)"],
    ["TA0003", "Persistence", "T1098", "Account Manipulation", "Maintaining persistent access post-detection"],
    ["TA0004", "Privilege Escalation", "T1068", "Exploitation for Privilege Escalation", "BYOVD via ListOpenedFileDrv_32.sys vulnerable driver"],
    ["TA0005", "Defense Evasion", "T1027", "Obfuscated Files or Information", "Encrypted wiper payloads, obfuscated AutoIt scripts"],
    ["TA0005", "Defense Evasion", "T1497.003", "Time-Based Evasion", "Sleep delays to evade sandbox analysis"],
    ["TA0005", "Defense Evasion", "T1562.004", "Disable/Modify Firewall", "Circumventing port blocking and access restrictions"],
    ["TA0006", "Credential Access", "T1110", "Brute Force", "SIP Register brute force on VoIP gateway (confirmed)"],
    ["TA0007", "Discovery", "T1046", "Network Service Scanning", "Live monitoring and enumeration of banking servers"],
    ["TA0007", "Discovery", "T1087", "Account Discovery", "/api/v1/users enumeration attempt (confirmed)"],
    ["TA0007", "Discovery", "T1590", "Gather Victim Network Info", "Reconnaissance via icanhazip.com for victim IP"],
    ["TA0009", "Collection", "T1005", "Data from Local System", "Screenshot capture of server dashboards and SCADA panels"],
    ["TA0010", "Exfiltration", "T1020", "Automated Exfiltration", "Telegram bot C2 + AWS S3 + Storj for data theft"],
    ["TA0011", "Command & Control", "T1071", "App Layer Protocol", "HTTPS on port 2515 to 31.192.237.207; Telegram bot API"],
    ["TA0040", "Impact", "T1561.002", "Disk Structure Wipe", "Hatef.exe (Windows) and Hamsa/update.sh (Linux) wipers"],
    ["TA0040", "Impact", "T1489", "Service Stop", "Solar inverter grid disconnection (Aqaba SEZ)"],
], [0.5, 0.8, 0.6, 1.2, 3.0])


# ═══════════════════════════════════════════════════════════════
# 8. DEFENSIVE RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════

section_heading("8. DEFENSIVE RECOMMENDATIONS")

subsection("8.1 Immediate Actions (0\u201324 Hours)")
for item in [
    "Conduct emergency audit of all third-party management systems deployed across banking infrastructure, with priority on contractor-deployed platforms.",
    "Implement emergency firewall rules (see Section 9) to block identified IOC IP ranges, Handala C2 (31.192.237.207), and non-standard ports (TCP/2515, TCP/4443, TCP/8443, TCP/9090).",
    "Block all malicious domains: api.ra-backup.com, shirideitch.com, sjc1.vultrobjects.com, web27.info, handala-hack.to.",
    "Deploy EDR hunting rules for Handala malware filenames: F5UPDATER.exe, Hatef.exe, Handala.exe, senvarservice-DC.exe, Naples.pif, OpenFileFinder.dll.",
    "Hunt for all 16 SHA256 malware hashes across all endpoints (see Section 5.7).",
    "Isolate and forensically image any systems associated with the compromised management platform.",
    "Engage Jordan CERT (JO-CERT) and relevant financial sector ISACs for coordinated response.",
    "Review and rotate all administrative credentials associated with banking core systems and SCADA/ICS platforms.",
    "Verify BELECTRIC solar system status \u2014 physical/network audit of Aqaba solar monitoring infrastructure.",
    "Validate whether examproctor.bankaletihad.com exists and assess its current patch status.",
    "Block external recon IPs: 178.79.148.111 (db.sql probe), 165.154.24.96 (HK phishing probe).",
    "Monitor outbound connections to api.telegram.org from non-standard hosts (Handala Telegram bot C2).",
]: bullet(item)

subsection("8.2 Short-Term Actions (1\u20137 Days)")
for item in [
    "Conduct comprehensive network traffic analysis for communications to/from identified IP ranges over the past 90 days.",
    "Perform full vulnerability assessment on all BELECTRIC and solar monitoring SCADA/ICS systems in the Aqaba Special Economic Zone.",
    "Deploy enhanced monitoring (EDR/NDR) on all banking core servers with specific detection rules for the TTPs identified in Section 7.",
    "Engage third-party incident response firm to conduct independent forensic analysis of the compromised management system.",
    "Review all contractor access and implement zero-trust network segmentation for third-party managed systems.",
    "Monitor Handala leak sites daily: handala-hack.to and .onion site for BELECTRIC/BAE data publication (avg 23.8 day delay).",
    "Investigate SIP brute force from internal branch IPs 10.0.6.69 and 10.0.6.71 targeting VoIP gateway.",
    "Confirm whether Shodan-reported IPs 13.41.61.138 and 23.21.102.205 are BAE assets.",
    "Hunt for BYOVD indicators: ListOpenedFileDrv_32.sys vulnerable driver loading across all Windows endpoints.",
]: bullet(item)

subsection("8.3 Long-Term Actions (1\u20133 Months)")
for item in [
    "Implement supply chain security assessment framework for all technology vendors and contractors.",
    "Deploy network micro-segmentation between IT and OT (SCADA/ICS) environments.",
    "Establish continuous threat monitoring for Iranian-nexus threat actors (Handala, CyberAv3ngers) targeting Jordanian financial and energy sectors.",
    "Conduct tabletop exercise simulating supply chain compromise scenario based on this incident.",
    "Review and enhance SCADA/ICS security posture across all critical infrastructure assets.",
    "Implement canary tokens and deception technology on high-value assets to detect persistent access.",
]: bullet(item)


# ═══════════════════════════════════════════════════════════════
# 9. FIREWALL BLOCK RULES
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
section_heading("9. FIREWALL BLOCK RULES (Multi-Format Export)")

code_block("iptables (Linux)", [
    "# Block Handala C2 infrastructure",
    "iptables -A INPUT -s 31.192.237.207 -j DROP",
    "iptables -A OUTPUT -d 31.192.237.207 -j DROP",
    "",
    "# Block suspected C2 infrastructure and detected recon IPs",
    "iptables -A INPUT -s 185.173.35.0/24 -j DROP",
    "iptables -A INPUT -s 91.243.44.0/24 -j DROP",
    "iptables -A INPUT -s 209.74.87.100 -j DROP",
    "iptables -A INPUT -s 157.20.182.49 -j DROP",
    "iptables -A INPUT -s 185.236.25.11 -j DROP",
    "iptables -A INPUT -s 38.180.239.161 -j DROP",
    "iptables -A INPUT -s 92.243.65.243 -j DROP",
    "iptables -A INPUT -s 178.79.148.111 -j DROP",
    "iptables -A INPUT -s 165.154.24.96 -j DROP",
    "iptables -A OUTPUT -d 185.173.35.0/24 -j DROP",
    "iptables -A OUTPUT -d 91.243.44.0/24 -j DROP",
    "iptables -A OUTPUT -d 209.74.87.100 -j DROP",
    "iptables -A OUTPUT -d 157.20.182.49 -j DROP",
    "iptables -A OUTPUT -d 185.236.25.11 -j DROP",
    "iptables -A OUTPUT -d 38.180.239.161 -j DROP",
    "iptables -A OUTPUT -d 92.243.65.243 -j DROP",
    "",
    "# Block non-standard backdoor ports observed",
    "iptables -A INPUT -p tcp --dport 2515 -j DROP",
    "iptables -A INPUT -p tcp --dport 4443 -j DROP",
    "iptables -A INPUT -p tcp --dport 8443 -j DROP",
    "iptables -A INPUT -p tcp --dport 9090 -j DROP",
    "",
    "# Block malicious domains (DNS-level or via resolved IPs)",
    "# api.ra-backup.com, shirideitch.com, sjc1.vultrobjects.com, web27.info",
    "# handala-hack.to, handala.to",
])

code_block("Palo Alto (PAN-OS)", [
    "# Block Handala C2 + threat infrastructure",
    "set rulebase security rules Block-Handala-C2 from any to any destination-address 31.192.237.207 action deny",
    "set rulebase security rules Block-Handala-C2 application any service any log-start yes log-end yes",
    "",
    "set rulebase security rules Block-BankAlEtihad-IOC from any to any source-address [185.173.35.0/24 91.243.44.0/24 209.74.87.100 157.20.182.49 185.236.25.11 38.180.239.161 92.243.65.243 178.79.148.111 165.154.24.96] action deny",
    "set rulebase security rules Block-BankAlEtihad-IOC application any service any log-start yes log-end yes",
    "",
    "set rulebase security rules Block-Backdoor-Ports from any to any service [tcp/2515 tcp/4443 tcp/8443 tcp/9090] action deny",
    "",
    "# Block malicious domains",
    "set rulebase security rules Block-Malicious-Domains from any to any destination [api.ra-backup.com www.shirideitch.com sjc1.vultrobjects.com web27.info handala-hack.to handala.to] action deny",
])

code_block("Windows Firewall (PowerShell)", [
    "# Block Handala C2",
    'New-NetFirewallRule -DisplayName "Block-Handala-C2" -Direction Outbound -RemoteAddress 31.192.237.207 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-Handala-C2-In" -Direction Inbound -RemoteAddress 31.192.237.207 -Action Block',
    "",
    "# Block inbound from suspected C2 ranges and recon IPs",
    'New-NetFirewallRule -DisplayName "Block-IOC-185.173.35.0" -Direction Inbound -RemoteAddress 185.173.35.0/24 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-91.243.44.0" -Direction Inbound -RemoteAddress 91.243.44.0/24 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-209.74.87.100" -Direction Inbound -RemoteAddress 209.74.87.100 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-157.20.182.49" -Direction Inbound -RemoteAddress 157.20.182.49 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-185.236.25.11" -Direction Inbound -RemoteAddress 185.236.25.11 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-38.180.239.161" -Direction Inbound -RemoteAddress 38.180.239.161 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-92.243.65.243" -Direction Inbound -RemoteAddress 92.243.65.243 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-178.79.148.111" -Direction Inbound -RemoteAddress 178.79.148.111 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-IOC-165.154.24.96" -Direction Inbound -RemoteAddress 165.154.24.96 -Action Block',
    'New-NetFirewallRule -DisplayName "Block-Backdoor-Ports" -Direction Inbound -Protocol TCP -LocalPort 2515,4443,8443,9090 -Action Block',
])

code_block("Cisco IOS ACL", [
    "ip access-list extended BLOCK-BANKETIHAD-IOC",
    "  deny ip host 31.192.237.207 any log",
    "  deny ip any host 31.192.237.207 log",
    "  deny ip 185.173.35.0 0.0.0.255 any",
    "  deny ip 91.243.44.0 0.0.0.255 any",
    "  deny ip host 209.74.87.100 any",
    "  deny ip host 157.20.182.49 any",
    "  deny ip host 185.236.25.11 any",
    "  deny ip host 38.180.239.161 any",
    "  deny ip host 92.243.65.243 any",
    "  deny ip host 178.79.148.111 any",
    "  deny ip host 165.154.24.96 any",
    "  deny tcp any any eq 2515",
    "  deny tcp any any eq 4443",
    "  deny tcp any any eq 8443",
    "  deny tcp any any eq 9090",
    "  permit ip any any",
])

code_block("Snort/Suricata", [
    '# Handala C2 Server',
    'alert tcp $HOME_NET any -> 31.192.237.207 any (msg:"IOC - Handala C2 Server (CRITICAL)"; sid:2026030701; rev:1;)',
    'alert tcp 31.192.237.207 any -> $HOME_NET any (msg:"IOC - Handala C2 Inbound"; sid:2026030702; rev:1;)',
    '',
    '# Suspected C2 ranges',
    'alert tcp $HOME_NET any -> 185.173.35.0/24 any (msg:"IOC - Suspected C2 to Bank al Etihad threat actor"; sid:2026030703; rev:1;)',
    'alert tcp $HOME_NET any -> 91.243.44.0/24 any (msg:"IOC - Suspected C2 range associated with SCADA compromise"; sid:2026030704; rev:1;)',
    '',
    '# Threat infrastructure IPs',
    'alert tcp $HOME_NET any -> 209.74.87.100 any (msg:"IOC - Threat infrastructure 209.74.87.100"; sid:2026030705; rev:1;)',
    'alert tcp $HOME_NET any -> 157.20.182.49 any (msg:"IOC - Threat infrastructure 157.20.182.49"; sid:2026030706; rev:1;)',
    'alert tcp $HOME_NET any -> 185.236.25.11 any (msg:"IOC - Threat infrastructure 185.236.25.11"; sid:2026030707; rev:1;)',
    'alert tcp $HOME_NET any -> 38.180.239.161 any (msg:"IOC - Threat infrastructure 38.180.239.161"; sid:2026030708; rev:1;)',
    'alert tcp $HOME_NET any -> 92.243.65.243 any (msg:"IOC - Threat infrastructure 92.243.65.243"; sid:2026030709; rev:1;)',
    '',
    '# Backdoor ports + Handala C2 port',
    'alert tcp any any -> $HOME_NET [2515,4443,8443,9090] (msg:"IOC - Backdoor/C2 port access attempt"; sid:2026030710; rev:1;)',
    '',
    '# WAF-detected recon IPs',
    'alert tcp $HOME_NET any -> 178.79.148.111 any (msg:"IOC - Detected db.sql recon probe"; sid:2026030711; rev:1;)',
    'alert tcp $HOME_NET any -> 165.154.24.96 any (msg:"IOC - Detected phishing recon HK"; sid:2026030712; rev:1;)',
    '',
    '# Malicious domain DNS lookups',
    'alert dns $HOME_NET any -> any any (msg:"IOC - Handala C2 domain api.ra-backup.com"; dns.query; content:"api.ra-backup.com"; sid:2026030713; rev:1;)',
    'alert dns $HOME_NET any -> any any (msg:"IOC - Malware delivery domain shirideitch.com"; dns.query; content:"shirideitch.com"; sid:2026030714; rev:1;)',
    'alert dns $HOME_NET any -> any any (msg:"IOC - Wiper payload host sjc1.vultrobjects.com"; dns.query; content:"sjc1.vultrobjects.com"; sid:2026030715; rev:1;)',
    'alert dns $HOME_NET any -> any any (msg:"IOC - Suspected phishing domain web27.info"; dns.query; content:"web27.info"; sid:2026030716; rev:1;)',
    'alert dns $HOME_NET any -> any any (msg:"IOC - Handala leak site handala-hack.to"; dns.query; content:"handala-hack.to"; sid:2026030717; rev:1;)',
])

code_block("Sigma Rule", [
    "title: Bank al Etihad / Handala IOC - Suspicious Network Connection",
    "status: experimental",
    "description: Detects connections to Handala C2, threat infrastructure, and malicious domains",
    "logsource:",
    "    category: firewall",
    "detection:",
    "    selection_handala_c2:",
    "        dst_ip: '31.192.237.207'",
    "    selection_ip_range:",
    "        dst_ip|startswith:",
    "            - '185.173.35.'",
    "            - '91.243.44.'",
    "    selection_ip_single:",
    "        dst_ip:",
    "            - '209.74.87.100'",
    "            - '157.20.182.49'",
    "            - '185.236.25.11'",
    "            - '38.180.239.161'",
    "            - '92.243.65.243'",
    "            - '178.79.148.111'",
    "            - '165.154.24.96'",
    "    selection_port:",
    "        dst_port:",
    "            - 2515",
    "            - 4443",
    "            - 8443",
    "            - 9090",
    "    condition: selection_handala_c2 or selection_ip_range or selection_ip_single or selection_port",
    "    level: critical",
    "    tags:",
    "        - attack.command_and_control",
    "        - attack.persistence",
    "        - attack.impact",
    "",
    "---",
    "",
    "title: Handala Malware Filename Detection",
    "status: experimental",
    "description: Detects Handala Hack Team malware filenames on endpoints",
    "logsource:",
    "    category: file_event",
    "    product: windows",
    "detection:",
    "    selection:",
    "        TargetFilename|endswith:",
    "            - '\\\\F5UPDATER.exe'",
    "            - '\\\\Hatef.exe'",
    "            - '\\\\Handala.exe'",
    "            - '\\\\senvarservice-DC.exe'",
    "            - '\\\\Naples.pif'",
    "            - '\\\\OpenFileFinder.dll'",
    "            - '\\\\ListOpenedFileDrv_32.sys'",
    "            - '\\\\Carroll.cmd'",
    "    condition: selection",
    "    level: critical",
    "    tags:",
    "        - attack.execution",
    "        - attack.impact",
    "        - attack.privilege_escalation",
])


# ═══════════════════════════════════════════════════════════════
# 10. CONSOLIDATED IOC SUMMARY
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
section_heading("10. CONSOLIDATED IOC SUMMARY")

subsection("Handala C2 & Threat Infrastructure")
table(["IP / Indicator", "Threat Type", "Source", "Confidence"], [
    ["31.192.237[.]207:2515", "Handala primary C2 server (Chelyabinsk, RU)", "Intezer/Splunk", "CRITICAL"],
    ["38.180.239[.]161", "Associated threat infrastructure", "Threat intel", "HIGH"],
    ["209.74.87[.]100", "Associated threat infrastructure", "Threat intel", "HIGH"],
    ["157.20.182[.]49", "Associated threat infrastructure", "Threat intel", "HIGH"],
    ["185.236.25[.]11", "Associated threat infrastructure", "Threat intel", "HIGH"],
    ["92.243.65[.]243", "Associated threat infrastructure", "Threat intel", "HIGH"],
    ["185.173.35[.]0/24", "Suspected C2 infrastructure", "Threat claim", "HIGH"],
    ["91.243.44[.]0/24", "SCADA/ICS callback range", "Threat claim", "MEDIUM"],
    ["TCP/2515, 4443, 8443, 9090", "C2 and backdoor ports", "Threat intel", "HIGH"],
], [1.8, 1.8, 1.5, 1.0])

subsection("Malicious Domains & URLs")
table(["Indicator", "Context", "Confidence"], [
    ["api[.]ra-backup[.]com", "C2 callback domain (RedAlert malware)", "CRITICAL"],
    ["shirideitch[.]com", "Malicious APK delivery domain", "HIGH"],
    ["sjc1[.]vultrobjects[.]com", "Linux wiper payload host", "CRITICAL"],
    ["web27[.]info", "Suspected malware staging / phishing infrastructure", "HIGH"],
    ["handala-hack[.]to", "Handala clearnet leak site (active)", "HIGH"],
    ["handala[.]to", "Handala clearnet leak site (down since 2025-06-26)", "HIGH"],
    ["hxxp://38[.]180[.]239[.]161", "Threat infrastructure HTTP endpoint", "HIGH"],
    ["hxxps://web27[.]info/0d894cce589ccd44", "Suspected malware staging", "HIGH"],
    ["hxxps://web27[.]info/b4cac8a7128d6507", "Suspected malware staging", "HIGH"],
    ["hxxps://web27[.]info/772b49c3c513b961", "Suspected malware staging", "HIGH"],
    ["hxxps://web27[.]info/b4f44ec4db01887aHOR", "Suspected malware staging", "HIGH"],
    ["hxxps://web27[.]info/ec714d155863dcd6", "Suspected malware staging", "HIGH"],
    ["hxxps://bit[.]ly/4tWJhQh", "Shortened URL \u2014 suspected malware delivery", "HIGH"],
], [2.5, 2.5, 1.0])

subsection("Malware Hashes (16 samples)")
table(["SHA256 (truncated)", "Description", "Severity"], [
    ["96dec6e07229...dcbc8", "update.zip \u2014 Phishing delivery archive", "CRITICAL"],
    ["fe07dca68f28...30bd2", "F5UPDATER.exe \u2014 .NET loader (signed)", "CRITICAL"],
    ["e28085e8d64b...3af35", "Hatef.exe \u2014 Windows wiper", "CRITICAL"],
    ["454e6d3782f2...9567", "Handala.exe \u2014 Delphi loader", "CRITICAL"],
    ["6f79c0e0e1aa...40ad", "update.sh \u2014 Linux wiper (Hamsa)", "CRITICAL"],
    ["83651b058966...8b72", "RedAlert.apk \u2014 Android malware", "HIGH"],
    ["7d9fb236607e...64c9", "Associated malware sample", "HIGH"],
], [2.0, 2.8, 1.0])
body("See Section 5.7 for complete hash list with all 16 SHA256 values.")

subsection("Confirmed External Threat IPs (WAF)")
table(["IP", "Threat Type", "Source", "Confidence"], [
    ["178.79.148.111", "db.sql database dump probe (UK)", "WAF logs", "HIGH"],
    ["165.154.24.96", "Phishing recon probe (Hong Kong)", "WAF logs", "HIGH"],
    ["194.238.28.128", "Security researcher recon", "WAF logs", "MEDIUM"],
], [1.8, 2.2, 1.5, 1.0])

subsection("Confirmed Internal Anomalies")
table(["IP", "Activity", "Risk Level"], [
    ["10.0.6.69", "SIP brute force against VoIP gateway (172.17.30.1:5060)", "MEDIUM"],
    ["10.0.6.71", "SIP brute force against VoIP gateway (172.17.30.1:5060)", "MEDIUM"],
], [1.5, 3.5, 1.5])

subsection("Infrastructure Assets")
table(["Asset", "IP / FQDN", "Status"], [
    ["Primary NGFW", "Primary-NGFW.bankaletihad.com", "Active, dropping threats"],
    ["Batch DB Server", "Batch-ADP-PROD.BankAlEtihad.Com (10.12.11.233)", "Active, normal operation"],
    ["S3 Bucket", "prod-bae-blog-frontend", "Probed, access denied (secure)"],
    ["AWS Account", "829210555596 (eu-west-2)", "Console login from Jordan \u2014 legitimate"],
], [1.5, 3.0, 2.0])


# ═══════════════════════════════════════════════════════════════
# 11. APPENDIX
# ═══════════════════════════════════════════════════════════════

section_heading("11. APPENDIX")

subsection("11.1 Threat Communique Translation (Farsi \u2192 English)")
body(
    "\u201cBank al Etihad (Bank al Etihad), one of the prominent and long-standing commercial banks of Jordan, "
    "has been the target of our successful penetration. Since yesterday, we have taken control of parts of this "
    "bank\u2019s technical infrastructure and have begun manipulation and persistent access operations into their "
    "primary servers.\u201d")
body(
    "\u201cThe initial penetration occurred through a management system deployed in Jordan that inadvertently "
    "created multiple backdoor access points to the bank\u2019s core servers. We sincerely thank the contractor "
    "responsible for deploying this system.\u201d")
body("[Truncated \u2014 additional content references solar energy systems in the Aqaba Special Economic Zone]")

subsection("11.2 BELECTRIC Dashboard Analysis")
body(
    "The compromised BELECTRIC solar monitoring dashboard reveals a solar installation with the following "
    "observed parameters at time of compromise: Monthly Energy output of 22.02 MWh, Yearly Energy of "
    "253.48 MWh, system availability of 133.85% (monthly, anomalous reading), ambient temperature of "
    "\u221240\u00b0C (indicative of sensor manipulation). All five inverters in OFF/Grid Loss state.")

subsection("11.3 BELECTRIC Gulf Ltd \u2014 Jordan Operations")
body(
    "BELECTRIC Gulf Ltd operates 7 solar plants in Jordan with approximately 100 MW total capacity. "
    "The South Amman Solar Plant alone comprises 46 MWp across 395,000 panels. A centralized SCADA control "
    "center monitors all plants remotely in real-time. BELECTRIC was previously listed as a Handala victim "
    "under the name \"BLEnergy\" on 2024-07-23 (confirmed via ransomware.live). The average delay between "
    "Handala attack execution and public claim posting is approximately 23.8 days \u2014 a public post "
    "regarding the March 2026 breach may be imminent.")

subsection("11.4 Handala Hack Team \u2014 Registration Identities")
body(
    "Known registration identities associated with Handala infrastructure: WordPress admin username \"vie6c\", "
    "domain registration name \"Roxie Collins\". These identities are used across Handala\u2019s leak site "
    "infrastructure and may appear in future domain registrations or web platform setups.")

subsection("11.5 Disclaimer")
body(
    "This report is produced by ScanWave.ai for threat intelligence purposes. The IOCs and analysis contained "
    "herein are based on open-source intelligence, threat actor communiques, visual evidence provided in the "
    "original disclosure, and log analysis from the Bank al Etihad deployment. IP ranges included "
    "as IOCs are assessed based on contextual analysis and may require further validation.")

# ── Full Legal Disclaimer (last page) ──
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Spacing before disclaimer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(24)

# Horizontal rule
rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(6)
rule_run = rule.add_run("_" * 95)
rule_run.font.size = Pt(7)
rule_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# DISCLAIMER heading
disc_heading = doc.add_paragraph()
disc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
disc_heading.paragraph_format.space_before = Pt(8)
disc_heading.paragraph_format.space_after = Pt(4)
dh_run = disc_heading.add_run("DISCLAIMER")
dh_run.font.size = Pt(9)
dh_run.font.bold = True
dh_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

# Disclaimer text
disclaimer_text = (
    "All information contained in this report was collected exclusively through publicly available "
    "Open Source Intelligence (OSINT), social media monitoring, and third-party threat intelligence feeds. "
    "No unauthorized access to any systems was performed during the preparation of this report. "
    "While this report contains security recommendations, ScanWave Cybersecurity does not guarantee the "
    "accuracy, completeness, or reliability of any information presented herein. This report is provided "
    'on an "as-is" basis for informational and defensive security purposes only and does not constitute '
    "legal or professional advice. ScanWave Cybersecurity, its employees, and its affiliates shall not be "
    "held liable for any damages, losses, or consequences \u2014 whether direct, indirect, incidental, or "
    "consequential \u2014 arising from the use of, reliance on, or actions taken based on the information "
    "contained in this report. Recipients assume full responsibility for any decisions made based on this material."
)
disc_para = doc.add_paragraph()
disc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
disc_para.paragraph_format.space_after = Pt(0)
d_run = disc_para.add_run(disclaimer_text)
d_run.font.size = Pt(7.5)
d_run.font.italic = True
d_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Save ──
output_path = os.path.join(BASE, "ScanWave_BAE_IOC_Combined_v3.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
